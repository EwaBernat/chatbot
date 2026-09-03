# -*- coding: utf-8 -*-
"""KPOF: HTML -> DOCX z zachowaniem stylu i kolorystyki EduPlaner 2026 / PCTP."""
import re, sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Mm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from dom import parse_html, Node, parse_style
from wordkit import *

# ---------------------------------------------------------------- paleta
V = {
    'purple': '2D1B69', 'purple-soft': 'EFEAF9', 'purple-line': 'D9D0F0', 'purple-mid': '5A4A94',
    'orange': 'E8450A', 'orange-soft': 'FDECE4', 'orange-line': 'F3CDBD',
    'ink': '2B2733', 'muted': '6F6A7D', 'paper': 'FFFFFF', 'hair': 'E4E1EC', 'zebra': 'FAF7F2',
    'lvl1': '1F8A5B', 'lvl1-bg': 'EAF6F0', 'lvl1-line': 'BFE3D1',
    'lvl2': 'C8811B', 'lvl2-bg': 'FBF3E3', 'lvl2-line': 'EED6A8',
    'lvl3': 'C0392B', 'lvl3-bg': 'FBEBE9', 'lvl3-line': 'F0C3BD',
}
FONT = 'Segoe UI'
SCRIPT_FONT = 'Segoe Script'
DOT = 'B7ADD6'

CW = 10658          # szerokość kolumny tekstu (188 mm) w twipach
CW_PX = CW / PX     # ~710 px

RATE_CH = {'1': '①', '2': '②', '3': '③', '4': '④', '5': '⑤', 'N': 'Ⓝ'}
RATE_COL = {'1': 'D93B30', '2': 'D9542F', '3': 'B3891A', '4': '5E9436', '5': '2E9E52', 'N': '8B8698'}
FILLED = {'1': '❶', '2': '❷', '3': '❸', '4': '❹', '5': '❺'}


def color(v, default=None):
    if not v:
        return default
    v = v.strip()
    m = re.match(r'var\(--([a-z0-9-]+)\)', v)
    if m:
        return V.get(m.group(1), default)
    if v.startswith('#'):
        v = v[1:]
        if len(v) == 3:
            v = ''.join(c * 2 for c in v)
        return v.upper()
    named = {'#fff': 'FFFFFF', 'white': 'FFFFFF', 'black': '000000'}
    return named.get(v.lower(), default)


def sizepx(v, default=None):
    if not v:
        return default
    m = re.match(r'([\d.]+)px', v.strip())
    return float(m.group(1)) if m else default


# ---------------------------------------------------------------- kontekst stylu
class Ctx(dict):
    def sub(self, **kw):
        c = Ctx(self)
        c.update(kw)
        return c


BASE = Ctx(size=9.0, color=V['ink'], bold=False, italic=False, caps=False,
           font=FONT, align=None, roomy=False)

# rozmiary px wg klas: (zwykłe, roomy)
CLS_PX = {
    'sec-note': (8.5, 10.5), 'box': (10, 12), 'meta': (8.7, 10.5), 'legal': (7.8, 9.5),
    'hint': (8.3, 8.3), 'obsline': (9.5, 9.5), 'blk': (9, 9), 'opt': (12, 12.5),
    'band': (8.5, 8.5), 'lvl': (9, 9), 'diag': (9.3, 11), 'rodo': (9, 12.5),
    'legend': (9, 9),
}


def ctx_for(node, ctx):
    """Aktualizuje kontekst na podstawie klas i stylu inline."""
    c = ctx
    for cl in node.cls:
        if cl in CLS_PX:
            px = CLS_PX[cl][1 if ctx['roomy'] else 0]
            c = c.sub(size=px2pt(px))
    if 'muted' in node.cls or 'mut' in node.cls or 'ex' in node.cls:
        c = c.sub(color=V['muted'])
    st = node.style
    if 'font-size' in st:
        px = sizepx(st['font-size'])
        if px:
            c = c.sub(size=px2pt(px))
    if 'color' in st:
        col = color(st['color'])
        if col:
            c = c.sub(color=col)
    if 'font-weight' in st:
        c = c.sub(bold=st['font-weight'] in ('700', '800', 'bold'))
    if 'font-style' in st:
        c = c.sub(italic=st['font-style'] == 'italic')
    if 'text-transform' in st:
        c = c.sub(caps=st['text-transform'] == 'uppercase')
    if 'line-height' in st:
        try:
            c = c.sub(lh=float(st['line-height']))
        except ValueError:
            pass
    if 'text-align' in st:
        c = c.sub(align={'center': AL.CENTER, 'right': AL.RIGHT,
                         'justify': AL.JUSTIFY, 'left': AL.LEFT}.get(st['text-align']))
    return c


# ---------------------------------------------------------------- sink
class Sink:
    """Cel renderowania: dokument albo komórka tabeli (z recyklingiem 1. akapitu)."""
    def __init__(self, container, cell=False):
        self.c = container
        self.pending = container.paragraphs[0] if cell and container.paragraphs else None
        self.cell = cell
    def par(self):
        if self.pending is not None:
            p, self.pending = self.pending, None
            return p
        return self.c.add_paragraph()
    def table(self, rows, cols):
        if self.pending is not None:
            # zostaw pusty, mikroskopijny akapit przed tabelą
            tune(self.pending, 0, 0, line=1, exact=True)
            run(self.pending, '', size=1)
            self.pending = None
        t = self.c.add_table(rows, cols) if not self.cell else self.c.add_table(rows, cols)
        t.style = None
        table_borders(t, top=None, left=None, bottom=None, right=None, insideH=None, insideV=None)
        cell_margins(t, 0, 0, 0, 0)
        return t
    def after_table(self):
        """Komórka musi kończyć się akapitem."""
        if self.cell:
            p = self.c.add_paragraph()
            tune(p, 0, 0, line=1, exact=True)
            run(p, '', size=1)

    def spacer(self, pt_h):
        p = self.par()
        tune(p, 0, 0, line=pt_h, exact=True)
        run(p, '', size=1)
        return p


def new_table(sink, rows, cols, widths):
    t = sink.table(rows, cols)
    set_widths(t, widths)
    return t


# ---------------------------------------------------------------- runy inline
def emit_inline(par, node, ctx):
    for ch in node.children:
        if ch.tag == '#text':
            txt = ch.text
            if txt is None:
                continue
            txt = txt.replace('\n', ' ')
            if not txt:
                continue
            run(par, txt, size=ctx['size'], bold=ctx['bold'], italic=ctx['italic'],
                color=ctx['color'], font=ctx['font'], caps=ctx['caps'])
        elif ch.tag == 'br':
            par.add_run().add_break()
        elif ch.tag in ('b', 'strong'):
            emit_inline(par, ch, ctx_for(ch, ctx.sub(bold=True,
                        color=ctx.get('bcolor', ctx['color']))))
        elif ch.tag in ('i', 'em'):
            emit_inline(par, ch, ctx_for(ch, ctx.sub(italic=True)))
        elif ch.tag in ('span', 'a', 'small', 'div', 'label'):
            emit_span(par, ch, ctx)
        else:
            emit_inline(par, ch, ctx)


NB = ' '

def dotted(par, ctx, px_width, size=None):
    n = max(4, int(px_width / 3.0))
    run(par, ' ' * n, size=size or ctx['size'], color=DOT, underline='dotted', font=ctx['font'])


def emit_span(par, node, ctx):
    cls = node.cls
    st = node.style
    c = ctx_for(node, ctx)

    if 'chk' in cls:
        run(par, '☑ ' if 'on' in cls else '☐ ', size=max(c['size'] + 1.5, 9),
            color=V['orange'] if 'on' in cls else '8A85A0', bold=True)
        return
    if 'rate' in cls:
        for rc in node.children:
            if rc.tag == '#text':
                continue
            v = rc.attrs.get('data-v', rc.all_text().strip())
            run(par, RATE_CH.get(v, v), size=13.5, bold=False,
                color=RATE_COL.get(v, V['ink']), font='Segoe UI Symbol')
            run(par, ' ', size=9)
        return
    if 'sdot' in cls:
        v = node.all_text().strip()
        w = sizepx(st.get('width'), 15)
        bg = color(st.get('background'), V['purple'])
        if v in FILLED:
            run(par, FILLED[v], size=max(9, w * 0.78), color=bg, font='Segoe UI Symbol')
        else:
            run(par, ' %s ' % v, size=max(7, w * 0.5), bold=True, color='FFFFFF', fill=bg)
        run(par, ' ', size=c['size'])
        return
    if 'dline' in cls or 'blank' in cls:
        w = sizepx(st.get('min-width'), 60)
        dotted(par, c, w)
        return
    if 'code' in cls:
        run(par, node.all_text().strip(), size=px2pt(8), bold=True,
            color=V['orange'], fill=V['orange-soft'])
        return
    if 'nowe' in cls:
        run(par, NB + '◆', size=px2pt(7.5), bold=True, color=V['orange'],
            fill='FBF3DF', font='Segoe UI Symbol')
        run(par, ' ' + node.all_text().strip() + NB, size=px2pt(8), bold=True,
            color='9A7412', fill='FBF3DF')
        return
    if 'lvpill' in cls or 'pill' in cls or 'qlvl' in cls:
        bg = V['purple'] if ('p' in cls or 'pill' not in cls) else V['orange']
        if 'lv1' in cls:
            bg = V['lvl1']
        elif 'lv2' in cls:
            bg = V['lvl2']
        elif 'lv3' in cls:
            bg = V['lvl3']
        elif 'qlvl' in cls:
            bg = 'C9C3D6'
        elif 'pill' in cls:
            bg = V['purple'] if 'p' in cls else V['orange']
        txt = node.all_text().strip()
        run(par, '  ' + (txt if txt else ' ' * 14) + '  ', size=c['size'],
            bold=True, color='FFFFFF', fill=bg)
        return
    if 'sig' in cls:
        run(par, node.all_text().strip(), size=px2pt(15), color='77737F', font=SCRIPT_FONT)
        return
    if 'rn' in cls or ('border-radius' in st and st.get('border-radius') == '50%' and 'background' in st):
        bg = color(st.get('background'), V['orange'])
        run(par, ' %s ' % node.all_text().strip(), size=max(6.5, c['size']),
            bold=True, color=color(st.get('color'), 'FFFFFF'), fill=bg)
        run(par, ' ', size=c['size'])
        return
    if 'background' in st and st['background'] not in ('none', 'transparent'):
        bg = color(st['background'], None)
        txt = node.all_text().strip()
        h = sizepx(st.get('height'), None)
        if bg and not txt and h is not None and h <= 3:
            w = sizepx(st.get('width'), 40)
            run(par, '─' * max(3, int(w / 6.0)), size=c['size'], color=bg)
            return
        if bg and not txt:
            n = max(6, int(sizepx(st.get('min-width'), 60) / 4.0))
            run(par, NB * n, size=c['size'], fill=bg,
                underline=None if bg == 'FFFFFF' else 'dotted', color=DOT)
            return
        if bg and bg != 'FFFFFF':
            run(par, NB + txt + NB, size=c['size'], bold=c['bold'],
                color=c['color'] if color(st.get('color')) is None else color(st['color']),
                fill=bg)
            return
    emit_inline(par, node, c)


# ---------------------------------------------------------------- pudełka
def box(sink, node, ctx, fill, border, size=None, pad_l=110, pad_r=110,
        before=3, after=0, left_accent=None, min_lines=0):
    """Ramka jako pojedyncza komórka tabeli (pozwala na blokowe dzieci)."""
    t = new_table(sink, 1, 1, [CW - 0])
    cell = t.cell(0, 0)
    cell_shade(cell, fill)
    spec = dict(top=(6, border), left=(6, border), bottom=(6, border), right=(6, border))
    if left_accent:
        spec['left'] = (24, left_accent)
    cell_borders(cell, **spec)
    cell_margins(t, 55, pad_l, 55, pad_r)
    table_indent(t, 0)
    inner = Sink(cell, cell=True)
    c = ctx if size is None else ctx.sub(size=size)
    render_children(inner, node, ctx_for(node, c))
    for _ in range(min_lines):
        p = inner.par()
        tune(p, 0, 0)
        run(p, '', size=c['size'])
    return t


# ---------------------------------------------------------------- komponenty
def comp_sec(sink, node, ctx):
    num = ''
    title = ''
    badge = None
    for ch in node.children:
        if ch.tag == '#text':
            continue
        if ch.has('n'):
            num = ch.all_text().strip()
        elif ch.tag == 'h2':
            title = ch.all_text().strip()
        elif ch.has('nowe'):
            badge = ch
    p = sink.par()
    tune(p, before=7 if ctx['roomy'] else 6, after=1.5, line=1.0, keep=True)
    par_borders(p, bottom=(4, V['hair']))
    if num:
        bg = V['purple'] if 'p' in [c for ch in node.children if ch.tag != '#text' and ch.has('n') for c in ch.cls] else V['orange']
        run(p, ' ' + num + ' ', size=9, bold=True, color='FFFFFF', fill=bg)
        run(p, '  ', size=9)
    run(p, title.upper(), size=px2pt(13.5 if ctx['roomy'] else 12.5), bold=True,
        color=V['purple'], spacing=12)
    if badge is not None:
        run(p, '   ', size=7)
        emit_span(p, badge, ctx)
    return p


def comp_head(sink, page_sub, pill_text, pill_purple, tag_text):
    t = new_table(sink, 1, 2, [int(CW * 0.62), int(CW * 0.38)])
    left, right = t.cell(0, 0), t.cell(0, 1)
    p = Sink(left, cell=True).par()
    tune(p, 0, 0, line=1.0)
    run(p, ' PCTP ', size=7, bold=True, color='FFFFFF', fill=V['purple'])
    run(p, '  ', size=8)
    run(p, 'EduPlaner 2026', size=px2pt(16), bold=True, color=V['purple'])
    p2 = left.add_paragraph()
    tune(p2, 0, 0, line=1.0)
    run(p2, page_sub.upper(), size=px2pt(8.5), bold=True, color=V['muted'], spacing=8)
    q = Sink(right, cell=True).par()
    tune(q, 0, 0, line=1.0, align=AL.RIGHT)
    run(q, '  ' + pill_text + '  ', size=px2pt(10), bold=True, color='FFFFFF',
        fill=V['purple'] if pill_purple else V['orange'])
    q2 = right.add_paragraph()
    tune(q2, 1, 0, line=1.0, align=AL.RIGHT)
    run(q2, tag_text.upper(), size=px2pt(8), color=V['muted'], spacing=10)
    return t


def comp_rule(sink):
    t = new_table(sink, 1, 2, [int(CW * 0.55), int(CW * 0.45)])
    row_height(t.rows[0], 2.6)
    for i, fill in enumerate((V['purple'], V['orange'])):
        c = t.cell(0, i)
        cell_shade(c, fill)
        p = c.paragraphs[0]
        tune(p, 0, 0, line=2.6, exact=True)
        run(p, '', size=1)
    return t


def comp_student(sink, node, ctx):
    fields = [ch for ch in node.children if ch.tag != '#text' and ch.has('field')]
    flexes = []
    for f in fields:
        st = f.style
        fl = st.get('flex', '1')
        try:
            flexes.append(float(fl))
        except ValueError:
            flexes.append(1.0)
    tot = sum(flexes)
    gap = 60
    avail = CW - gap * (len(fields) - 1)
    widths = []
    for i, f in enumerate(flexes):
        widths.append(int(avail * f / tot))
    cols = len(fields) * 2 - 1
    w = []
    for i, x in enumerate(widths):
        if i:
            w.append(gap)
        w.append(x)
    t = new_table(sink, 1, cols, w)
    for i, f in enumerate(fields):
        cell = t.cell(0, i * 2)
        cell_shade(cell, V['purple-soft'])
        cell_borders(cell, top=(6, V['purple-line']), left=(6, V['purple-line']),
                     bottom=(6, V['purple-line']), right=(6, V['purple-line']))
        p = cell.paragraphs[0]
        tune(p, 1.5, 1.5, line=1.0, left=110, right=110)
        for ch in f.children:
            if ch.tag == '#text':
                continue
            if ch.has('lab'):
                run(p, ch.all_text().strip().upper() + '  ', size=px2pt(7.5), bold=True,
                    color=V['purple'], spacing=10)
            elif ch.has('blank'):
                dotted(p, ctx, 90, size=px2pt(11.5))
            elif ch.has('val'):
                st = ch.style
                run(p, ch.all_text().strip(), size=px2pt(11.5),
                    bold=st.get('font-weight') != '400',
                    color=color(st.get('color'), V['ink']))
    return t


def comp_blk(sink, node, ctx):
    p = sink.par()
    tune(p, before=4, after=2, line=1.0, keep=True)
    dot = V['orange']
    if 'g' in node.cls:
        dot = V['lvl1']
    elif 'pp' in node.cls:
        dot = V['purple']
    run(p, '■ ', size=px2pt(9), color=dot)
    run(p, node.all_text().strip().upper(), size=px2pt(9), bold=True,
        color=V['purple'], spacing=8)
    return p


def comp_grid(sink, node, ctx, ncols, cell_fn, gap=100, widths=None):
    items = [ch for ch in node.children if ch.tag != '#text']
    if not items:
        return None
    nrows = (len(items) + ncols - 1) // ncols
    colw = widths or [int((CW - gap * (ncols - 1)) / ncols)] * ncols
    w = []
    for i, x in enumerate(colw):
        if i:
            w.append(gap)
        w.append(x)
    t = new_table(sink, nrows, ncols * 2 - 1, w)
    for i, item in enumerate(items):
        r, c = divmod(i, ncols)
        cell_fn(t.cell(r, c * 2), item, ctx)
    return t


def fill_opt(cell, item, ctx):
    on = 'on' in item.cls
    cell_shade(cell, V['orange-soft'] if on else 'FFFFFF')
    bc = V['orange-line'] if on else V['hair']
    cell_borders(cell, top=(6, bc), left=(6, bc), bottom=(6, bc), right=(6, bc))
    p = cell.paragraphs[0]
    tune(p, 2.5, 2.5, line=1.05, left=100, right=100)
    emit_inline(p, item, ctx_for(item, ctx.sub(size=px2pt(12.5 if ctx['roomy'] else 12))))


def fill_fg(cell, item, ctx):
    cell_shade(cell, 'FAF9FD')
    cell_borders(cell, top=(6, V['hair']), left=(6, V['hair']),
                 bottom=(6, V['hair']), right=(6, V['hair']))
    s = Sink(cell, cell=True)
    for ch in item.children:
        if ch.tag == '#text':
            continue
        if ch.has('l'):
            p = s.par()
            tune(p, 2.5, 0.5, line=1.0, left=100, right=100)
            run(p, ch.all_text().strip().upper(), size=px2pt(7.5), bold=True,
                color=V['purple'], spacing=8)
        elif ch.has('blankline'):
            p = s.par()
            tune(p, 1, 3, line=px2pt(22 if ctx['roomy'] else 15), exact=True,
                 left=100, right=100)
            par_borders(p, bottom=(6, DOT, 'dotted'))
            run(p, '', size=8)
        elif ch.has('v'):
            p = s.par()
            tune(p, 1, 2.5, line=1.05, left=100, right=100)
            mut = 'mut' in ch.cls
            run(p, ch.all_text().strip(), size=px2pt(12.5 if ctx['roomy'] else 11.5),
                bold=not mut, color=V['muted'] if mut else V['ink'])
        else:
            p = s.par()
            tune(p, 1, 2, line=1.05, left=100, right=100)
            emit_inline(p, ch, ctx)


def fill_band(cell, item, ctx):
    kind = 'g'
    for k in ('g2', 'g', 'a', 'r'):
        if k in item.cls:
            kind = k
            break
    pal = {'g': ('EEF6EF', 'C9E2CD', '2E7D46'), 'g2': ('F1F7EA', 'D3E6BF', '4E7A2A'),
           'a': ('FBF6E8', 'ECDCAE', 'B3891A'), 'r': ('FDF1F0', 'F0CFCB', 'C0392B')}[kind]
    cell_shade(cell, pal[0])
    cell_borders(cell, top=(6, pal[1]), left=(6, pal[1]), bottom=(6, pal[1]), right=(6, pal[1]))
    s = Sink(cell, cell=True)
    p = s.par()
    tune(p, 2, 0, line=1.05, left=90, right=90)
    for ch in item.children:
        if ch.tag == '#text':
            continue
        if ch.has('bn'):
            run(p, ch.all_text().strip() + '  ', size=px2pt(11), bold=True, color=pal[2])
        elif ch.has('bt'):
            run(p, ch.all_text().strip(), size=px2pt(9), bold=True, color=pal[2])
        else:
            q = s.par()
            tune(q, 0.5, 2, line=1.05, left=90, right=90)
            run(q, ch.all_text().strip(), size=px2pt(8.5), color=V['ink'])


def fill_plain(cell, item, ctx):
    s = Sink(cell, cell=True)
    render_children(s, item, ctx_for(item, ctx), as_block=True)


# ---------------------------------------------------------------- tabele .tbl
def col_widths(tbl, total=CW_PX):
    ths = []
    thead = None
    for ch in tbl.children:
        if ch.tag == 'thead':
            thead = ch
    if thead is not None:
        for tr in thead.children:
            if tr.tag == 'tr':
                ths = [c for c in tr.children if c.tag in ('th', 'td')]
                break
    if not ths:
        for tr in tbl.find_all('tr'):
            ths = [c for c in tr.children if c.tag in ('th', 'td')]
            break
    n = len(ths) or 1
    specs = []
    for th in ths:
        w = th.style.get('width')
        if w and w.endswith('px'):
            specs.append(('px', float(w[:-2])))
        elif w and w.endswith('%'):
            specs.append(('pc', float(w[:-1])))
        else:
            specs.append(('auto', None))
    specs = [(k, (max(v, 58.0) if (k == 'px' and v is not None and v < 58) else v))
             for k, v in specs]
    used = 0.0
    for kind, val in specs:
        if kind == 'px':
            used += val
        elif kind == 'pc':
            used += total * val / 100.0
    autos = [i for i, s in enumerate(specs) if s[0] == 'auto']
    rest = max(total - used, total * 0.1)
    out = []
    for kind, val in specs:
        if kind == 'px':
            out.append(val)
        elif kind == 'pc':
            out.append(total * val / 100.0)
        else:
            out.append(rest / len(autos) if autos else total / n)
    scale = total / sum(out)
    return [int(x * scale * PX) for x in out]


def td_style(td, ctx, roomy):
    c = ctx.sub(size=px2pt(12.5 if roomy else 12))
    if td.has('itm'):
        c = c.sub(size=px2pt(11))
    if td.has('lp'):
        c = c.sub(color=V['orange'], bold=True, align=AL.CENTER)
    if td.has('b'):
        c = c.sub(bold=True, color=V['ink'])
    if td.has('mut') or td.has('ex'):
        c = c.sub(color=V['muted'])
    return ctx_for(td, c)


def comp_table(sink, tbl, ctx):
    roomy = ctx['roomy']
    widths = col_widths(tbl)
    head_rows, body_rows = [], []
    for part in tbl.children:
        if part.tag == 'thead':
            head_rows += [r for r in part.children if r.tag == 'tr']
        elif part.tag == 'tbody':
            body_rows += [r for r in part.children if r.tag == 'tr']
        elif part.tag == 'tr':
            body_rows.append(part)
    rows = head_rows + body_rows
    ncols = len(widths)
    is_qr = tbl.has('qrtbl')
    if not rows:
        return None
    t = new_table(sink, len(rows), ncols, widths)
    table_borders(t, top=(6, V['purple-line']), left=(6, V['purple-line']),
                  bottom=(6, V['purple-line']), right=(6, V['purple-line']),
                  insideH=(4, V['hair']), insideV=None)
    pad_v = 120 if (roomy and not is_qr) else 60
    if tbl.has('memtbl'):
        pad_v = 200
    cell_margins(t, pad_v, 120, pad_v, 120)
    table_indent(t, 0)
    for ri, tr in enumerate(rows):
        cells = [c for c in tr.children if c.tag in ('th', 'td')]
        is_head = ri < len(head_rows)
        if is_head:
            cant_split(t.rows[ri])
            repeat_header(t.rows[ri])
        bi = ri - len(head_rows)
        for ci, td in enumerate(cells[:ncols]):
            cell = t.cell(ri, ci)
            if is_head:
                cell_shade(cell, V['purple'])
                cell_borders(cell, top=None, left=None, bottom=None, right=None)
            elif bi % 2 == 1:
                cell_shade(cell, V['zebra'])
            if td.has('lp') or td.has('rate-c') or 'text-align:center' in td.attrs.get('style', ''):
                valign(cell, 'center')
            s = Sink(cell, cell=True)
            if is_head:
                p = s.par()
                tune(p, 0, 0, line=1.05)
                run(p, td.all_text().strip().upper(), size=px2pt(10.5 if roomy else 10),
                    bold=True, color='FFFFFF', spacing=8)
            else:
                c = td_style(td, ctx, roomy)
                blocks = [x for x in td.children if x.tag in ('div', 'p', 'table', 'ul')]
                if blocks:
                    render_children(s, td, c, as_block=True)
                else:
                    p = s.par()
                    tune(p, 0, 0, line=1.12, align=c['align'])
                    if td.has('code'):
                        run(p, td.all_text().strip(), size=px2pt(8), bold=True, color=V['orange'])
                    else:
                        emit_inline(p, td, c)
    return t


# ---------------------------------------------------------------- inne komponenty
def comp_areasum(sink, node, ctx):
    """Pasek wyniku obszaru: tabela, aby pola do wpisania nie znikały na końcu wiersza."""
    asl = ''
    abs_ = []
    for ch in node.children:
        if ch.tag == '#text':
            continue
        if ch.has('asl'):
            asl = ch.all_text().strip()
        elif ch.has('ab'):
            abs_.append(' '.join(ch.all_text().split()))
    # teksty typu "Σ  pkt", "ocenionych  / 5", "średnia"
    widths = [int(x * PX) for x in (150, 118, 150, 62, 80, 150)]
    t = new_table(sink, 1, 6, widths)
    table_borders(t, top=(6, V['purple-line']), left=(6, V['purple-line']),
                  bottom=(6, V['purple-line']), right=(6, V['purple-line']),
                  insideH=None, insideV=None)
    cell_margins(t, 70, 110, 70, 60)
    table_indent(t, 0)
    for i in range(6):
        cell_shade(t.cell(0, i), V['purple-soft'])
        valign(t.cell(0, i), 'center')

    def cellpar(i, align=None):
        p = t.cell(0, i).paragraphs[0]
        tune(p, 0, 0, line=1.1, align=align)
        return p

    p = cellpar(0)
    run(p, asl.upper(), size=px2pt(9), bold=True, color=V['purple'], spacing=8)
    # Σ ... pkt
    p = cellpar(1)
    run(p, 'Σ ', size=px2pt(10), color='463F5C')
    dotted(p, ctx, 40, size=px2pt(11))
    run(p, ' pkt', size=px2pt(10), color='463F5C')
    p = cellpar(2)
    run(p, 'ocenionych ', size=px2pt(10), color='463F5C')
    dotted(p, ctx, 40, size=px2pt(11))
    tot = '5'
    if abs_ and '/' in abs_[1]:
        tot = abs_[1].split('/')[-1].strip() or '5'
    run(p, ' / ' + tot, size=px2pt(10), color='463F5C')
    p = cellpar(3, AL.RIGHT)
    run(p, 'średnia', size=px2pt(10), color='463F5C')
    p = cellpar(4)
    dotted(p, ctx, 60, size=px2pt(11))
    cell_shade(t.cell(0, 4), 'FFFFFF')
    p = cellpar(5, AL.CENTER)
    cell_shade(t.cell(0, 5), 'C9C3D6')
    run(p, NB, size=px2pt(9.5))
    return t


def comp_qsumo(sink, node, ctx):
    """Pasek wyniku ogólnego (fioletowy)."""
    title, subtitle = 'Wynik ogólny', ''
    tail = []
    for ch in node.children:
        if ch.tag == '#text':
            continue
        txt = ' '.join(ch.all_text().split())
        st = ch.style
        if st.get('font-size', '').startswith('15'):
            title = txt
        elif st.get('font-size', '').startswith('9'):
            subtitle = txt
        else:
            tail.append(txt)
    widths = [int(x * PX) for x in (260, 105, 75, 120, 150)]
    t = new_table(sink, 1, 5, widths)
    cell_margins(t, 110, 130, 110, 60)
    table_indent(t, 0)
    for i in range(5):
        cell_shade(t.cell(0, i), V['purple'])
        valign(t.cell(0, i), 'center')

    def cellpar(i, align=None):
        p = t.cell(0, i).paragraphs[0]
        tune(p, 0, 0, line=1.15, align=align)
        return p

    p = cellpar(0)
    run(p, title, size=px2pt(15), bold=True, color='FFFFFF')
    if subtitle:
        q = t.cell(0, 0).add_paragraph()
        tune(q, 1, 0, line=1.1)
        run(q, subtitle, size=px2pt(9), color='DAD3EE')
    p = cellpar(1, AL.RIGHT)
    run(p, 'Średnia ogólna:', size=px2pt(11), color='FFFFFF')
    p = cellpar(2, AL.CENTER)
    cell_shade(t.cell(0, 2), 'FFFFFF')
    dotted(p, ctx, 55, size=px2pt(11))
    p = cellpar(3, AL.RIGHT)
    run(p, '/ 5', size=px2pt(11), bold=True, color='FFFFFF')
    run(p, '  ·  Kwalifikacja:', size=px2pt(11), color='FFFFFF')
    p = cellpar(4, AL.CENTER)
    cell_shade(t.cell(0, 4), 'FFFFFF')
    dotted(p, ctx, 110, size=px2pt(11))
    return t


def comp_sg(sink, node, ctx):
    items = [ch for ch in node.children if ch.tag != '#text']
    boxes = [i for i in items if i.has('sgb')]
    caps = [i for i in items if i.has('sgc')]
    n = max(len(boxes), 1)
    gap = 400
    colw = int((CW - gap * (n - 1)) / n)
    w = []
    for i in range(n):
        if i:
            w.append(gap)
        w.append(colw)
    t = new_table(sink, 2, n * 2 - 1, w)
    for i in range(n):
        c1 = t.cell(0, i * 2)
        p = c1.paragraphs[0]
        tune(p, 20, 0, line=1.0)
        par_borders(p, bottom=(10, DOT))
        run(p, '', size=9)
        c2 = t.cell(1, i * 2)
        q = c2.paragraphs[0]
        tune(q, 3, 0, line=1.0)
        if i < len(caps):
            run(q, caps[i].all_text().strip().upper(), size=px2pt(9.5), bold=True,
                color=V['muted'], spacing=8)
    return t


AREA_NAMES = [
    ('I', 'Uczenie się i stosowanie wiedzy'), ('II', 'Ogólne zadania i wymagania'),
    ('III', 'Porozumiewanie się'), ('IV', 'Poruszanie się'), ('V', 'Samoobsługa'),
    ('VI', 'Życie domowe'), ('VII', 'Wzajemne kontakty i związki międzyludzkie'),
    ('VIII', 'Główne obszary życia'), ('IX', 'Życie społeczne, lokalne i obywatelskie'),
]


def comp_chart(sink, node, ctx, area_names):
    """Karty wykresu: siatka do ręcznego naniesienia profilu + miejsce na mapę radarową."""
    gridcard, radarcard = None, None
    for card in node.find_all('div', 'chartcard'):
        if card.find_all('div', 'barchart'):
            gridcard = card
        elif card.find_all('div', 'radarchart'):
            radarcard = card
    gap = 170
    colw = [int((CW - gap) * 0.52), int((CW - gap) * 0.48)]
    t = new_table(sink, 1, 3, [colw[0], gap, colw[1]])
    for idx, card in ((0, gridcard), (2, radarcard)):
        if card is None:
            continue
        cell = t.cell(0, idx)
        cell_borders(cell, top=(6, V['purple-line']), left=(6, V['purple-line']),
                     bottom=(6, V['purple-line']), right=(6, V['purple-line']))
        s = Sink(cell, cell=True)
        ttl = card.find_all('div', 'chartttl')
        p = s.par()
        tune(p, 3, 2, line=1.05, left=110, right=110)
        run(p, '▌ ', size=9, color=V['orange'])
        run(p, (ttl[0].all_text().strip() if ttl else '').upper(), size=px2pt(10),
            bold=True, color=V['purple'], spacing=8)
        inner_w = colw[0 if idx == 0 else 1] - 260
        if idx == 0:
            lab = int(inner_w * 0.13)
            cw = int((inner_w - lab) / 9)
            g = s.table(6, 10)
            set_widths(g, [lab] + [cw] * 9)
            table_borders(g, top=(4, V['purple-line']), left=(4, V['purple-line']),
                          bottom=(4, V['purple-line']), right=(4, V['purple-line']),
                          insideH=(2, 'E9E5F2'), insideV=(2, 'E9E5F2'))
            cell_margins(g, 0, 20, 0, 20)
            table_indent(g, 110)
            for r, v in enumerate(['5', '4', '3', '2', '1']):
                row_height(g.rows[r], 19)
                pp = g.cell(r, 0).paragraphs[0]
                tune(pp, 0, 0, line=1.0, align=AL.CENTER)
                run(pp, v, size=6.5, bold=True, color='8A85A0')
                for cc in range(1, 10):
                    q = g.cell(r, cc).paragraphs[0]
                    tune(q, 0, 0, line=1.0)
                    run(q, '', size=6)
            row_height(g.rows[5], 12)
            pp = g.cell(5, 0).paragraphs[0]
            tune(pp, 0, 0, line=1.0)
            run(pp, '', size=6)
            for cc in range(1, 10):
                q = g.cell(5, cc).paragraphs[0]
                tune(q, 0.5, 0, line=1.0, align=AL.CENTER)
                run(q, 'd%d' % cc, size=6.5, bold=True, color=V['purple'])
                cell_borders(g.cell(5, cc), top=(4, V['purple-line']), left=None,
                             bottom=None, right=None)
            s.after_table()
            note = s.par()
            tune(note, 2, 3, line=1.15, left=110, right=110)
            run(note, 'Zaznacz kropką średnią każdego obszaru (d1–d9) i połącz linią — '
                      'profil funkcjonowania dziecka.', size=px2pt(8.5), color=V['muted'],
                italic=True)
        else:
            for _ in range(9):
                q = s.par()
                tune(q, 0, 0, line=11, exact=True, left=110, right=110)
                run(q, '', size=8)
            note = s.par()
            tune(note, 2, 3, line=1.15, left=110, right=110, align=AL.CENTER)
            run(note, 'Miejsce na mapę radarową\n', size=px2pt(9), bold=True, color=V['purple'])
            note2 = s.par()
            tune(note2, 0, 3, line=1.15, left=110, right=110, align=AL.CENTER)
            run(note2, '(generowana automatycznie w wersji interaktywnej HTML)',
                size=px2pt(8.5), color=V['muted'], italic=True)
    return t


def comp_qrtbl(sink, tbl, ctx):
    """Szybki odczyt — w druku 9 wierszy do wypełnienia."""
    widths = col_widths(tbl)
    t = new_table(sink, 10, 4, widths)
    table_borders(t, top=(6, V['purple-line']), left=(6, V['purple-line']),
                  bottom=(6, V['purple-line']), right=(6, V['purple-line']),
                  insideH=(4, V['hair']), insideV=None)
    cell_margins(t, 55, 120, 55, 120)
    table_indent(t, 0)
    heads = ['Obszar', 'Nazwa', 'Średnia', 'Kwalifikacja']
    for ci, h in enumerate(heads):
        cell = t.cell(0, ci)
        cell_shade(cell, V['purple'])
        p = cell.paragraphs[0]
        tune(p, 0, 0, line=1.05)
        run(p, h.upper(), size=px2pt(8.5), bold=True, color='FFFFFF', spacing=8)
    for i, (num, name) in enumerate(AREA_NAMES, start=1):
        if i % 2 == 0:
            for ci in range(4):
                cell_shade(t.cell(i, ci), V['zebra'])
        p = t.cell(i, 0).paragraphs[0]
        tune(p, 0, 0, line=1.05)
        run(p, num, size=px2pt(10), bold=True, color=V['orange'])
        p = t.cell(i, 1).paragraphs[0]
        tune(p, 0, 0, line=1.05)
        run(p, name, size=px2pt(9.5), color=V['ink'])
        for ci in (2, 3):
            p = t.cell(i, ci).paragraphs[0]
            tune(p, 0, 0, line=1.05)
            dotted(p, ctx, 55 if ci == 2 else 90, size=px2pt(9.5))
    return t


# ---------------------------------------------------------------- dispatcher
SKIP = ('head', 'rule', 'student', 'foot')
BLOCK_TAGS = {'div', 'p', 'table', 'section', 'h1', 'h2', 'h3', 'ul', 'ol', 'li', 'br'}


def is_block(n):
    if n.tag == '#text':
        return False
    if n.tag in ('span', 'b', 'i', 'em', 'strong', 'a', 'small'):
        return False
    if n.tag == 'div':
        st = n.style
        if st.get('display') in ('inline-flex', 'inline-block'):
            return False
    return n.tag in BLOCK_TAGS


def render_children(sink, node, ctx, as_block=False):
    buf = []
    def flush():
        if not buf:
            return
        if all(x.tag == '#text' and not x.text.strip() for x in buf):
            del buf[:]
            return
        p = sink.par()
        tune(p, 1.5, 1.5, line=1.25, align=ctx['align'])
        holder = Node('div')
        holder.children = list(buf)
        emit_inline(p, holder, ctx)
        del buf[:]
    for ch in node.children:
        if is_block(ch):
            flush()
            render_block(sink, ch, ctx)
        else:
            buf.append(ch)
    flush()


def render_block(sink, n, ctx):
    cls = n.cls
    c = ctx_for(n, ctx)
    if any(k in cls for k in SKIP):
        return
    if n.tag == 'table':
        if 'qrtbl' in cls:
            comp_qrtbl(sink, n, c)
        else:
            comp_table(sink, n, c)
        sink.after_table()
        return
    if 'sec' in cls and n.tag == 'div':
        comp_sec(sink, n, c)
        return
    if 'blk' in cls:
        comp_blk(sink, n, c)
        return
    if 'eyebrow' in cls:
        for ch in n.children:
            if ch.tag == '#text':
                continue
            p = sink.par()
            tune(p, 5, 2, line=1.1, align=AL.CENTER)
            if ch.has('e1'):
                run(p, '  ' + ch.all_text().strip().upper() + '  ', size=px2pt(9),
                    bold=True, color='FFFFFF', fill=V['orange'], spacing=12)
            else:
                run(p, ch.all_text().strip().upper(), size=px2pt(9), bold=True,
                    color=V['purple-mid'], spacing=15)
        return
    if 'areasum' in cls:
        comp_areasum(sink, n, c)
        return
    if 'qsum-o' in cls:
        comp_qsumo(sink, n, c)
        sink.after_table()
        return
    if 'kpchart' in cls:
        for ch in n.children:
            if ch.tag == '#text':
                continue
            if ch.has('chartgrid'):
                comp_chart(sink, ch, c, AREA_NAMES)
                sink.after_table()
            else:
                render_block(sink, ch, c)
        return
    if 'chartttl' in cls:
        p = sink.par()
        tune(p, 5, 2, line=1.05, keep=True)
        run(p, '▌ ', size=9, color=V['orange'])
        run(p, n.all_text().strip().upper(), size=px2pt(10), bold=True,
            color=V['purple'], spacing=8)
        return
    if 'scaleleg' in cls:
        items = [x for x in n.children if x.tag != '#text']
        def fill_chip(cell, item, cx):
            cell_borders(cell, top=(6, V['hair']), left=(6, V['hair']),
                         bottom=(6, V['hair']), right=(6, V['hair']))
            p = cell.paragraphs[0]
            tune(p, 2, 2, line=1.05, left=80, right=80)
            emit_inline(p, item, cx.sub(size=px2pt(9)))
        comp_grid(sink, n, c, 3 if len(items) > 4 else len(items), fill_chip, gap=70)
        sink.after_table()
        return
    if 'optgrid' in cls:
        ncols = 2
        gtc = n.style.get('grid-template-columns', '')
        m = re.search(r'repeat\((\d+)', gtc)
        if m:
            ncols = int(m.group(1))
        elif gtc.count('fr') > 0:
            ncols = gtc.count('fr')
        comp_grid(sink, n, c, ncols, fill_opt, gap=90)
        sink.after_table()
        return
    if 'fgrid' in cls:
        comp_grid(sink, n, c, 2, fill_fg, gap=120)
        sink.after_table()
        return
    if 'fg' in cls:
        t = new_table(sink, 1, 1, [CW])
        table_indent(t, 0)
        cell_margins(t, 60, 110, 60, 110)
        fill_fg(t.cell(0, 0), n, c)
        sink.after_table()
        return
    if 'bands' in cls:
        comp_grid(sink, n, c, 4, fill_band, gap=110)
        sink.after_table()
        return
    if 'two' in cls:
        comp_grid(sink, n, c, 2, fill_plain, gap=150)
        sink.after_table()
        return
    if 'sg' in cls:
        comp_sg(sink, n, c)
        sink.after_table()
        return
    if 'opsum' in cls:
        box(sink, n, c.sub(size=px2pt(9.5)), V['purple-soft'], V['purple-line'],
            left_accent=V['purple'])
        sink.after_table()
        return
    if 'legal' in cls:
        box(sink, n, c, 'F6F4FB', V['purple-line'], size=px2pt(9.5 if c['roomy'] else 7.8))
        sink.after_table()
        return
    if 'obsline' in cls:
        box(sink, n, c, 'FAF8FF', 'D3C8EF', size=px2pt(9.5), min_lines=1)
        sink.after_table()
        return
    if 'approve' in cls:
        box(sink, n, c.sub(italic=True, size=px2pt(11)), V['purple-soft'],
            V['purple-line'], left_accent=V['purple'])
        sink.after_table()
        return
    if 'diag' in cls:
        box(sink, n, c, V['purple-soft'], V['purple-line'], left_accent=V['purple'])
        sink.after_table()
        return
    if 'meta' in cls:
        box(sink, n, c, 'FAF9FD', V['hair'])
        sink.after_table()
        return
    if 'box' in cls:
        fill = 'FAF9FD' if 'soft' in cls else 'FFFFFF'
        ml = 0
        if 'empty' in cls:
            mh = sizepx(n.style.get('min-height'), 44)
            ml = max(1, int(mh / 16))
        box(sink, n, c, fill, V['hair'], min_lines=ml)
        sink.after_table()
        return
    if 'chartcard' in cls or 'chartgrid' in cls:
        render_children(sink, n, c)
        return
    if n.tag in ('h1', 'h2', 'h3'):
        p = sink.par()
        tune(p, 4, 2, line=1.1, align=c['align'], keep=True)
        run(p, n.all_text().strip(), size=c['size'], bold=True, color=V['purple'])
        return
    fgs = [x for x in n.children if x.tag != '#text' and x.has('fg')]
    if len(fgs) > 1 and len(fgs) == len([x for x in n.children if x.tag != '#text']):
        comp_grid(sink, n, c, len(fgs), fill_fg, gap=120)
        sink.after_table()
        return
    # generyczny div/p
    kids_block = any(is_block(x) for x in n.children)
    if kids_block:
        render_children(sink, n, c)
    else:
        if not n.all_text().strip():
            st = n.style
            if 'min-height' in st or 'height' in st:
                sink.spacer(px2pt(sizepx(st.get('min-height') or st.get('height'), 6)))
            return
        p = sink.par()
        mt = sizepx(n.style.get('margin-top'), None)
        p_before = px2pt(mt) if mt is not None else 2.0
        tune(p, p_before, 2, line=c.get('lh', 1.3), align=c['align'])
        emit_inline(p, n, c)


# ---------------------------------------------------------------- dokument
def setup_section(sec):
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(9)
    sec.bottom_margin = Mm(10)
    sec.left_margin = Mm(11)
    sec.right_margin = Mm(11)
    sec.header_distance = Mm(6)
    sec.footer_distance = Mm(5)


def build_footer(sec, foot_node):
    sec.footer.is_linked_to_previous = False
    p = sec.footer.paragraphs[0]
    tune(p, 3, 0, line=1.0)
    par_borders(p, top=(4, V['hair']))
    tabstop(p, CW, 'right')
    spans = [ch for ch in foot_node.children if ch.tag == 'span']
    if spans:
        emit_inline(p, spans[0], BASE.sub(size=px2pt(8), color=V['muted']))
    p.add_run('\t')
    if len(spans) > 1:
        emit_inline(p, spans[1], BASE.sub(size=px2pt(8), color=V['muted'],
                                          bcolor=V['purple']))
    # pusty akapit domykający
    return p


def sheet_meta(sheet):
    sub = pill = tag = ''
    pill_purple = False
    head = None
    for d in sheet.find_all('div', 'head'):
        head = d
        break
    if head is not None:
        s = head.find_all('div', 'sub')
        if s:
            sub = s[0].all_text().strip()
        pl = head.find_all('span', 'pill')
        if pl:
            pill = pl[0].all_text().strip()
            pill_purple = 'p' in pl[0].cls
        tg = head.find_all('div', 'tag')
        if tg:
            tag = tg[0].all_text().strip()
    return sub, pill, pill_purple, tag


def fix_blank_lines(doc):
    """Akapity złożone wyłącznie z kropkowanej podkreślonej spacji zamień na
    akapit z kropkowaną krawędzią dolną (LibreOffice/Word przycinają końcowe spacje)."""
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    for p_el in body.iter(qn('w:p')):
        par = Paragraph(p_el, None)
        txt = ''.join(r.text for r in par.runs)
        if txt.strip():
            continue
        dotted_runs = []
        for r in par.runs:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is None:
                continue
            u = rPr.find(qn('w:u'))
            if u is not None and u.get(qn('w:val')) == 'dotted':
                dotted_runs.append((r, u, rPr))
        if not dotted_runs:
            continue
        for r, u, rPr in dotted_runs:
            rPr.remove(u)
        pPr = p_el.find(qn('w:pPr'))
        if pPr is None:
            pPr = p_el.makeelement(qn('w:pPr'), {})
            p_el.insert(0, pPr)
        if pPr.find(qn('w:pBdr')) is None:
            par_borders(par, bottom=(6, DOT, 'dotted'))


def fix_settings(doc):
    st = doc.settings.element
    for z in st.findall(qn('w:zoom')):
        if z.get(qn('w:percent')) is None:
            z.set(qn('w:percent'), '100')


def build(html_path, out_path):
    html = open(html_path, encoding='utf-8').read()
    html = re.sub(r'<script.*?</script>', '', html, flags=re.S)
    root = parse_html(html)
    sheets = [s for s in root.find_all('section') if s.has('sheet')]

    doc = Document()
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(9)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.line_spacing = 1.15

    for i, sheet in enumerate(sheets):
        if i == 0:
            sec = doc.sections[0]
        else:
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
        setup_section(sec)
        foots = sheet.find_all('div', 'foot')
        if foots:
            build_footer(sec, foots[0])

        ctx = BASE.sub(roomy=sheet.has('roomy'))
        sink = Sink(doc)
        sub, pill, pill_purple, tag = sheet_meta(sheet)
        comp_head(sink, sub, pill, pill_purple, tag)
        comp_rule(sink)
        studs = sheet.find_all('div', 'student')
        if studs:
            sink.spacer(2)
            comp_student(sink, studs[0], ctx)
        render_children(sink, sheet, ctx)

    fix_blank_lines(doc)
    fix_settings(doc)
    doc.save(out_path)
    return len(sheets)


if __name__ == '__main__':
    src, dst = sys.argv[1], sys.argv[2]
    n = build(src, dst)
    print('OK', dst, n, 'stron źródłowych')
