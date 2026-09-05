# -*- coding: utf-8 -*-
"""
Nanosi na oryginalny skrypt przedszkolny poprawki z audytu Strażnika Prawa
(5 września 2026 r.). Zmiany wpisywane są kolorem niebieskim, w miejscach,
których dotyczą; oryginalne brzmienie pozostaje nienaruszone.

Wejście : Skrypt_dla_nauczycieli.docx (wydanie 1, przedszkole)
Wyjście : Skrypt_dla_nauczycieli_PRZEDSZKOLE_wydanie2_po_audycie.docx
"""
import copy
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

BLUE = "0B4F9E"


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for t2 in cell.tables:
                    for r2 in t2.rows:
                        for c2 in r2.cells:
                            for p2 in c2.paragraphs:
                                yield p2


def find(doc, anchor):
    for p in iter_paragraphs(doc):
        if anchor in p.text:
            return p
    raise SystemExit('NIE ZNALEZIONO KOTWICY: ' + anchor)


def blue_run(paragraph, text, size=None, template=None):
    """Dopisuje na końcu akapitu niebieski fragment korekty."""
    src = template if template is not None else (paragraph.runs[-1] if paragraph.runs else None)
    r = paragraph.add_run(text)
    if src is not None:
        r._element.insert(0, copy.deepcopy(src._element.rPr)) if src._element.rPr is not None else None
    r.font.name = 'Arial'
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    if size:
        r.font.size = Pt(size)
    elif src is not None and src.font.size is not None:
        r.font.size = src.font.size
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(a), 'Arial')
    return r


def clone_bullet(paragraph, text):
    """Klonuje akapit listy (zachowując jego formatowanie) i wstawia go poniżej."""
    new = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new)
    from docx.text.paragraph import Paragraph
    np = Paragraph(new, paragraph._parent)
    for r in list(np.runs)[1:]:
        r._element.getparent().remove(r._element)
    if not np.runs:
        raise SystemExit('Akapit wzorcowy nie ma biegów: ' + paragraph.text[:60])
    r0 = np.runs[0]
    r0.text = text
    r0.font.bold = True
    r0.font.color.rgb = RGBColor.from_string(BLUE)
    return np


def main(src, dst):
    doc = Document(src)

    # ------------------------------------------------------------------ 1
    # Strona tytułowa — znacznik wydania 2
    p = find(doc, 'Dokumentacja przedszkolna')
    blue_run(p, '   ·   WYDANIE 2 po audycie podstaw prawnych z 5 września 2026 r. — '
                'fragmenty dopisane w audycie oznaczono kolorem niebieskim.')

    # ------------------------------------------------------------------ 2
    # Nota „Jak korzystać ze skryptu" — legenda koloru
    p = find(doc, 'Skrypt można drukować w całości albo częściami')
    blue_run(p, '  W tym wydaniu kolorem niebieskim oznaczono uzupełnienia i sprostowania wprowadzone '
                'w audycie Strażnika Prawa 5 września 2026 r. Oryginalne brzmienie wydania 1 pozostawiono '
                'bez zmian — audyt nic nie usuwa, tylko dopowiada.')

    # ------------------------------------------------------------------ 3
    # Część 1 — podstawa programowa: dopisanie aktu zmieniającego
    p = find(doc, 'Podstawa programowa wychowania przedszkolnego — rozporządzenie ME z 11 marca')
    blue_run(p, ' Akt zmieniony rozporządzeniem Dz.U. 2026 poz. 958, również od 1.09.2026 — zmiana '
                'nie dotyczy załącznika nr 1, czyli podstawy programowej wychowania przedszkolnego. '
                'Podstawa obejmuje od razu wszystkie grupy wiekowe.')

    # ------------------------------------------------------------------ 4
    # Część 1 — orzeczenia: status aktu i akt uchylony
    p = find(doc, 'Orzeczenia i opinie zespołów orzekających — rozporządzenie ME z 2 marca 2026 r.')
    blue_run(p, ' Akt obowiązuje od 14.04.2026 i z tym dniem uchylił rozporządzenie MEN z 7 września '
                '2017 r. (t.j. Dz.U. 2023 poz. 2061) — tego publikatora nie wolno już cytować jako '
                'obowiązującego. Zakres opinii do 31.08.2026 określa § 33.')

    # ------------------------------------------------------------------ 5
    # Część 1 — dwa brakujące akty w podstawie prawnej części
    p = find(doc, 'Prawo oświatowe — ustawa z 14 grudnia 2016 r.')
    clone_bullet(p,
        'Świadectwa i druki — rozporządzenie MEiN z 7 czerwca 2023 r. w sprawie świadectw, dyplomów '
        'państwowych i innych druków (Dz.U. 2023 poz. 1120, z późn. zm.): wzór informacji o gotowości '
        'dziecka do podjęcia nauki w szkole podstawowej. To jedyna podstawa druku wymienianego '
        'w częściach 3 i 6 oraz w kalendarzu.')
    clone_bullet(p,
        'Organizacja przedszkoli — rozporządzenie MEN z 28 lutego 2019 r. w sprawie szczegółowej '
        'organizacji publicznych szkół i publicznych przedszkoli (t.j. Dz.U. 2023 poz. 2736, z późn. zm.): '
        'godziny pracy przedszkola, oddziały, arkusz organizacji — dane sekcji II metryczki.')

    # ------------------------------------------------------------------ 6
    # Część 1, akapit 21 — ostrzeżenie przed cytowaniem publikatorów pierwotnych
    p = find(doc, 'Warto zapamiętać: pozycja tysiąc pięćset siedemdziesiąt osiem')
    blue_run(p, ' I od razu zastrzeżenie, bo to zdanie bywa źle rozumiane. Pozycje tysiąc pięćset '
                'siedemdziesiąt osiem i tysiąc pięćset dziewięćdziesiąt jeden to publikatory pierwotne. '
                'Służą do rozpoznania, o który akt chodzi, a nie do cytowania. W dokumencie dziecka '
                'wpisujemy wyłącznie obowiązujące teksty jednolite: tysiąc trzysta dziewięć dla '
                'kształcenia specjalnego i tysiąc siedemset dziewięćdziesiąt osiem dla pomocy '
                'psychologiczno-pedagogicznej.')

    # ------------------------------------------------------------------ 7
    # Część 3 — podstawa prawna: podawanie leków
    p = find(doc, 'RODO — art. 5 ust. 1 lit. c (minimalizacja danych)')
    clone_bullet(p,
        'Podawanie leków w przedszkolu nie ma odrębnej podstawy ustawowej. Ustawa z 12 kwietnia 2019 r. '
        'o opiece zdrowotnej nad uczniami (Dz.U. 2019 poz. 1078) dotyczy uczniów, a nie dzieci objętych '
        'wychowaniem przedszkolnym. Podstawą działania są: pisemne upoważnienie rodziców ze wskazaniem '
        'leku, dawki i godzin, dobrowolna pisemna zgoda nauczyciela, przeszkolenie oraz procedura '
        'wprowadzona zarządzeniem dyrektora. Nauczyciel nie ma obowiązku podania leku.')

    # ------------------------------------------------------------------ 8
    # Część 3, akapit 16 — doprecyzowanie „na jakiej podstawie"
    p = find(doc, 'kto podaje lek, na jakiej podstawie, gdzie lek jest przechowywany')
    blue_run(p, ' Doprecyzujmy, co znaczy „na jakiej podstawie”, bo nie chodzi o przepis ustawy — takiego '
                'przepisu dla przedszkola nie ma. Chodzi o trzy dokumenty: pisemne upoważnienie rodziców '
                'ze wskazaniem leku, dawki i godzin podania, dobrowolną pisemną zgodę nauczyciela, który '
                'lek podaje, oraz procedurę przyjętą zarządzeniem dyrektora. Bez tego kompletu leku '
                'nie podajemy.')

    # ------------------------------------------------------------------ 9
    # Część 4 — sprostowanie przypisania obowiązku diagnozy przedszkolnej
    p = find(doc, 'Obserwacja pedagogiczna i analiza gotowości szkolnej — rozporządzenie o dokumentacji')
    blue_run(p, ' Sprostowanie: rozporządzenie o dokumentacji jest podstawą dla arkuszy obserwacji '
                'jako dokumentacji badań i czynności uzupełniających. Sam obowiązek prowadzenia '
                'obserwacji pedagogicznych i przeprowadzenia analizy gotowości dziecka do nauki w szkole '
                'wynika z zadań przedszkola określonych w podstawie programowej (Dz.U. 2026 poz. 378), '
                'a wzór informacji o gotowości — z rozporządzenia o świadectwach (Dz.U. 2023 poz. 1120).')

    # ----------------------------------------------------------------- 10
    # Część 6 — numeracja ustępów § 6 do sprawdzenia w ogłoszonym tekście
    p = find(doc, 'ust. 9–11 (ocena co najmniej dwa razy w roku, prawa rodziców)')
    blue_run(p, ' Uwaga redakcyjna: numerację ustępów sprawdź w ogłoszonym tekście jednolitym przed '
                'przepisaniem do druku. Materie są w § 6 na pewno, ale okresowa ocena, prawo rodziców '
                'do udziału w spotkaniach zespołu, obowiązek pisemnego zawiadomienia o terminie '
                'i prawo do kopii oceny oraz programu to cztery odrębne ustępy, a nie trzy. '
                'Bezpieczny zapis w dokumencie dziecka: „§ 6 rozporządzenia” bez numeru ustępu.')

    # ----------------------------------------------------------------- 11
    # Część 6 — podstawa informacji o gotowości szkolnej
    p = find(doc, 'Informacja o gotowości szkolnej raz, do końca kwietnia')
    blue_run(p, ' Dodajmy podstawę tego druku, bo w wydaniu pierwszym jej zabrakło. Wzór informacji '
                'o gotowości dziecka do podjęcia nauki w szkole podstawowej określa rozporządzenie '
                'Ministra Edukacji i Nauki z siódmego czerwca dwa tysiące dwudziestego trzeciego roku '
                'w sprawie świadectw, dyplomów państwowych i innych druków, Dziennik Ustaw pozycja '
                'tysiąc sto dwadzieścia. Informację sporządzamy na tym wzorze, a nie na druku własnym.')

    # ----------------------------------------------------------------- 12
    # Dodatek — kalendarz: podstawa druku informacji o gotowości
    for p in iter_paragraphs(doc):
        if p.text.strip() == 'Informacja o gotowości szkolnej':
            blue_run(p, '  (wzór — Dz.U. 2023 poz. 1120)', size=8)
            break

    doc.save(dst)
    print('Zapisano:', dst)


if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
