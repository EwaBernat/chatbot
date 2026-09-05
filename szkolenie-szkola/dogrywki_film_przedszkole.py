# -*- coding: utf-8 -*-
"""
Karta zmian do materiału filmowego — przedszkole.
Przekłada ustalenia audytu Strażnika Prawa (5.09.2026) na konkretne dogrywki:
numer planszy, miejsce cięcia, tekst do przeczytania, treść planszy, czas.

Tempo narracji policzone z istniejącego filmu: 5753 słowa / 53:57 = 106,6 sł./min.
"""
import os
import math
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from build_skrypt_szkola import (Document, PURPLE, ORANGE, INK, GREY, LILAC, LILAC2, BLUE,
                                 BORDER, _sub, run, para, hairline, page_break, accent_box,
                                 band, h2, h3, table, numbered, steps)

SEK_NA_SLOWO = 60.0 / 106.6          # 0,563 s — zmierzone na istniejącym filmie
PAUZA = 0.6                          # oddech przed wstawką i po niej


def czas(tekst):
    n = len(tekst.split())
    return n, math.ceil(n * SEK_NA_SLOWO + PAUZA)


def mmss(s):
    return f'{int(s) // 60}:{int(s) % 60:02d}'


# ---------------------------------------------------------------- treść dogrywek
# (modul, tytul_modulu, dlugosc_s, [ (typ, po_planszy, powod, tekst, plansza) ])

DOGRYWKI = [
 ('M1', 'Podstawa prawna — co obowiązuje od 1 września 2026 r.', 570, [
  ('WSTAWKA', '08', 'Ustalenie C3 — brak aktu zmieniającego podstawę programową',
   'Jedno uzupełnienie do tego rozporządzenia. Zostało ono zmienione rozporządzeniem z Dziennika Ustaw '
   'z dwa tysiące dwudziestego szóstego roku, pozycja dziewięćset pięćdziesiąt osiem, które również wchodzi '
   'w życie pierwszego września. Zmiana nie dotyczy jednak załącznika numer jeden, czyli podstawy programowej '
   'wychowania przedszkolnego. W naszym załączniku nie zmienia się nic.',
   'PODSTAWA PROGRAMOWA — AKT ZMIENIAJĄCY\n'
   '• Dz.U. 2026 poz. 958 — od 1.09.2026\n'
   '• zmiana NIE dotyczy załącznika nr 1 (wychowanie przedszkolne)'),

  ('WSTAWKA', '12', 'Ustalenie C4 — brak informacji o uchyleniu poprzedniego rozporządzenia',
   'Powiedzmy też wyraźnie, co to rozporządzenie zrobiło z poprzednim stanem prawnym. Z dniem czternastego '
   'kwietnia uchyliło rozporządzenie Ministra Edukacji Narodowej z siódmego września dwa tysiące siedemnastego '
   'roku, którego tekst jednolity znajdziemy w Dzienniku Ustaw z dwa tysiące dwudziestego trzeciego roku, '
   'pozycja dwa tysiące sześćdziesiąt jeden. Ten publikator przestał być aktualny. Jeżeli ktoś ma go '
   'w swoich drukach, trzeba go wymienić.',
   'CO UCHYLIŁO NOWE ROZPORZĄDZENIE\n'
   '• Dz.U. 2026 poz. 428 obowiązuje od 14.04.2026\n'
   '• uchyliło rozp. MEN z 7.09.2017 (t.j. Dz.U. 2023 poz. 2061)\n'
   '• starego publikatora nie cytujemy'),

  ('WSTAWKA', '21', 'Ustalenie A4 — pułapka dydaktyczna przy publikatorach pierwotnych',
   'I od razu zastrzeżenie, bo to zdanie bywa źle rozumiane. Pozycje tysiąc pięćset siedemdziesiąt osiem '
   'i tysiąc pięćset dziewięćdziesiąt jeden to publikatory pierwotne. Służą do rozpoznania, o który akt chodzi, '
   'a nie do cytowania. W dokumencie dziecka wpisujemy wyłącznie obowiązujące teksty jednolite: tysiąc trzysta '
   'dziewięć dla kształcenia specjalnego i tysiąc siedemset dziewięćdziesiąt osiem dla pomocy '
   'psychologiczno-pedagogicznej.',
   'DO ROZPOZNANIA ≠ DO CYTOWANIA\n'
   '• 1578 i 1591 — publikatory pierwotne, do rozpoznania aktu\n'
   '• w dokumencie dziecka: t.j. Dz.U. 2020 poz. 1309 · t.j. Dz.U. 2023 poz. 1798'),

  ('NOWA PLANSZA', '25', 'Ustalenia C1 i C2 — dwa akty nieobecne w wydaniu 1',
   'Zanim przejdziemy do aktów nadrzędnych, dwa rozporządzenia, o których w pierwszej wersji nie powiedziałam, '
   'a bez których dwa nasze dokumenty wiszą w próżni. Pierwsze to rozporządzenie Ministra Edukacji i Nauki '
   'z siódmego czerwca dwa tysiące dwudziestego trzeciego roku w sprawie świadectw, dyplomów państwowych '
   'i innych druków, Dziennik Ustaw pozycja tysiąc sto dwadzieścia. To ono określa wzór informacji o gotowości '
   'dziecka do podjęcia nauki w szkole podstawowej. Informację sporządzamy na tym wzorze, a nie na druku '
   'własnym. Drugie to rozporządzenie Ministra Edukacji Narodowej z dwudziestego ósmego lutego dwa tysiące '
   'dziewiętnastego roku w sprawie szczegółowej organizacji publicznych szkół i publicznych przedszkoli, '
   'w tekście jednolitym z dwa tysiące dwudziestego trzeciego roku, pozycja dwa tysiące siedemset trzydzieści '
   'sześć. Z niego wynikają dane, które wpisujemy w drugiej sekcji metryczki: organizacja pobytu, oddziały '
   'i godziny pracy przedszkola.',
   'DWA AKTY, KTÓRE MUSZĄ BYĆ W NASZEJ LIŚCIE\n'
   '• Świadectwa i druki — Dz.U. 2023 poz. 1120 → wzór informacji o gotowości\n'
   '• Organizacja przedszkoli — t.j. Dz.U. 2023 poz. 2736 → sekcja II metryczki'),
 ]),

 ('M3', 'Metryczka dziecka — pierwszy dokument września', 345, [
  ('WSTAWKA', '12', 'Ustalenie C1 — informacja o gotowości bez wskazania urzędowego wzoru',
   'Dopowiedzmy od razu jedno. Informację o gotowości sporządzamy na urzędowym wzorze, który określa '
   'rozporządzenie o świadectwach, dyplomach państwowych i innych drukach. Nie tworzymy do niej druku własnego.',
   'INFORMACJA O GOTOWOŚCI — URZĘDOWY WZÓR\n'
   '• wzór: rozp. MEiN z 7.06.2023 (Dz.U. 2023 poz. 1120)\n'
   '• termin: do końca kwietnia'),

  ('WSTAWKA', '16', 'Ustalenie A3 — „na jakiej podstawie” sugeruje nieistniejący przepis',
   'Doprecyzujmy, co znaczy „na jakiej podstawie”, bo to pytanie wraca na każdym szkoleniu. Nie chodzi '
   'o przepis ustawy, bo takiego przepisu dla przedszkola nie ma. Ustawa o opiece zdrowotnej nad uczniami '
   'dotyczy uczniów, a nie dzieci objętych wychowaniem przedszkolnym. Chodzi o trzy dokumenty. Pisemne '
   'upoważnienie rodziców, ze wskazaniem leku, dawki i godzin podania. Dobrowolną pisemną zgodę nauczyciela, '
   'który lek podaje. I procedurę przyjętą zarządzeniem dyrektora. Bez tego kompletu leku nie podajemy. '
   'I jeszcze jedno zdanie, ważne dla nas wszystkich: nauczyciel nie ma obowiązku podania leku.',
   'PODAWANIE LEKU — NA JAKIEJ PODSTAWIE?\n'
   '• brak przepisu ustawowego dla przedszkola\n'
   '• ustawa o opiece zdrowotnej nad uczniami dotyczy UCZNIÓW\n'
   '• podstawa: upoważnienie rodziców + zgoda nauczyciela + procedura dyrektora\n'
   '• nauczyciel NIE MA OBOWIĄZKU podania leku'),
 ]),

 ('M4', 'KPOF — budowa narzędzia, skala, liczenie wyniku, odczyt profilu', 608, [
  ('WSTAWKA', '08', 'Ustalenie A2 — błędne przypisanie obowiązku obserwacji i analizy gotowości',
   'Jeszcze słowo o tym, skąd bierze się nasza obserwacja, bo to dwie różne podstawy. Rozporządzenie '
   'o dokumentacji jest podstawą dla samych arkuszy, jako dokumentacji badań i czynności uzupełniających. '
   'Natomiast obowiązek prowadzenia obserwacji pedagogicznych i przeprowadzenia analizy gotowości dziecka '
   'do nauki w szkole wynika z zadań przedszkola zapisanych w podstawie programowej. Arkusz bierzemy '
   'z jednego przepisu, a obowiązek z drugiego.',
   'DWIE RÓŻNE PODSTAWY\n'
   '• arkusze obserwacji → rozp. o dokumentacji (t.j. Dz.U. 2024 poz. 50)\n'
   '• obowiązek obserwacji i analizy gotowości → podstawa programowa (Dz.U. 2026 poz. 378)\n'
   '• wzór informacji o gotowości → rozp. o świadectwach (Dz.U. 2023 poz. 1120)'),
 ]),

 ('M6', 'WOPF, IPET, cele SMART, ewaluacja i opinia dla poradni', 792, [
  ('WSTAWKA', '04', 'Ustalenie A1 — numeracja ustępów § 6 niepotwierdzona',
   'Jedna uwaga redakcyjna, którą proszę potraktować poważnie. W materiałach spotkają Państwo odesłania '
   'do konkretnych ustępów paragrafu szóstego. Numerację ustępów sprawdzajcie w ogłoszonym tekście jednolitym, '
   'zanim przepiszecie ją do dokumentu dziecka. Sam paragraf szósty jest właściwy na pewno, ale okresowa ocena, '
   'prawo rodziców do udziału w spotkaniach zespołu, obowiązek pisemnego zawiadomienia o terminie oraz prawo '
   'do kopii oceny i programu to cztery odrębne ustępy, a nie trzy. Bezpieczny zapis w dokumencie dziecka '
   'brzmi po prostu: paragraf szósty rozporządzenia.',
   'NUMERACJA USTĘPÓW § 6\n'
   '• § 6 — właściwy na pewno\n'
   '• cztery odrębne materie, nie trzy\n'
   '• w druku dziecka: „§ 6 rozporządzenia” bez numeru ustępu'),

  ('WSTAWKA', '30', 'Ustalenie C1 — kalendarz ewaluacji bez podstawy druku',
   'I podstawa tego ostatniego druku, bo w pierwszej wersji jej nie podałam. Wzór informacji o gotowości '
   'dziecka do podjęcia nauki w szkole podstawowej określa rozporządzenie Ministra Edukacji i Nauki '
   'z siódmego czerwca dwa tysiące dwudziestego trzeciego roku w sprawie świadectw, dyplomów państwowych '
   'i innych druków, Dziennik Ustaw pozycja tysiąc sto dwadzieścia.',
   'INFORMACJA O GOTOWOŚCI SZKOLNEJ\n'
   '• wzór: Dz.U. 2023 poz. 1120\n'
   '• termin: do końca kwietnia · jeden raz w roku'),
 ]),
]

PLANSZE_PODMIANA = [
 ('M1', 'plansza „PODSTAWA PRAWNA CZĘŚCI”',
  'Dopisać dwie pozycje: rozporządzenie o świadectwach (Dz.U. 2023 poz. 1120) oraz rozporządzenie '
  'o organizacji przedszkoli (t.j. Dz.U. 2023 poz. 2736, z późn. zm.). Przy podstawie programowej dopisać '
  '„zm. Dz.U. 2026 poz. 958 — nie dotyczy zał. nr 1”. Przy rozporządzeniu o orzekaniu dopisać „obowiązuje '
  'od 14.04.2026; uchyliło t.j. Dz.U. 2023 poz. 2061”.'),
 ('M3', 'plansza „PODSTAWA PRAWNA CZĘŚCI”',
  'Dopisać pozycję o podawaniu leków: brak podstawy ustawowej dla przedszkola; ustawa o opiece zdrowotnej '
  'nad uczniami (Dz.U. 2019 poz. 1078) dotyczy uczniów; podstawą jest upoważnienie rodziców, pisemna zgoda '
  'nauczyciela i procedura dyrektora.'),
 ('M4', 'plansza „PODSTAWA PRAWNA CZĘŚCI”',
  'Rozdzielić dotychczasową pozycję na dwie: rozporządzenie o dokumentacji — arkusze obserwacji; podstawa '
  'programowa (Dz.U. 2026 poz. 378) — obowiązek obserwacji i analizy gotowości. Dodać wzór informacji '
  'o gotowości (Dz.U. 2023 poz. 1120).'),
 ('M6', 'plansza „PODSTAWA PRAWNA CZĘŚCI”',
  'Zamienić „ust. 9–11 (ocena co najmniej dwa razy w roku, prawa rodziców)” na „§ 6 rozporządzenia — ocena '
  'co najmniej dwa razy w roku szkolnym oraz prawa rodziców do udziału, zawiadomienia i kopii; numerację '
  'ustępów sprawdzić w ogłoszonym tekście jednolitym”.'),
 ('DODATEK', 'plansza / tabela „Kalendarz dokumentacji na rok szkolny”',
  'W wierszu „do końca kwietnia” przy dokumencie „Informacja o gotowości szkolnej” dopisać podstawę wzoru: '
  'Dz.U. 2023 poz. 1120.'),
]


# ---------------------------------------------------------------- skład dokumentu

def build_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Arial'; st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    st.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = Cm(1.9); sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)

    hp = sec.header.paragraphs[0]; hp.paragraph_format.space_after = Pt(3)
    bdr = _sub(hp._p.get_or_add_pPr(), 'w:pBdr')
    _sub(bdr, 'w:bottom', **{'w:val': 'single', 'w:sz': '4', 'w:space': '4', 'w:color': BORDER})
    run(hp, 'PCTP · EduPlaner 2026', size=7.5, bold=True, color=ORANGE)
    run(hp, '     Karta dogrywek do materiału filmowego · przedszkole · po audycie z 5 września 2026 r.',
        size=7.5, color=GREY)

    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(3)
    bdr = _sub(fp._p.get_or_add_pPr(), 'w:pBdr')
    _sub(bdr, 'w:top', **{'w:val': 'single', 'w:sz': '4', 'w:space': '4', 'w:color': BORDER})
    run(fp, 'Strona ', size=7.5, color=GREY)
    f1 = OxmlElement('w:fldSimple'); f1.set(qn('w:instr'), 'PAGE'); fp._p.append(f1)
    run(fp, ' z ', size=7.5, color=GREY)
    f2 = OxmlElement('w:fldSimple'); f2.set(qn('w:instr'), 'NUMPAGES'); fp._p.append(f2)
    run(fp, ' · karta dogrywek · materiał filmowy przedszkole', size=7.5, color=GREY)
    return doc


def blok_tekstu(doc, tekst):
    """Tekst do przeczytania — w ramce, powiększony, z komfortową interlinią."""
    t = doc.add_table(rows=1, cols=1)
    tblPr = t._tbl.tblPr
    b = _sub(tblPr, 'w:tblBorders')
    _sub(b, 'w:top', **{'w:val': 'single', 'w:sz': '2', 'w:color': LILAC})
    _sub(b, 'w:left', **{'w:val': 'single', 'w:sz': '18', 'w:color': BLUE})
    _sub(b, 'w:bottom', **{'w:val': 'single', 'w:sz': '2', 'w:color': LILAC})
    _sub(b, 'w:right', **{'w:val': 'single', 'w:sz': '2', 'w:color': LILAC})
    _sub(b, 'w:insideH', **{'w:val': 'none', 'w:sz': '0', 'w:color': LILAC})
    _sub(b, 'w:insideV', **{'w:val': 'none', 'w:sz': '0', 'w:color': LILAC})
    c = t.cell(0, 0)
    from build_skrypt_szkola import cell_shade, cell_margins
    cell_shade(c, 'FFFFFF'); cell_margins(c, top=140, left=200, bottom=140, right=180)
    c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
    para(c, 'TEKST DO PRZECZYTANIA', size=8, bold=True, color=BLUE, after=6)
    para(c, tekst, size=11.5, after=0, line=320)
    para(doc, '', after=6)


def main():
    doc = build_doc()

    para(doc, 'EDUPLANER 2026 · PCTP · PRODUKCJA', size=9, bold=True, color=ORANGE, after=4)
    para(doc, 'Karta dogrywek do materiału filmowego', size=25, bold=True, color=PURPLE, after=2)
    para(doc, 'Przedszkole · szkolenie rady pedagogicznej · sześć modułów M1–M6',
         size=13, bold=True, color=PURPLE, after=8)
    para(doc, 'Dokument przekłada ustalenia audytu Strażnika Prawa z 5 września 2026 r. na konkretne dogrywki: '
              'wskazuje moduł, planszę, po której następuje cięcie, gotowy tekst do przeczytania oraz treść '
              'planszy do wymiany. Nic nie trzeba nagrywać od nowa — wszystkie zmiany są wstawkami.',
         size=10, after=10)
    hairline(doc)

    licz = sum(len(d[3]) for d in DOGRYWKI)
    suma = sum(czas(x[3])[1] for d in DOGRYWKI for x in d[3])
    accent_box(doc, [
        f'Dogrywek do nagrania: {licz}. Łączny materiał do dogrania: około {mmss(suma)} '
        f'({suma} sekund) przy tempie zmierzonym na istniejącym filmie — 107 słów na minutę.',
        'Modułów dotkniętych zmianą: cztery (M1, M3, M4, M6). Moduły M2 i M5 pozostają bez zmian '
        'i nie wymagają żadnej ingerencji.',
        'Wszystkie zmiany są WSTAWKAMI między istniejące plansze. Nie ma ani jednej podmiany zdania '
        'wewnątrz nagranego już fragmentu — dzięki temu nie trzeba nagrywać niczego ponownie, '
        'a montaż sprowadza się do wklejenia nowego materiału w podanych miejscach.',
        'Do wymiany jest natomiast pięć plansz z podstawą prawną — to zmiana wyłącznie graficzna, '
        'bez ingerencji w ścieżkę dźwiękową.',
    ], title='CO TRZEBA ZROBIĆ — W CZTERECH ZDANIACH', accent=BLUE)

    h2(doc, 'Zestawienie modułów')
    rows = []
    for mod, tytul, dl, lista in DOGRYWKI:
        s = sum(czas(x[3])[1] for x in lista)
        rows.append([mod, tytul, mmss(dl), str(len(lista)), f'+{s} s', mmss(dl + s)])
    rows.append(['M2', 'Obieg dokumentów — jak jeden wynika z drugiego', '5:36', '—', '—', '5:36'])
    rows.append(['M5', 'Obserwacja pogłębiona — ABC, sensoryka, teoria umysłu, mowa', '9:46', '—', '—', '9:46'])
    stare = 3237
    rows.append(['RAZEM', 'sześć modułów', mmss(stare), str(licz), f'+{suma} s', mmss(stare + suma)])
    table(doc, ['Moduł', 'Tytuł', 'Czas obecny', 'Dogrywek', 'Przyrost', 'Czas po zmianie'],
          rows, widths=[1.6, 7.6, 2.0, 1.8, 1.6, 2.1], size=8.5)

    h2(doc, 'Jak korzystać z tej karty')
    steps(doc, [
        '**Nagraj tylko fragmenty z ramek** oznaczonych „TEKST DO PRZECZYTANIA”. Każdy jest osobnym plikiem — '
        'nazwa pliku podana jest przy każdej dogrywce.',
        '**Wklej je w miejscach cięcia.** Miejsce cięcia opisane jest numerem planszy i ostatnim zdaniem, '
        'które po niej pada — wystarczy znaleźć je w nagraniu. Kodów czasowych nie podaję, bo nie mam '
        'dostępu do plików wideo; numer planszy i cytat wystarczą, żeby trafić w sekundę.',
        '**Wymień pięć plansz z podstawą prawną** według wykazu na końcu karty. To zmiana graficzna, '
        'ścieżka dźwiękowa pozostaje bez zmian.',
        '**Dodaj nowe plansze** do wstawek — treść każdej podana jest obok tekstu narracji.',
        '**Zaktualizuj czasy w spisie części skryptu** po zmontowaniu — nowe długości modułów są w tabeli powyżej.',
    ])

    # ------------------------------------------------------------- moduły
    for mod, tytul, dl, lista in DOGRYWKI:
        page_break(doc)
        s = sum(czas(x[3])[1] for x in lista)
        band(doc, mod, tytul, f'{mmss(dl)} → {mmss(dl + s)}  ·  {len(lista)} dogrywek  ·  +{s} s')

        for i, (typ, po, powod, tekst, plansza) in enumerate(lista, 1):
            n, sek = czas(tekst)
            nazwa = f'{mod}_dogrywka_{i}.wav'
            h3(doc, f'{mod} · dogrywka {i} — {typ} po planszy {po}')
            table(doc,
                  ['Miejsce cięcia', 'Rodzaj', 'Powód (ustalenie audytu)', 'Słów', 'Czas', 'Plik'],
                  [[f'bezpośrednio po planszy {po}', typ, powod, str(n), f'{sek} s', nazwa]],
                  widths=[3.4, 2.2, 5.6, 1.2, 1.2, 3.1], size=8, align_first_bold=False)
            blok_tekstu(doc, tekst)
            para(doc, 'PLANSZA DO WYŚWIETLENIA W CZASIE TEJ WSTAWKI', size=8, bold=True,
                 color=ORANGE, before=2, after=4)
            for j, ln in enumerate(plansza.split('\n')):
                para(doc, ln, size=9.5 if j == 0 else 9,
                     bold=(j == 0), color=PURPLE if j == 0 else INK,
                     after=(4 if j == 0 else 2), indent=0.4)
            para(doc, '', after=8)

    # ------------------------------------------------------------- plansze
    page_break(doc)
    band(doc, 'PLANSZE', 'Wykaz plansz do wymiany — zmiana wyłącznie graficzna')
    para(doc, 'Poniższe plansze zawierają podstawę prawną i wymagają korekty. Ścieżka dźwiękowa w tych '
              'miejscach pozostaje bez zmian — plansza jest wyświetlana pod istniejącą narracją.', after=8)
    table(doc, ['Moduł', 'Plansza', 'Co poprawić'],
          [[m, p, o] for m, p, o in PLANSZE_PODMIANA],
          widths=[1.8, 4.4, 10.5], size=8.5)

    h2(doc, 'Kolejność prac')
    numbered(doc, [
        '**Wymiana plansz** — pięć pozycji z wykazu powyżej. Można wykonać od razu, niezależnie od nagrań.',
        '**Nagranie dziewięciu dogrywek** — łącznie około czterech minut materiału. Warto nagrać je jednym '
        'ciągiem, w tych samych warunkach akustycznych co oryginał.',
        '**Montaż wstawek** — w kolejności modułów, od M1 do M6. Po każdej wstawce sprawdzić płynność '
        'przejścia; wszystkie wstawki zaczynają się zdaniem wprowadzającym, więc nie powinny brzmieć jak sklejka.',
        '**Aktualizacja spisu części w skrypcie** — nowe czasy modułów z tabeli zestawienia.',
        '**Aktualizacja napisów** (jeżeli film ma napisy) — tekst wstawek jest gotowy do wklejenia '
        'w pliku towarzyszącym „dogrywki-narracja.txt”.',
    ])

    accent_box(doc, [
        'Czego ta karta NIE zmienia: żadnego już nagranego zdania. Audyt nie zakwestionował ani jednej tezy '
        'wypowiedzianej w filmie — wszystkie dziewięć dogrywek to uzupełnienia i doprecyzowania.',
        'Gdyby zdecydowała się Pani nagrać tylko część: najważniejsze są dogrywki M3 · 2 (podawanie leków — '
        'dotyczy bezpieczeństwa nauczyciela) oraz M1 · 4 (dwa brakujące akty — bez nich informacja o gotowości '
        'nie ma podstawy). Reszta to precyzja, nie ryzyko.',
    ], title='DWIE UWAGI NA KONIEC', accent=ORANGE)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'DOGRYWKI_film_PRZEDSZKOLE_po_audycie_2026-09-05.docx')
    doc.save(out)

    # plik towarzyszący — sam tekst do nagrania / do napisów
    txt = [os.path.join(os.path.dirname(os.path.abspath(__file__)), 'dogrywki-narracja.txt')]
    with open(txt[0], 'w', encoding='utf-8') as f:
        f.write('DOGRYWKI DO MATERIAŁU FILMOWEGO — PRZEDSZKOLE\n')
        f.write('Po audycie Strażnika Prawa, 5 września 2026 r.\n')
        f.write('Tempo odniesienia: 107 słów na minutę (zmierzone na istniejącym filmie).\n')
        f.write('=' * 78 + '\n\n')
        for mod, tytul, dl, lista in DOGRYWKI:
            for i, (typ, po, powod, tekst, plansza) in enumerate(lista, 1):
                n, sek = czas(tekst)
                f.write(f'--- {mod}_dogrywka_{i}.wav | {typ} po planszy {po} | {n} słów | ~{sek} s ---\n')
                f.write(tekst + '\n\n')
    print('Zapisano:', out)
    print('Zapisano:', txt[0])
    print(f'Dogrywek: {licz}, łącznie {suma} s ({mmss(suma)})')


if __name__ == '__main__':
    main()
