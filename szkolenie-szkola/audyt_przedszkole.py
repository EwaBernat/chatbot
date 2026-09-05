# -*- coding: utf-8 -*-
"""
Raport Strażnika Prawa — audyt podstaw prawnych skryptu przedszkolnego
(EduPlaner 2026 · PCTP). Korzysta z tej samej warstwy składu co skrypt szkolny.
"""
import os
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from build_skrypt_szkola import (Document, PURPLE, ORANGE, INK, GREY, LILAC, BLUE, BORDER,
                                 _sub, run, para, hairline, page_break, accent_box, band,
                                 h2, h3, table, numbered, steps, rich)


def build_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Arial'
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    st.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = Cm(1.9); sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)

    hp = sec.header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(3)
    bdr = _sub(hp._p.get_or_add_pPr(), 'w:pBdr')
    _sub(bdr, 'w:bottom', **{'w:val': 'single', 'w:sz': '4', 'w:space': '4', 'w:color': BORDER})
    run(hp, 'PCTP · EduPlaner 2026', size=7.5, bold=True, color=ORANGE)
    run(hp, '     Raport Strażnika Prawa · audyt skryptu przedszkolnego · 5 września 2026 r.',
        size=7.5, color=GREY)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(3)
    bdr = _sub(fp._p.get_or_add_pPr(), 'w:pBdr')
    _sub(bdr, 'w:top', **{'w:val': 'single', 'w:sz': '4', 'w:space': '4', 'w:color': BORDER})
    run(fp, 'Strona ', size=7.5, color=GREY)
    from docx.oxml import OxmlElement
    f1 = OxmlElement('w:fldSimple'); f1.set(qn('w:instr'), 'PAGE'); fp._p.append(f1)
    run(fp, ' z ', size=7.5, color=GREY)
    f2 = OxmlElement('w:fldSimple'); f2.set(qn('w:instr'), 'NUMPAGES'); fp._p.append(f2)
    run(fp, ' · audyt podstaw prawnych · skrypt przedszkolny', size=7.5, color=GREY)
    return doc


def main():
    doc = build_doc()

    para(doc, 'EDUPLANER 2026 · PCTP · STRAŻNIK PRAWA', size=9, bold=True, color=ORANGE, after=4)
    para(doc, 'Raport Strażnika Prawa', size=26, bold=True, color=PURPLE, after=2)
    para(doc, 'Audyt podstaw prawnych skryptu dla nauczycieli — przedszkole',
         size=13, bold=True, color=PURPLE, after=8)
    para(doc, 'Przedmiot: „Skrypt dla nauczycieli. Transkrypcja filmu · podstawy prawne · sposób przygotowania '
              'dokumentów”, dokumentacja przedszkolna, rok szkolny 2026/2027, sześć części, 53 minuty 57 sekund.',
         size=10, after=2)
    para(doc, 'Data audytu: 5 września 2026 r.   ·   Metoda: ta sama, którą zastosowano do skryptu szkolnego.',
         size=10, color=GREY, after=10)
    hairline(doc)

    accent_box(doc, [
        'Skrypt przedszkolny przeszedł audyt w stanie znacznie lepszym niż szkolny. '
        '[[Nie zawiera ani jednego nieaktualnego publikatora]] — wszystkie sześć cytowanych aktów wskazano '
        'w obowiązujących tekstach jednolitych albo w aktualnym publikatorze pierwotnym.',
        'Nie powtórzyły się w nim żadne z pięciu błędów wykrytych w skrypcie szkolnym. Powód jest prosty: '
        'skrypt przedszkolny nie wchodzi w te obszary (ocenianie, egzamin, nauczyciel współorganizujący, '
        'finansowanie), w których szkolny się pomylił.',
        'Audyt wykrył natomiast [[cztery sprostowania merytoryczne i pięć luk]] — miejsc, w których skrypt '
        'wymaga dokumentu albo czynności, nie podając jej podstawy prawnej.',
        'Nie znaleziono w skrypcie żadnego z trzech „krytycznych rozjazdów” odnotowanych w rejestrze placówki '
        '(Dz.U. 2023 poz. 2029 — prawo geologiczne i górnicze; Dz.U. 2024 poz. 1640 — ERTMS kolejowy; '
        'Dz.U. 2026 poz. 770). Do skryptu nie przeniknęły.',
    ], title='WYNIK W CZTERECH ZDANIACH')

    h2(doc, 'Zestawienie ilościowe')
    table(doc,
          ['Kategoria', 'Skrypt przedszkolny', 'Skrypt szkolny (dla porównania)'],
          [['Cytowane akty prawne', '6 + RODO + ICF', '12 + RODO + ICF + komunikat CKE'],
           ['Nieaktualne publikatory', '[[0]]', '3 (ocenianie, nauczanie indywidualne, finansowanie)'],
           ['Błędy merytoryczne w treści przepisu', '[[0]]', '2 (nauczyciel współorganizujący, art. 8 ust. 16–17)'],
           ['Sprostowania merytoryczne (nieścisłości)', '4', '—'],
           ['Luki — czynność bez wskazanej podstawy', '5', '1 (podstawa programowa)'],
           ['Druki bez podstawy w rozporządzeniu, opisane uczciwie', '[[1 — metryczka]]', '2 (metryczka, PWES)'],
           ['Rozjazdy z rejestru placówki obecne w dokumencie', '[[0]]', '0']],
          widths=[6.2, 4.6, 5.9], size=8.5)

    # ------------------------------------------------------------------ A
    page_break(doc)
    band(doc, 'CZĘŚĆ A', 'Sprostowania merytoryczne — cztery ustalenia', 'naniesione na wydanie 2')
    table(doc,
          ['#', 'Gdzie', 'Zapis wydania 1', 'Ustalenie audytu', 'Co dopisano na niebiesko'],
          [['A1', 'Cz. 6 — podstawa prawna części',
            '„§ 6 rozporządzenia (…): ust. 1 pkt 1–8 (zawartość programu), ust. 4 (program po ocenie), '
            'ust. 9–11 (ocena co najmniej dwa razy w roku, prawa rodziców)”.',
            'Trzy wskazane ustępy obejmują cztery odrębne materie: okresową ocenę, prawo rodziców do udziału '
            'w spotkaniach, obowiązek pisemnego zawiadomienia o terminie oraz prawo do kopii oceny i programu. '
            'Numeracji nie udało się potwierdzić w ogłoszonym tekście — dostępne źródła wtórne podają '
            'rozbieżnie ust. 9 albo ust. 10 dla oceny okresowej. Audyt nie rozstrzyga tego domysłem.',
            'Uwaga redakcyjna: numerację sprawdzić w ogłoszonym tekście jednolitym; bezpieczny zapis '
            'w dokumencie dziecka to „§ 6 rozporządzenia” bez numeru ustępu.'],
           ['A2', 'Cz. 4 — podstawa prawna części',
            '„Obserwacja pedagogiczna i analiza gotowości szkolnej — rozporządzenie o dokumentacji '
            '(t.j. Dz.U. 2024 poz. 50) jako dokumentacja badań i czynności uzupełniających”.',
            'Przypisanie częściowo błędne. Rozporządzenie o dokumentacji jest podstawą dla arkuszy obserwacji '
            'jako dokumentacji badań i czynności uzupełniających — i tu zapis jest trafny. Ale sam obowiązek '
            'prowadzenia obserwacji i przeprowadzenia analizy gotowości wynika z zadań przedszkola w podstawie '
            'programowej, a wzór informacji o gotowości — z rozporządzenia o świadectwach.',
            'Sprostowanie rozdzielające trzy podstawy: dokumentacja (arkusze), podstawa programowa (obowiązek), '
            'świadectwa (wzór druku).'],
           ['A3', 'Cz. 3 — akapit 16 i instrukcja „krok po kroku”, sekcja VI',
            '„W polu zalecenia dotyczące postępowania zapisujemy konkretnie, kto podaje lek, '
            'na jakiej podstawie (…)”.',
            'Sformułowanie sugeruje istnienie podstawy prawnej, której dla przedszkola nie ma. Ustawa '
            'o opiece zdrowotnej nad uczniami (Dz.U. 2019 poz. 1078) dotyczy uczniów, nie dzieci objętych '
            'wychowaniem przedszkolnym. Podawanie leków opiera się na upoważnieniu rodziców, dobrowolnej '
            'pisemnej zgodzie nauczyciela, przeszkoleniu i procedurze dyrektora. Nauczyciel nie ma obowiązku '
            'podania leku — to istotne dla jego bezpieczeństwa.',
            'Wyjaśnienie, że „na jakiej podstawie” oznacza trzy dokumenty, a nie przepis ustawy; '
            'w podstawie prawnej części 3 dodano osobny punkt.'],
           ['A4', 'Cz. 1 — akapit 21',
            '„Warto zapamiętać: pozycja tysiąc pięćset siedemdziesiąt osiem to kształcenie specjalne, '
            'a pozycja tysiąc pięćset dziewięćdziesiąt jeden to pomoc psychologiczno-pedagogiczna”.',
            'Zdanie jest prawdziwe, ale postawione dwa zdania po regule „cytujemy tekst jednolity, nie '
            'publikator pierwotny”. To pułapka dydaktyczna: nauczyciel zapamiętuje właśnie te dwie liczby '
            'i wpisuje je do druku. Rejestr placówki pokazuje, że taki błąd realnie się zdarzał.',
            'Zastrzeżenie, że 1578 i 1591 służą do rozpoznania aktu, a nie do cytowania; w dokumencie '
            'dziecka wpisujemy 1309 i 1798.']],
          widths=[0.8, 2.9, 4.0, 5.2, 3.8], size=7.5)

    # ------------------------------------------------------------------ B
    page_break(doc)
    band(doc, 'CZĘŚĆ B', 'Zapisy zweryfikowane i potwierdzone jako prawidłowe', 'bez zmian')
    para(doc, 'Ta część jest w audycie równie ważna jak lista błędów. Każdy z poniższych zapisów sprawdzono '
              'u źródła i każdy pozostaje bez zmian. Gdyby ktoś zakwestionował je „z pamięci”, poniższa tabela '
              'jest odpowiedzią.', after=8)
    table(doc,
          ['#', 'Zapis wydania 1', 'Wynik weryfikacji'],
          [['B1', 'Nowa podstawa programowa wychowania przedszkolnego — rozporządzenie ME z 11 marca 2026 r. '
            '(Dz.U. 2026 poz. 378); dziewięć obszarów: społeczny, osobisty, językowy, matematyczny, '
            'przyrodniczy, techniczny, cyfrowy, artystyczny i ruchowy; obowiązuje od 1.09.2026 i obejmuje '
            'od razu wszystkie grupy wiekowe.',
            'POTWIERDZONE w całości — łącznie z listą dziewięciu obszarów i z tym, że w przedszkolu '
            'nowa podstawa wchodzi jednorazowo dla wszystkich grup. (W szkole podstawowej wchodzi etapami, '
            'od klas I i IV — to różnica, którą skrypt szkolny musiał uwzględnić, a przedszkolny nie.)'],
           ['B2', 'Rozporządzenie ME z 2 marca 2026 r. w sprawie orzeczeń i opinii (Dz.U. 2026 poz. 428); '
            '§ 7 ust. 6–7 i § 8 od 1.09.2026; § 7 ust. 3 — opinia w terminie 10 dni od dnia otrzymania prośby '
            'przez dyrektora; kopia dla rodziców; wejście w życie 14 kwietnia.',
            'POTWIERDZONE w całości. Akt istnieje, obowiązuje od 14.04.2026 i z tym dniem uchylił '
            'rozporządzenie z 7 września 2017 r. (t.j. Dz.U. 2023 poz. 2061). Termin 10 dni i wymóg opisu '
            'w kategoriach aktywności i uczestniczenia według ICF są w tekście aktu.'],
           ['B3', 'Prawo oświatowe — t.j. Dz.U. 2026 poz. 820.',
            'POTWIERDZONE. Wcześniejsze teksty jednolite Dz.U. 2024 poz. 737 (wygasł 31.07.2025) i Dz.U. 2025 '
            'poz. 1043 (wygasł 21.06.2026) nie są już właściwe.'],
           ['B4', 'Kształcenie specjalne — t.j. Dz.U. 2020 poz. 1309. Pomoc psychologiczno-pedagogiczna — '
            't.j. Dz.U. 2023 poz. 1798. Dokumentacja przebiegu nauczania — t.j. Dz.U. 2024 poz. 50.',
            'POTWIERDZONE — wszystkie trzy w obowiązujących tekstach jednolitych, zgodnie z rejestrem placówki.'],
           ['B5', 'Terminy: IPET do 30 września albo 30 dni od złożenia orzeczenia w przedszkolu; '
            'wielospecjalistyczna ocena co najmniej dwa razy w roku szkolnym.',
            'POTWIERDZONE co do treści obowiązku (zastrzeżenie do numeracji ustępów — ustalenie A1).'],
           ['B6', 'Metryczka dziecka opisana jako narzędzie wewnętrzne, wprowadzane zarządzeniem dyrektora; '
            'reguły przekierowania, progi i karta decyzyjna opisane jako decyzja rady pedagogicznej; '
            'wprost napisane, że cele SMART nie są nazwane w rozporządzeniu.',
            'POTWIERDZONE i wskazane jako wzorzec. To jest dokładnie ten sposób opisywania druków własnych, '
            'którego zabrakło w skrypcie szkolnym przy planie wsparcia (PWES) — i który tam trzeba było '
            'wprowadzić w audycie.'],
           ['B7', 'Publikatory pierwotne 1578 (kształcenie specjalne) i 1591 (pomoc pp) jako dwa różne akty '
            'z tej samej daty.',
            'POTWIERDZONE co do faktu (zastrzeżenie dydaktyczne — ustalenie A4).']],
          widths=[0.8, 7.3, 8.6], size=7.5)

    # ------------------------------------------------------------------ C
    page_break(doc)
    band(doc, 'CZĘŚĆ C', 'Luki — czynność wymagana przez skrypt bez wskazanej podstawy', 'pięć ustaleń')
    table(doc,
          ['#', 'Luka', 'Podstawa, której zabrakło', 'Status'],
          [['C1', 'Informacja o gotowości dziecka do podjęcia nauki w szkole podstawowej pojawia się w skrypcie '
            'trzykrotnie — w części 3, w części 6 i w kalendarzu — i ani razu nie ma wskazanej podstawy '
            'ani informacji, że sporządza się ją na urzędowym wzorze.',
            'Rozporządzenie MEiN z 7 czerwca 2023 r. w sprawie świadectw, dyplomów państwowych i innych druków '
            '(Dz.U. 2023 poz. 1120, z późn. zm.) — określa wzór informacji. Termin: do końca kwietnia.',
            '[[Dopisane w cz. 1, cz. 6 i w kalendarzu]]'],
           ['C2', 'Sekcja II metryczki (grupa, godziny pobytu, posiłki, roczne obowiązkowe przygotowanie '
            'przedszkolne) opiera się na organizacji przedszkola, której akt nie jest w skrypcie wymieniony.',
            'Rozporządzenie MEN z 28 lutego 2019 r. w sprawie szczegółowej organizacji publicznych szkół '
            'i publicznych przedszkoli (t.j. Dz.U. 2023 poz. 2736, z późn. zm.).',
            '[[Dopisane w cz. 1]]'],
           ['C3', 'Podana podstawa programowa nie uwzględnia aktu zmieniającego.',
            'Rozporządzenie Dz.U. 2026 poz. 958, również od 1.09.2026 — zmiana nie dotyczy załącznika nr 1, '
            'czyli wychowania przedszkolnego. Warto to napisać wprost, żeby nikt nie szukał zmian '
            'w swoim załączniku.',
            '[[Dopisane w cz. 1]]'],
           ['C4', 'Skrypt nie informuje, że nowe rozporządzenie o orzekaniu uchyliło poprzednie.',
            'Dz.U. 2026 poz. 428 uchyliło z dniem 14.04.2026 rozporządzenie MEN z 7 września 2017 r. '
            '(t.j. Dz.U. 2023 poz. 2061). Bez tej informacji nauczyciel, który zna stary akt, może go dalej '
            'cytować w dobrej wierze.',
            '[[Dopisane w cz. 1]]'],
           ['C5', 'Trzy akty obecne w rejestrze placówki nie występują w skrypcie, choć dotyczą materii, '
            'którą skrypt omawia: teczki dziecka, bezpieczeństwa i procedur.',
            'Ustawa z 13 maja 2016 r. o przeciwdziałaniu zagrożeniom przestępczością na tle seksualnym '
            'i ochronie małoletnich (t.j. Dz.U. 2026 poz. 110) — standardy ochrony małoletnich; '
            'rozporządzenie MENiS z 31 grudnia 2002 r. w sprawie bezpieczeństwa i higieny '
            '(t.j. Dz.U. 2020 poz. 1604); ustawa z 10 maja 2018 r. o ochronie danych osobowych '
            '(t.j. Dz.U. 2019 poz. 1781).',
            'DO DECYZJI AUTORKI — poza zakresem szkolenia dokumentacyjnego, ale warte jednego zdania '
            'w części 1.']],
          widths=[0.8, 5.6, 6.9, 3.4], size=7.5)

    # ------------------------------------------------------------------ D
    h2(doc, 'Ustalenie wspólne dla obu skryptów')
    accent_box(doc, [
        'Zastrzeżenie z ustalenia A1 — numeracja ustępów § 6 rozporządzenia o kształceniu specjalnym — '
        'dotyczy także skryptu szkolnego, który w wydaniu 2 cytuje „§ 6 ust. 9–12”. W obu dokumentach '
        'materia jest wskazana prawidłowo, niepewna pozostaje wyłącznie numeracja ustępów.',
        'Rekomendacja jednolita dla obu skryptów i dla wszystkich druków: w dokumencie dziecka piszemy '
        '„§ 6 rozporządzenia MEN z 9 sierpnia 2017 r. (t.j. Dz.U. 2020 poz. 1309)” i opisujemy obowiązek '
        'słowami. Numer ustępu dodajemy dopiero po sprawdzeniu w ogłoszonym tekście jednolitym — i wtedy '
        'wpisujemy go do rejestru przepisów wraz z datą sprawdzenia.',
    ], title='NUMERACJA USTĘPÓW § 6 — DOTYCZY TAKŻE SKRYPTU SZKOLNEGO', accent=ORANGE)

    h2(doc, 'Czego audyt nie obejmował')
    steps(doc, [
        '**Treści narzędzia KPOF** — liczby twierdzeń, skali, progów i sposobu liczenia średniej. To narzędzie '
        'własne placówki i skrypt uczciwie tak je opisuje.',
        '**Reguł przekierowania i karty decyzyjnej** — skrypt sam pisze, że nie wynikają wprost z przepisu, '
        'lecz są decyzją rady pedagogicznej wpisaną do procedury. To zapis wzorcowy.',
        '**Numeracji sekcji metryczki** (sekcja VI, VII, XI) — to układ druku własnego, nie przepisu.',
        '**Zgodności skryptu z kodem aplikacji** — rejestr przepisów placówki odnotowuje trzy rozjazdy '
        'naprawione w kodzie 22 sierpnia 2026 r.; audyt sprawdził tylko, że nie ma ich w skrypcie.',
    ])

    h2(doc, 'Zalecenia po audycie')
    numbered(doc, [
        '**Przyjąć wydanie 2 skryptu** z naniesionymi na niebiesko uzupełnieniami — nic nie zostało usunięte '
        'ani przeredagowane wbrew autorce, wszystkie zmiany są dopowiedzeniami.',
        '**Rozstrzygnąć numerację ustępów § 6** w ogłoszonym tekście jednolitym i wpisać wynik do rejestru '
        'przepisów. Do tego czasu w druku dziecka stosować zapis bez numeru ustępu.',
        '**Dopisać do rejestru przepisów placówki dwa akty**, których w nim brakuje po stronie przedszkolnej: '
        'rozporządzenie o świadectwach (Dz.U. 2023 poz. 1120) i rozporządzenie o organizacji przedszkoli '
        '(t.j. Dz.U. 2023 poz. 2736, z późn. zm.).',
        '**Zdecydować w sprawie ustalenia C5** — czy standardy ochrony małoletnich, BHP i krajowa ustawa '
        'o ochronie danych mają wejść do części 1, czy pozostać poza zakresem szkolenia dokumentacyjnego.',
        '**Ustalić rytm przeglądu.** Rejestr placówki był weryfikowany 21 sierpnia 2026 r., audyt skryptów '
        '5 września 2026 r. Kolejny przegląd przed styczniową oceną wielospecjalistyczną zamknie cykl.',
    ])

    hairline(doc)
    para(doc, 'Raport sporządzono metodą Strażnika Prawa: każde ustalenie ma wskazane źródło i datę '
              'sprawdzenia, a tam, gdzie źródła nie udało się osiągnąć, audyt mówi o tym wprost zamiast '
              'zgadywać. Ustalenie A1 jest tego przykładem.',
         size=8.5, italic=True, color=GREY)

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'AUDYT_Skryptu_PRZEDSZKOLNEGO_Straznik_Prawa_2026-09-05.docx')
    doc.save(out)
    print('Zapisano:', out)


if __name__ == '__main__':
    main()
