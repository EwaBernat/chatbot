# -*- coding: utf-8 -*-
"""
Generator dokumentu: "Skrypt dla nauczycieli — szkoła podstawowa"
EduPlaner 2026 · PCTP · szkolenie rady pedagogicznej.

Odwzorowuje układ i markę skryptu przedszkolnego:
Arial, fiolet #2D1B69, pomarańcz #E8450A, tła #F2F0F7 / #F7F6FA.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE = "2D1B69"
ORANGE = "E8450A"
INK    = "1A1A1A"
GREY   = "8A8A8A"
LILAC  = "F2F0F7"
LILAC2 = "F7F6FA"
BORDER = "D6D1E4"
BLUE   = "0B4F9E"   # kolor korekty — audyt Strażnika Prawa 05.09.2026

# ---------------------------------------------------------------- low level

# Kolejność elementów potomnych wymagana przez schemat OOXML — Word i LibreOffice
# odrzucają plik, w którym właściwości są zapisane w innej kolejności.
_PPR = ['w:pStyle', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore', 'w:framePr', 'w:widowControl',
        'w:numPr', 'w:suppressLineNumbers', 'w:pBdr', 'w:shd', 'w:tabs', 'w:suppressAutoHyphens',
        'w:kinsoku', 'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing',
        'w:mirrorIndents', 'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment',
        'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr', 'w:pPrChange']
_RPR = ['w:rStyle', 'w:rFonts', 'w:b', 'w:bCs', 'w:i', 'w:iCs', 'w:caps', 'w:smallCaps', 'w:strike',
        'w:dstrike', 'w:outline', 'w:shadow', 'w:emboss', 'w:imprint', 'w:noProof', 'w:snapToGrid',
        'w:vanish', 'w:webHidden', 'w:color', 'w:spacing', 'w:w', 'w:kern', 'w:position', 'w:sz',
        'w:szCs', 'w:highlight', 'w:u', 'w:effect', 'w:bdr', 'w:shd', 'w:fitText', 'w:vertAlign',
        'w:rtl', 'w:cs', 'w:em', 'w:lang', 'w:eastAsianLayout', 'w:specVanish', 'w:oMath']
_TBLPR = ['w:tblStyle', 'w:tblpPr', 'w:tblOverlap', 'w:bidiVisual', 'w:tblStyleRowBandSize',
          'w:tblStyleColBandSize', 'w:tblW', 'w:jc', 'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders',
          'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook', 'w:tblCaption', 'w:tblDescription',
          'w:tblPrChange']
_TCPR = ['w:cnfStyle', 'w:tcW', 'w:gridSpan', 'w:hMerge', 'w:vMerge', 'w:tcBorders', 'w:shd',
         'w:noWrap', 'w:tcMar', 'w:textDirection', 'w:tcFitText', 'w:vAlign', 'w:hideMark',
         'w:headers', 'w:tcPrChange']
_ORDERS = {'pPr': _PPR, 'rPr': _RPR, 'tblPr': _TBLPR, 'tcPr': _TCPR}


def _local(el):
    return 'w:' + el.tag.split('}')[-1]


def _sub(parent, tag, **attrs):
    """Wstawia element właściwości we właściwym miejscu; istniejący element
    o tym samym tagu uzupełnia zamiast go duplikować."""
    order = _ORDERS.get(parent.tag.split('}')[-1])
    if order is not None:
        for child in parent:
            if _local(child) == tag:
                for k, v in attrs.items():
                    child.set(qn(k), str(v))
                return child
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    if order is not None and tag in order:
        idx = order.index(tag)
        for child in parent:
            name = _local(child)
            if name in order and order.index(name) > idx:
                child.addprevious(el)
                return el
    parent.append(el)
    return el

def shade(el, fill):
    pr = el.get_or_add_pPr() if el.tag.endswith('}p') else el
    _sub(pr, 'w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': fill})

def cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    _sub(tcPr, 'w:shd', **{'w:val': 'clear', 'w:color': 'auto', 'w:fill': fill})

def cell_margins(cell, top=80, left=110, bottom=80, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = _sub(tcPr, 'w:tcMar')
    for tag, v in (('w:top', top), ('w:left', left), ('w:bottom', bottom), ('w:right', right)):
        _sub(mar, tag, **{'w:w': v, 'w:type': 'dxa'})

def run(p, text, size=10, bold=False, italic=False, color=INK, caps=False, space_after=None):
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    if caps:
        rPr = r._element.get_or_add_rPr()
        _sub(rPr, 'w:caps', **{'w:val': '1'})
    rPr = r._element.get_or_add_rPr()
    _sub(rPr, 'w:rFonts', **{'w:ascii': 'Arial', 'w:hAnsi': 'Arial',
                             'w:cs': 'Arial', 'w:eastAsia': 'Arial'})
    return r

def rich(p, text, size=10, bold=False, italic=False, color=INK):
    """Tekst z korektami: fragment w [[...]] renderuje się na niebiesko (audyt)."""
    parts = re.split(r'\[\[(.*?)\]\]', text, flags=re.S)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:
            run(p, part, size=size, bold=True, italic=italic, color=BLUE)
        else:
            run(p, part, size=size, bold=bold, italic=italic, color=color)


def para(doc_or_cell, text='', size=10, bold=False, italic=False, color=INK,
         before=0, after=6, align=None, indent=None, hanging=None, fill=None,
         line=264, keep=False):
    p = doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pPr = p._p.get_or_add_pPr()
        _sub(pPr, 'w:spacing', **{'w:line': line, 'w:lineRule': 'auto'})
    if align is not None:
        p.alignment = align
    if indent is not None:
        pf.left_indent = Cm(indent)
    if hanging is not None:
        pf.first_line_indent = Cm(-hanging)
    if fill:
        shade(p._p, fill)
    if keep:
        pf.keep_with_next = True
    if text:
        rich(p, text, size=size, bold=bold, italic=italic, color=color)
    return p

def hairline(doc, color=BORDER):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    bdr = _sub(pPr, 'w:pBdr')
    _sub(bdr, 'w:bottom', **{'w:val': 'single', 'w:sz': '4', 'w:space': '2', 'w:color': color})
    return p

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)

def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = _sub(tblPr, 'w:tblBorders')
    for tag in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
        _sub(borders, tag, **{'w:val': 'none', 'w:sz': '0', 'w:color': 'auto'})

def grid_borders(table, color=BORDER, sz='4'):
    tblPr = table._tbl.tblPr
    borders = _sub(tblPr, 'w:tblBorders')
    for tag in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
        _sub(borders, tag, **{'w:val': 'single', 'w:sz': sz, 'w:color': color})

def accent_box(doc, lines, title=None, fill=LILAC, accent=PURPLE, size=9):
    """Ramka z grubą lewą krawędzią — jak 'PODSTAWA PRAWNA CZĘŚCI'."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = t._tbl.tblPr
    borders = _sub(tblPr, 'w:tblBorders')
    _sub(borders, 'w:top',     **{'w:val': 'single', 'w:sz': '2',  'w:color': fill})
    _sub(borders, 'w:left',    **{'w:val': 'single', 'w:sz': '18', 'w:color': accent})
    _sub(borders, 'w:bottom',  **{'w:val': 'single', 'w:sz': '2',  'w:color': fill})
    _sub(borders, 'w:right',   **{'w:val': 'single', 'w:sz': '2',  'w:color': fill})
    _sub(borders, 'w:insideH', **{'w:val': 'none', 'w:sz': '0', 'w:color': fill})
    _sub(borders, 'w:insideV', **{'w:val': 'none', 'w:sz': '0', 'w:color': fill})
    c = t.cell(0, 0)
    cell_shade(c, fill)
    cell_margins(c, top=140, left=200, bottom=140, right=180)
    c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
    if title:
        para(c, title, size=8.5, bold=True, color=accent, after=5)
    for i, ln in enumerate(lines):
        p = para(c, '', size=size, after=(0 if i == len(lines) - 1 else 4))
        run(p, '· ', size=size, bold=True, color=accent)
        rich(p, ln, size=size, color=INK)
    para(doc, '', after=4)
    return t

def band(doc, tag, title, meta=''):
    """Pasek tytułowy części."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    shade(p._p, LILAC)
    run(p, '  %s  ' % tag, size=10, bold=True, color=ORANGE)
    run(p, '   %s' % title, size=13, bold=True, color=PURPLE)
    if meta:
        run(p, '   ·   %s' % meta, size=10, bold=True, color=ORANGE)
    return p

def h2(doc, text, color=PURPLE, size=11.5, before=12, after=6):
    return para(doc, text, size=size, bold=True, color=color,
                before=before, after=after, keep=True)

def h3(doc, text, color=ORANGE, size=9.5, before=10, after=4):
    p = para(doc, '', before=before, after=after, keep=True)
    run(p, text, size=size, bold=True, color=color, caps=True)
    return p

def narration(doc, items):
    """Numerowane akapity transkrypcji: 01, 02, ..."""
    for i, txt in enumerate(items, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(5)
        pf.left_indent = Cm(1.0)
        pf.first_line_indent = Cm(-1.0)
        pPr = p._p.get_or_add_pPr()
        _sub(pPr, 'w:spacing', **{'w:line': '264', 'w:lineRule': 'auto'})
        run(p, '%02d' % i, size=9, bold=True, color=ORANGE)
        run(p, '\t', size=10)
        rich(p, txt, size=10, color=INK)
        tabs = _sub(pPr, 'w:tabs')
        _sub(tabs, 'w:tab', **{'w:val': 'left', 'w:pos': '567'})

def steps(doc, items, marker='▸', color=PURPLE, size=10, after=5):
    for txt in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(after)
        pf.left_indent = Cm(0.65)
        pf.first_line_indent = Cm(-0.65)
        run(p, marker + '  ', size=size, bold=True, color=color)
        # obsługa pogrubionego wstępu: "Lead — reszta"
        if txt.startswith('**'):
            head, rest = txt[2:].split('**', 1)
            rich(p, head, size=size, bold=True, color=INK)
            rich(p, rest, size=size, color=INK)
        else:
            rich(p, txt, size=size, color=INK)

def numbered(doc, items, color=ORANGE, size=10):
    for i, txt in enumerate(items, 1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(5)
        pf.left_indent = Cm(0.8)
        pf.first_line_indent = Cm(-0.8)
        run(p, '%d.  ' % i, size=size, bold=True, color=color)
        if txt.startswith('**'):
            head, rest = txt[2:].split('**', 1)
            rich(p, head, size=size, bold=True, color=INK)
            rich(p, rest, size=size, color=INK)
        else:
            rich(p, txt, size=size, color=INK)

def table(doc, headers, rows, widths=None, size=8.5, head_fill=PURPLE,
          head_color='FFFFFF', zebra=True, align_first_bold=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    grid_borders(t, color='FFFFFF', sz='6')
    hdr = t.rows[0]
    for i, htxt in enumerate(headers):
        c = hdr.cells[i]
        cell_shade(c, head_fill)
        cell_margins(c)
        c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
        p = para(c, '', after=0, line=250)
        run(p, htxt, size=size, bold=True, color=head_color)
    for ri, rowvals in enumerate(rows):
        cells = t.add_row().cells
        fill = LILAC2 if (zebra and ri % 2 == 0) else 'FFFFFF'
        for ci, val in enumerate(rowvals):
            c = cells[ci]
            cell_shade(c, fill)
            cell_margins(c)
            c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
            for j, sub in enumerate(str(val).split('\n')):
                p = para(c, '', after=(0 if j == 0 else 0), line=250)
                rich(p, sub, size=size,
                     bold=(align_first_bold and ci == 0),
                     color=INK)
    if widths:
        t.autofit = False
        _sub(t._tbl.tblPr, 'w:tblLayout', **{'w:type': 'fixed'})
        _sub(t._tbl.tblPr, 'w:tblW',
             **{'w:w': int(sum(widths) * 567), 'w:type': 'dxa'})
        for i, w in enumerate(widths):
            t.columns[i].width = Cm(w)
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    para(doc, '', after=6)
    return t

def form_lines(doc, label, n=1, size=9, width_hint=''):
    """Etykieta + linie do wypełnienia ręcznego."""
    para(doc, label, size=size, bold=True, color=PURPLE, before=6, after=2, keep=True)
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        bdr = _sub(pPr, 'w:pBdr')
        _sub(bdr, 'w:bottom', **{'w:val': 'dotted', 'w:sz': '4', 'w:space': '2', 'w:color': BORDER})
        run(p, ' ', size=size)

def checkboxes(doc, items, cols=2, size=9):
    """Siatka pól wyboru ☐."""
    rows = (len(items) + cols - 1) // cols
    t = doc.add_table(rows=rows, cols=cols)
    no_borders(t)
    for i, it in enumerate(items):
        c = t.cell(i // cols, i % cols)
        cell_margins(c, top=30, bottom=30, left=60, right=60)
        c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
        p = para(c, '', after=0, line=240)
        run(p, '☐  ', size=size + 1, color=PURPLE)
        run(p, it, size=size, color=INK)
    for i in range(len(items), rows * cols):
        c = t.cell(i // cols, i % cols)
        c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)
        para(c, '', after=0)
    para(doc, '', after=6)
    return t

def form_header(doc, sygnatura, tytul, podtytul=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    shade(p._p, LILAC)
    run(p, '  %s  ' % sygnatura, size=9, bold=True, color=ORANGE)
    run(p, '   %s' % tytul, size=12, bold=True, color=PURPLE)
    if podtytul:
        para(doc, podtytul, size=8.5, italic=True, color=GREY, after=8)
    return p


# ================================================================ dokument

def build():
    doc = Document()

    st = doc.styles['Normal']
    st.font.name = 'Arial'
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    st.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = Cm(1.9); sec.bottom_margin = Cm(1.9)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)

    # nagłówek strony
    hp = sec.header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(3)
    pPr = hp._p.get_or_add_pPr()
    bdr = _sub(pPr, 'w:pBdr')
    _sub(bdr, 'w:bottom', **{'w:val': 'single', 'w:sz': '4', 'w:space': '4', 'w:color': BORDER})
    run(hp, 'PCTP · EduPlaner 2026', size=7.5, bold=True, color=ORANGE)
    run(hp, '     Skrypt dla nauczycieli · szkoła podstawowa — transkrypcja, podstawy prawne, przygotowanie dokumentów',
        size=7.5, color=GREY)

    # stopka strony z numeracją
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(3)
    pPr = fp._p.get_or_add_pPr()
    bdr = _sub(pPr, 'w:pBdr')
    _sub(bdr, 'w:top', **{'w:val': 'single', 'w:sz': '4', 'w:space': '4', 'w:color': BORDER})
    run(fp, 'Strona ', size=7.5, color=GREY)
    fr = fp.add_run(); fr.font.name = 'Arial'; fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor.from_string(GREY)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE'); fp._p.append(fld)
    run(fp, ' z ', size=7.5, color=GREY)
    fld2 = OxmlElement('w:fldSimple'); fld2.set(qn('w:instr'), 'NUMPAGES'); fp._p.append(fld2)
    run(fp, ' · skrypt dla nauczycieli · szkoła podstawowa', size=7.5, color=GREY)

    # ------------------------------------------------------ strona tytułowa
    para(doc, 'EDUPLANER 2026 · PCTP · SZKOLENIE RADY PEDAGOGICZNEJ',
         size=9, bold=True, color=ORANGE, after=4)
    para(doc, 'Skrypt dla nauczycieli', size=28, bold=True, color=PURPLE, after=2)
    para(doc, 'Szkoła podstawowa — kształcenie specjalne i pomoc psychologiczno-pedagogiczna',
         size=13, bold=True, color=PURPLE, after=8)
    para(doc, 'Transkrypcja filmu · podstawy prawne · uzasadnienie zmian w dokumentacji · sposób przygotowania druków',
         size=10, color=INK, after=2)
    para(doc, 'Dokumentacja szkolna · rok szkolny 2026/2027 · siedem części · 71 minut 40 sekund · 18 druków w załączniku',
         size=10, color=GREY, after=6)
    p = para(doc, '', after=10)
    run(p, 'WYDANIE 2 · po audycie podstaw prawnych — 5 września 2026 r.  ', size=9.5, bold=True, color=BLUE)
    run(p, 'Fragmenty zmienione w audycie oznaczono kolorem niebieskim. Wykaz zmian — Raport Strażnika Prawa '
            'na następnej stronie; wykaz aktów wraz z datą weryfikacji — załącznik F.', size=9, color=GREY)
    hairline(doc)

    h2(doc, 'JAK KORZYSTAĆ ZE SKRYPTU', before=6)
    para(doc, 'Skrypt odpowiada dokładnie temu, co słychać w filmie. Każda część ma trzy warstwy: podstawę prawną '
              '(żeby wiedzieć, z czego wynika obowiązek), pełną transkrypcję narracji (żeby wrócić do zdania, które umknęło) '
              'oraz instrukcję przygotowania dokumentu krok po kroku (żeby usiąść i zrobić). Numery akapitów transkrypcji '
              'odpowiadają kolejnym planszom w filmie.')
    para(doc, 'Skrypt można drukować w całości albo częściami — każda część zaczyna się od nowej strony. '
              'Załącznik C zawiera osiemnaście gotowych druków do powielenia; można je drukować niezależnie od skryptu.')

    h2(doc, 'NOTA REDAKCYJNA — ZASADA STRAŻNIKA PRAWA')
    accent_box(doc, [
        '[[Wszystkie publikatory w tym wydaniu zweryfikowano 5 września 2026 r. i wpisano do rejestru przepisów '
        '(załącznik F). Znak ⚑ pozostaje tylko przy jednym źródle, które zmienia się co roku i którego nie da się '
        'zweryfikować raz na zawsze: komunikat dyrektora Centralnej Komisji Egzaminacyjnej o sposobach dostosowania '
        'warunków egzaminu ósmoklasisty.]]',
        'W dokumencie ucznia cytujemy zawsze obowiązujący tekst jednolity, a nie pierwotny publikator '
        '(np. „t.j. Dz.U. 2020 poz. 1309”, a nie „Dz.U. 2017 poz. 1578”). '
        '[[Uwaga odwrotna, równie ważna: tekst jednolity także wygasa. Prawo oświatowe miało w latach 2024–2026 trzy '
        'kolejne teksty jednolite — Dz.U. 2024 poz. 737, Dz.U. 2025 poz. 1043 i obowiązujący Dz.U. 2026 poz. 820.]]',
        'Progi, reguły przekierowania i skale przeliczeniowe opisane w częściach 5 i 6 nie wynikają wprost z przepisu. '
        'To decyzja rady pedagogicznej, którą wpisujemy do procedury szkoły — po to, aby decyzja nie zależała od tego, '
        'kto danego dnia patrzy na arkusz.',
        'Dane ucznia użyte w przykładach (Zofia Lewandowska, klasa III A) są danymi z dokumentacji wzorcowej '
        'EduPlaner 2026 i służą wyłącznie celom szkoleniowym.',
    ], title='CZTERY ZASADY, KTÓRE OBOWIĄZUJĄ W CAŁYM SKRYPCIE')

    h2(doc, 'Spis części')
    table(doc,
          ['Część', 'Tytuł', 'Czas', 'Plik'],
          [['Część 1', 'Podstawa prawna — co obowiązuje w szkole od 1 września 2026 r.', '10:40', 'S1.mp4'],
           ['Część 2', 'Dlaczego szkoła musi zmienić dokumentację kształcenia specjalnego — uzasadnienie i zakres zmian', '12:20', 'S2.mp4'],
           ['Część 3', 'Obieg dokumentów w szkole — jak jeden wynika z drugiego', '6:10', 'S3.mp4'],
           ['Część 4', 'Metryczka i teczka ucznia — pierwszy dokument września', '6:30', 'S4.mp4'],
           ['Część 5', 'KSzOF — budowa narzędzia, skala, steny, liczenie wyniku, odczyt profilu', '11:20', 'S5.mp4'],
           ['Część 6', 'Obserwacja pogłębiona — ABC i FBA, profil sensoryczny, teoria umysłu, karta mowy', '10:30', 'S6.mp4'],
           ['Część 7', 'WOPF-SP, IPET, PWES, cele SMART, ewaluacja i opinia dla poradni', '14:10', 'S7.mp4'],
           ['Załącznik A', 'Kalendarz dokumentacji na rok szkolny', '—', '—'],
           ['Załącznik B', 'Język funkcjonalny — cztery pytania do każdego zdania', '—', '—'],
           ['Załącznik C', 'Osiemnaście druków kwestionariuszy i kart do powielenia', '—', '—'],
           ['Załącznik D', 'Wzór zarządzenia dyrektora i procedura dokumentacyjna', '—', '—'],
           ['Załącznik E', 'Checklista wdrożenia — 30 / 60 / 90 dni', '—', '—'],
           ['Załącznik F', '[[Rejestr przepisów cytowanych w skrypcie — status i data weryfikacji]]', '—', '—']],
          widths=[2.3, 10.5, 1.6, 1.6])

    para(doc, 'Placówka wzorcowa użyta w przykładach: Szkoła Podstawowa nr 7 im. Jana Brzechwy w Koszalinie. '
              'Uczennica: Zofia Lewandowska, klasa III A, orzeczenie o potrzebie kształcenia specjalnego '
              'nr PPP.4223.18.2026 z dnia 12.06.2026 r. wydane przez Poradnię Psychologiczno-Pedagogiczną w Koszalinie '
              'ze względu na niepełnosprawność sprzężoną. Numer programu: IPET/2026-2027/III A/07.',
          size=8.5, italic=True, color=GREY, before=4)

    return doc, sec


# ================================================================ CZĘŚĆ 1

def czesc_1(doc):
    page_break(doc)
    band(doc, 'CZĘŚĆ 1', 'Podstawa prawna — co obowiązuje w szkole od 1 września 2026 r.', '10:40  ·  S1.mp4')

    accent_box(doc, [
        'Prawo oświatowe — ustawa z 14 grudnia 2016 r., art. 1 pkt 5–7 (dostosowanie treści, metod i organizacji '
        'do możliwości psychofizycznych ucznia) oraz art. 127 (kształcenie specjalne) — '
        '[[t.j. Dz.U. 2026 poz. 820, z późn. zm.]] (poprzednie teksty jednolite Dz.U. 2024 poz. 737 i Dz.U. 2025 '
        'poz. 1043 wygasły).',
        'Kształcenie specjalne — rozporządzenie MEN z 9 sierpnia 2017 r., t.j. Dz.U. 2020 poz. 1309: § 5 (zajęcia '
        'rewalidacyjne), § 6 (WOPFU i IPET), [[§ 7 ust. 2 — obowiązek zatrudnienia nauczyciela współorganizującego '
        'wyłącznie przy autyzmie, w tym zespole Aspergera, oraz niepełnosprawnościach sprzężonych; § 7 ust. 3 — '
        'pozostałe przypadki, w tym niedostosowanie społeczne, tylko za zgodą organu prowadzącego]].',
        'Pomoc psychologiczno-pedagogiczna — rozporządzenie MEN z 9 sierpnia 2017 r., t.j. Dz.U. 2023 poz. 1798: '
        'formy pomocy w szkole oraz § 12 — zindywidualizowana ścieżka kształcenia.',
        'Orzeczenia i opinie zespołów orzekających — rozporządzenie ME z 2 marca 2026 r., Dz.U. 2026 poz. 428: '
        '§ 7 ust. 2–3 (opinia szkoły w terminie 10 dni), § 7 ust. 6–7 i § 8 — obowiązują od 1.09.2026. '
        '[[Akt zweryfikowany: obowiązuje od 14.04.2026 i z tym dniem uchylił rozporządzenie MEN z 7 września 2017 r. '
        '(t.j. Dz.U. 2023 poz. 2061), którego nie wolno już cytować jako obowiązującego.]]',
        'Dokumentacja przebiegu nauczania — rozporządzenie MEN z 25 sierpnia 2017 r., t.j. Dz.U. 2024 poz. 50: '
        'księga uczniów, dziennik lekcyjny, dzienniki zajęć, arkusze ocen, dokumentacja badań i czynności uzupełniających. '
        '[[Uwaga na rozjazd spotykany w drukach: Dz.U. 2024 poz. 1640 to akt o zupełnie innej treści.]]',
        'Ocenianie, klasyfikowanie i promowanie — rozporządzenie MEN z 22 lutego 2019 r., '
        '[[t.j. Dz.U. 2023 poz. 2572, z późn. zm. (m.in. Dz.U. 2025 poz. 778)]]: dostosowanie wymagań edukacyjnych, '
        'ocena zachowania ucznia z orzeczeniem, zwolnienia.',
        'Indywidualne nauczanie — rozporządzenie MEN z 9 sierpnia 2017 r., '
        '[[t.j. Dz.U. 2023 poz. 2468, z późn. zm. (zm. Dz.U. 2024 poz. 1714)]].',
        'Organizacja publicznych szkół — rozporządzenie MEN z 28 lutego 2019 r., '
        '[[t.j. Dz.U. 2023 poz. 2736, z późn. zm. (Dz.U. 2025 poz. 849, Dz.U. 2026 poz. 130, a od 1.09.2026 '
        'Dz.U. 2026 poz. 1090)]].',
        'Podstawa programowa — [[rozporządzenie ME z 11 marca 2026 r. w sprawie podstawy programowej wychowania '
        'przedszkolnego oraz podstawy programowej kształcenia ogólnego dla szkoły podstawowej, w tym dla uczniów '
        'z niepełnosprawnością intelektualną w stopniu umiarkowanym lub znacznym (Dz.U. 2026 poz. 378, zm. Dz.U. 2026 '
        'poz. 958). Wdrażana sukcesywnie: od 1.09.2026 w klasach I i IV, w kolejnych latach w następnych klasach; '
        'dla uczniów z niepełnosprawnością intelektualną w stopniu umiarkowanym lub znacznym — od 1.09.2026. '
        'Dla pozostałych klas nadal obowiązuje rozporządzenie MEN z 14 lutego 2017 r. (Dz.U. 2017 poz. 356, z późn. zm.).]]',
        'Zapewnianie dostępności osobom ze szczególnymi potrzebami — ustawa z 19 lipca 2019 r., '
        't.j. Dz.U. 2024 poz. 1411.',
        'Finansowanie zadań oświatowych — ustawa z 27 października 2017 r., [[t.j. Dz.U. 2026 poz. 650, art. 8 '
        'ust. 1: jednostka samorządu terytorialnego przeznacza na zadania wymagające stosowania specjalnej '
        'organizacji nauki i metod pracy środki w wysokości nie niższej niż kwota wynikająca z podziału części '
        'oświatowej subwencji ogólnej. Obowiązek adresowany jest do organu prowadzącego, nie do szkoły.]]',
        'RODO — rozporządzenie (UE) 2016/679, art. 5 ust. 1 lit. c (minimalizacja danych) oraz art. 9 '
        '(dane o zdrowiu jako dane szczególnej kategorii). ICF — Międzynarodowa Klasyfikacja Funkcjonowania, WHO 2001.',
    ], title='PODSTAWA PRAWNA CZĘŚCI')

    h3(doc, 'Transkrypcja narracji')
    narration(doc, [
        'Dzień dobry Państwu. Zanim otworzymy pierwszy druk, otwórzmy razem przepis.',
        'Ten moduł odpowiada na jedno, bardzo praktyczne pytanie. Z czego wynika każdy dokument, który wypełniamy '
        'w szkole od pierwszego września dwa tysiące dwudziestego szóstego roku? Bo kiedy znamy źródło rubryki, '
        'wypełniamy ją ze zrozumieniem. I wtedy dokumentacja przestaje być ciężarem, a staje się narzędziem.',
        'Zanim przejdziemy do przepisów, odpowiedzmy na pytanie, które w szkołach pada najczęściej. Mamy już teczki '
        'uczniów, mamy oceny wielospecjalistyczne i programy. Czy naprawdę musimy to zmieniać? Odpowiedź brzmi: '
        'nie zaczynamy od zera, ale aktualizujemy. Dotychczasowe arkusze, oceny i programy pozostają ważnym źródłem '
        'danych o uczniu. Zmienia się jednak język, w którym opisujemy funkcjonowanie, i sposób, w jaki jeden dokument '
        'zasila drugi. Uzasadnieniu tej zmiany poświęcimy cały moduł drugi.',
        'W szkole podstawowej pracujemy w ramach dwóch odrębnych ścieżek i ta różnica organizuje całą naszą dokumentację. '
        'Ścieżka pierwsza to kształcenie specjalne. Dotyczy ucznia, który posiada orzeczenie o potrzebie kształcenia '
        'specjalnego. Prowadzi do wielospecjalistycznej oceny poziomu funkcjonowania i do indywidualnego programu '
        'edukacyjno-terapeutycznego. Ścieżka druga to pomoc psychologiczno-pedagogiczna. Dotyczy ucznia z opinią poradni '
        'albo rozpoznany przez nauczycieli, bez orzeczenia. Prowadzi do planu wsparcia edukacyjnego i do oceny '
        'efektywności udzielanej pomocy. Ten sam kwestionariusz obserwacji obsługuje obie ścieżki. Rozgałęzienie '
        'następuje dopiero na etapie oceny.',
        'Zacznijmy od ustawy.',
        'Prawo oświatowe, ustawa z czternastego grudnia dwa tysiące szesnastego roku, w artykule pierwszym mówi, '
        'że system oświaty zapewnia dostosowanie treści, metod i organizacji nauczania do możliwości psychofizycznych '
        'uczniów. To jest zdanie fundamentalne. Nie mówi ono o obniżaniu wymagań. Mówi o dostosowaniu drogi, którą uczeń '
        'do tych wymagań dochodzi.',
        'Artykuł sto dwudziesty siódmy ustawy wprowadza kształcenie specjalne dla uczniów posiadających orzeczenie '
        'o potrzebie kształcenia specjalnego i odsyła do rozporządzenia, które jest dla nas najważniejsze.',
        'To rozporządzenie Ministra Edukacji Narodowej z dziewiątego sierpnia dwa tysiące siedemnastego roku w sprawie '
        'warunków organizowania kształcenia, wychowania i opieki dla dzieci i młodzieży niepełnosprawnych, niedostosowanych '
        'społecznie i zagrożonych niedostosowaniem społecznym. Obowiązujący tekst jednolity znajdziemy w Dzienniku Ustaw '
        'z dwa tysiące dwudziestego roku, pod pozycją tysiąc trzysta dziewięć. I właśnie tak, z tekstem jednolitym, '
        'cytujemy je w dokumentach ucznia. Pierwotny publikator, pozycja tysiąc pięćset siedemdziesiąt osiem, '
        'przestał być właściwym adresem.',
        'Z tego rozporządzenia wynikają cztery nasze obowiązki. Paragraf piąty: zajęcia rewalidacyjne dla ucznia '
        'z orzeczeniem. Paragraf szósty ustęp pierwszy: indywidualny program edukacyjno-terapeutyczny i jego osiem '
        'obowiązkowych elementów. Paragraf szósty ustęp czwarty: program opracowujemy po dokonaniu wielospecjalistycznej '
        'oceny poziomu funkcjonowania ucznia. I paragraf szósty ustęp dziewiąty: ocenę dokonujemy co najmniej dwa razy '
        'w roku szkolnym.',
        '[[Zapamiętajmy też paragraf siódmy, bo jest najczęściej cytowany błędnie. Ustęp drugi mówi, że '
        'w przedszkolach i szkołach ogólnodostępnych, w których kształceniem specjalnym objęci są uczniowie '
        'posiadający orzeczenie wydane ze względu na autyzm, w tym zespół Aspergera, albo niepełnosprawności '
        'sprzężone, zatrudnia się dodatkowo nauczyciela współorganizującego kształcenie. Tylko te dwie podstawy '
        'orzeczenia. To nie jest dobra wola dyrektora, to obowiązek. Ustęp trzeci mówi o wszystkich pozostałych '
        'przypadkach: przy innych niepełnosprawnościach, przy niedostosowaniu społecznym i przy zagrożeniu '
        'niedostosowaniem takiego nauczyciela można zatrudnić dodatkowo, ale wyłącznie za zgodą organu prowadzącego. '
        'Nie jest to obowiązek szkoły. Bezwzględny obowiązek przy niedostosowaniu społecznym dotyczy szkół '
        'specjalnych i młodzieżowych ośrodków wychowawczych, a nie szkoły ogólnodostępnej. Wymiar pracy nauczyciela '
        'współorganizującego zapisujemy w programie, w konkretnych godzinach i konkretnych zajęciach.]]',
        'Terminy programu są dwa. Do trzydziestego września dla ucznia, który rozpoczyna kształcenie z orzeczeniem '
        'w danym roku szkolnym. Albo trzydzieści dni od dnia złożenia w szkole orzeczenia, niezależnie od miesiąca. '
        'Termin liczymy od daty wpływu orzeczenia do szkoły, dlatego datę wpływu odnotowujemy w metryczce ucznia.',
        'Drugie rozporządzenie z tej samej daty dotyczy pomocy psychologiczno-pedagogicznej.',
        'Obowiązujący tekst jednolity to Dziennik Ustaw z dwa tysiące dwudziestego trzeciego roku, pozycja tysiąc '
        'siedemset dziewięćdziesiąt osiem. Warto zapamiętać: dwa różne akty, ta sama data dziewiątego sierpnia. '
        'Kształcenie specjalne to pozycja tysiąc trzysta dziewięć, pomoc psychologiczno-pedagogiczna to pozycja tysiąc '
        'siedemset dziewięćdziesiąt osiem.',
        'Co z niego bierzemy? Katalog form pomocy, który w szkole jest szerszy niż w przedszkolu. Zajęcia '
        'dydaktyczno-wyrównawcze. Zajęcia korekcyjno-kompensacyjne. Zajęcia rozwijające umiejętności uczenia się. '
        'Zajęcia rozwijające kompetencje emocjonalno-społeczne. Zajęcia logopedyczne. Zajęcia związane z wyborem '
        'kierunku kształcenia i zawodu. Porady, konsultacje i warsztaty. Oraz zindywidualizowana ścieżka kształcenia, '
        'opisana w paragrafie dwunastym.',
        'Zatrzymajmy się przy zindywidualizowanej ścieżce, bo to forma czysto szkolna i często mylona z nauczaniem '
        'indywidualnym. Ścieżka polega na tym, że uczeń realizuje część zajęć wspólnie z klasą, a część indywidualnie. '
        'Objęcie ucznia ścieżką wymaga opinii publicznej poradni, a wniosek składają rodzice. Ścieżki nie stosujemy '
        'wobec ucznia objętego kształceniem specjalnym ani indywidualnym nauczaniem. Nauczanie indywidualne to zupełnie '
        'inna instytucja: wymaga orzeczenia o potrzebie indywidualnego nauczania i osobnego rozporządzenia z dwa tysiące '
        'siedemnastego roku, pozycja tysiąc sześćset szesnaście.',
        'Trzeci akt jest w tym roku najważniejszy, bo to on zmienia sposób, w jaki opisujemy ucznia.',
        'To rozporządzenie Ministra Edukacji z drugiego marca dwa tysiące dwudziestego szóstego roku w sprawie orzeczeń '
        'i opinii wydawanych przez zespoły orzekające. Dziennik Ustaw z dwa tysiące dwudziestego szóstego roku, '
        'pozycja czterysta dwadzieścia osiem. Przepisy, które dotyczą bezpośrednio szkoły, czyli paragraf siódmy '
        'ustęp szósty i siódmy oraz paragraf ósmy, obowiązują od pierwszego września.',
        'Co to oznacza dla nas? Po pierwsze, ocena funkcjonalna ucznia staje się obowiązkowym etapem poprzedzającym '
        'wydanie orzeczenia. Po drugie, szkoła, na prośbę przewodniczącego zespołu orzekającego, wydaje opinię '
        'o funkcjonowaniu ucznia w szkole. Opinia opisuje trudności ucznia, ale równie starannie jego mocne strony '
        'i uzdolnienia rozpoznane przez nauczycieli i specjalistów. Przepis wprost mówi, że o uczniu piszemy także dobrze.',
        'Termin określa paragraf siódmy ustęp trzeci. Zacytujmy go w całości. Opinię, o której mowa w ustępie drugim, '
        'wydaje się w terminie dziesięciu dni od dnia otrzymania przez dyrektora prośby o jej wydanie. Dziesięć dni, '
        'liczonych od dnia, w którym prośba dotarła do dyrektora. Kopię opinii otrzymują rodzice ucznia.',
        'Podział ról jest jasny i wygodny. Formalną ocenę funkcjonalną sporządza zespół w poradni. My obserwujemy ucznia '
        'na lekcjach, na przerwach i w świetlicy, i opisujemy to, co widzimy. Nie stawiamy diagnoz. Dostarczamy '
        'rzetelnych, uporządkowanych danych.',
        'I tu jest sedno całego szkolenia. Dziesięć dni to bardzo mało, jeśli obserwację zaczynamy dopiero po wpłynięciu '
        'prośby. W szkole to jeszcze trudniejsze niż w przedszkolu, bo ucznia klasy szóstej uczy dziewięcioro nauczycieli '
        'i żaden z nich nie widzi go przez cały dzień. Dlatego kwestionariusz obserwacji wypełniamy we wrześniu, '
        'a nie w kwietniu. Nie po to, żeby leżał w segregatorze. Po to, żeby w dowolnym dniu roku móc odpowiedzieć '
        'poradni na podstawie danych, spokojnie i na czas.',
        'Czwarty akt porządkuje dokumentację przebiegu nauczania.',
        'Rozporządzenie z dwudziestego piątego sierpnia dwa tysiące siedemnastego roku, w obowiązującym tekście '
        'jednolitym. Wymienia księgę uczniów, dziennik lekcyjny, dzienniki zajęć, arkusze ocen oraz dokumentację badań '
        'i czynności uzupełniających prowadzonych przez nauczycieli, wychowawców i specjalistów. W tej ostatniej '
        'kategorii mieszczą się nasze arkusze obserwacji, karty ABC i profile sensoryczne. To jest ich podstawa prawna, '
        'i tak je opisujemy.',
        'Metryczka ucznia, którą omówimy w module czwartym, jest naszym narzędziem wewnętrznym. Wprowadzamy ją '
        'zarządzeniem dyrektora, bo porządkuje dane, które i tak gromadzimy z innych tytułów. Dzięki niej wszystko, '
        'co dotyczy ucznia, jest w jednym miejscu, a każdy dokument dziedziczy dane, zamiast je powielać.',
        'Piąty akt jest specyficznie szkolny i w przedszkolu nie ma odpowiednika. To rozporządzenie o ocenianiu, '
        'klasyfikowaniu i promowaniu uczniów, z dwudziestego drugiego lutego dwa tysiące dziewiętnastego roku.',
        'Z niego wynika, że nauczyciel dostosowuje wymagania edukacyjne do indywidualnych potrzeb ucznia. Dla ucznia '
        'z orzeczeniem podstawą dostosowania jest orzeczenie i program. Dla ucznia z opinią poradni — opinia. Dla ucznia '
        'objętego pomocą bez opinii — rozpoznanie dokonane przez nauczycieli i zapisane w dokumentacji. To ostatnie zdanie '
        'jest kluczowe: jeżeli dostosowujemy wymagania na podstawie własnego rozpoznania, to rozpoznanie musi być zapisane. '
        'Rozmowa w pokoju nauczycielskim nie jest podstawą.',
        'Z tego samego rozporządzenia wynika, że przy ustalaniu oceny zachowania ucznia posiadającego orzeczenie '
        'lub opinię uwzględnia się wpływ stwierdzonych zaburzeń lub zaburzeń rozwojowych na jego zachowanie. '
        'To zdanie ratuje wielu uczniów przed oceną naganną za objaw.',
        '[[Zanim przejdziemy dalej, jedno zdanie o podstawie programowej, bo w tym roku wchodzi ona etapami. '
        'Rozporządzenie Ministra Edukacji z jedenastego marca dwa tysiące dwudziestego szóstego roku, Dziennik Ustaw '
        'pozycja trzysta siedemdziesiąt osiem, wprowadza nową podstawę programową wychowania przedszkolnego '
        'i kształcenia ogólnego dla szkoły podstawowej. Od pierwszego września obowiązuje ona w klasach pierwszej '
        'i czwartej, a w latach następnych wchodzi w kolejnych klasach. Dla uczniów z niepełnosprawnością '
        'intelektualną w stopniu umiarkowanym lub znacznym obowiązuje od pierwszego września w całej szkole '
        'podstawowej. W pozostałych klasach nadal stosujemy podstawę z dwa tysiące siedemnastego roku, pozycja '
        'trzysta pięćdziesiąt sześć. Ma to bardzo praktyczny skutek dla naszej dokumentacji. W programie uczennicy '
        'klasy trzeciej, o której będziemy mówić w kolejnych modułach, powołujemy się na podstawę z dwa tysiące '
        'siedemnastego roku, a nie na nową. Nowa obejmie ją dopiero wtedy, gdy dojdzie do klasy objętej wdrożeniem.]]',
        'Szósty obszar to egzamin ósmoklasisty. Dostosowanie warunków i form przeprowadzania egzaminu przysługuje '
        'uczniowi na podstawie orzeczenia, opinii poradni albo — w niektórych przypadkach — pozytywnej opinii rady '
        'pedagogicznej. Szczegółowe sposoby dostosowania ogłasza corocznie komunikat dyrektora Centralnej Komisji '
        'Egzaminacyjnej. Terminy i katalog dostosowań sprawdzamy w komunikacie na dany rok. Dla nas płynie z tego '
        'jeden wniosek praktyczny: dostosowanie, którego nie ma w dokumentacji ucznia w listopadzie, nie pojawi się '
        'na egzaminie w maju.',
        'I trzy akty, które stoją nad wszystkimi. Ustawa o zapewnianiu dostępności osobom ze szczególnymi potrzebami. '
        '[[Ustawa o finansowaniu zadań oświatowych, której artykuł ósmy ustęp pierwszy nakłada na jednostkę samorządu '
        'terytorialnego obowiązek przeznaczenia na zadania wymagające stosowania specjalnej organizacji nauki i metod '
        'pracy środków w wysokości nie niższej niż kwota wynikająca z podziału części oświatowej subwencji ogólnej. '
        'Adresatem tego obowiązku jest organ prowadzący, a nie szkoła. Dla nas płynie z tego wniosek pośredni, '
        'ale ważny: to dokumentacja ucznia pokazuje, jakie zadania szkoła faktycznie realizuje, i to na jej podstawie '
        'organ prowadzący planuje oraz rozlicza te środki.]] Oraz rozporządzenie o ochronie danych osobowych. Pamiętajmy, że dane o zdrowiu i rozwoju '
        'ucznia to dane szczególnej kategorii. Teczka ucznia ma swoje bezpieczne miejsce, a klauzula informacyjna '
        'jest podpisana przez rodzica.',
        'W naszej szkole działa Strażnik Prawa. To nie jest osoba, która zna przepisy na pamięć. To osoba, która '
        'przy każdej decyzji zadaje jedno proste pytanie: z czego to wynika i gdzie to jest zapisane? Funkcja jest '
        'rotacyjna i pełni ją w ciągu roku każdy członek zespołu.',
        'Na zamknięcie tego modułu trzy zdania, które warto zabrać ze sobą.',
        'Każdy druk ma swój przepis, i my go znamy. Obserwacja wyprzedza pismo z poradni, bo wrześniowy arkusz daje nam '
        'spokój na cały rok. I na koniec: przepisy sprawdzamy w Dzienniku Ustaw, zanim wpiszemy je do dokumentu ucznia. '
        'Tak pracuje Strażnik Prawa. Zapraszam do modułu drugiego, w którym odpowiemy na pytanie, dlaczego zmiana '
        'dokumentacji jest w tym roku konieczna.',
    ])

    h2(doc, 'Co z którego aktu bierzemy — mapa dla wypełniającego')
    table(doc,
          ['Akt prawny', 'Co z niego wynika dla szkoły', 'Gdzie to widać w druku'],
          [['Prawo oświatowe\n[[t.j. Dz.U. 2026 poz. 820]]\nart. 1 pkt 5–7, art. 127',
            'Obowiązek dostosowania treści, metod i organizacji do możliwości ucznia; kształcenie specjalne na podstawie orzeczenia.',
            'Metryczka — sekcja „podstawa objęcia wsparciem”; IPET — część wstępna.'],
           ['Kształcenie specjalne\nt.j. Dz.U. 2020 poz. 1309, § 5–7',
            'Zajęcia rewalidacyjne; osiem elementów IPET; WOPFU co najmniej dwa razy w roku; nauczyciel '
            'współorganizujący — [[obowiązkowo tylko przy autyzmie/zespole Aspergera i niepełnosprawnościach '
            'sprzężonych (§ 7 ust. 2); w pozostałych przypadkach za zgodą organu prowadzącego (§ 7 ust. 3)]].',
            'WOPF-SP — całość; IPET — sekcje 2–7; karta kontrolna rozporządzenia.'],
           ['Pomoc pp\nt.j. Dz.U. 2023 poz. 1798',
            'Rozpoznawanie potrzeb przez nauczycieli, katalog form pomocy, zindywidualizowana ścieżka, ocena efektywności.',
            'PWES; karta oceny efektywności; KSzOF jako narzędzie rozpoznania.'],
           ['Orzeczenia i opinie\nDz.U. 2026 poz. 428\n[[obowiązuje od 14.04.2026]]',
            'Ocena funkcjonalna przed orzeczeniem; opinia szkoły o funkcjonowaniu ucznia w terminie 10 dni; język ICF.',
            'Opinia o funkcjonowaniu ucznia — siedem punktów; KSzOF (kody ICF przy twierdzeniach).'],
           ['Dokumentacja przebiegu nauczania\nt.j. Dz.U. 2024 poz. 50',
            'Arkusze obserwacji i karty specjalistów jako dokumentacja badań i czynności uzupełniających.',
            'Stopka każdego druku obserwacyjnego; wykaz dokumentacji w metryczce.'],
           ['Ocenianie\n[[t.j. Dz.U. 2023 poz. 2572, z późn. zm.]]',
            'Dostosowanie wymagań edukacyjnych; uwzględnienie zaburzeń przy ocenie zachowania; zwolnienia.',
            'IPET — sekcja dostosowań; karta dostosowań przedmiotowych (druk 11).'],
           ['Egzamin ósmoklasisty\nkomunikat dyrektora CKE ⚑',
            'Dostosowanie warunków i form egzaminu na podstawie orzeczenia, opinii albo opinii rady pedagogicznej.',
            'Karta dostosowań egzaminacyjnych (druk 12); protokół rady pedagogicznej.'],
           ['Finansowanie zadań oświatowych\n[[t.j. Dz.U. 2026 poz. 650, art. 8 ust. 1]]',
            '[[Obowiązek jednostki samorządu terytorialnego: przeznaczyć na zadania wymagające specjalnej '
            'organizacji nauki środki nie niższe niż kwota z subwencji. Szkoła nie jest adresatem tego przepisu.]]',
            'IPET — wymiar godzin przy każdym zaleceniu; arkusz organizacyjny jako źródło danych dla organu '
            'prowadzącego.'],
           ['RODO\nart. 5 ust. 1 lit. c, art. 9',
            'Minimalizacja danych; dane o zdrowiu jako dane szczególnej kategorii.',
            'Metryczka (bez PESEL w drukach pochodnych); klauzula informacyjna w każdym druku.']],
          widths=[4.3, 6.6, 5.1])

    h2(doc, 'Jak przygotować szkołę do zmian — krok po kroku')
    steps(doc, [
        '**Zrób przegląd wzorów druków używanych w szkole.** Sprawdź w każdym, czy w podstawie prawnej stoi obowiązujący '
        'tekst jednolity, a nie pierwotny publikator. Najczęstszy błąd: „Dz.U. 2017 poz. 1578” zamiast „t.j. Dz.U. 2020 poz. 1309”.',
        '**Rozdziel dwie ścieżki w dokumentacji.** Uczeń z orzeczeniem: WOPF-SP → IPET. Uczeń bez orzeczenia: '
        'rozpoznanie → PWES → ocena efektywności. Jeden kwestionariusz obserwacji obsługuje obie ścieżki.',
        '**Przemapuj narzędzia obserwacji na dziewięć obszarów ICF.** Kwestionariusz KSzOF ma ten układ wbudowany '
        'i przy każdym twierdzeniu podaje kod klasyfikacji.',
        '**Wpisz do kalendarza szkoły wrześniową obserwację.** To ona daje dane, gdy poradnia poprosi o opinię '
        'z terminem dziesięciu dni.',
        '**Ustal, kto pełni funkcję Strażnika Prawa** i w jakim rytmie funkcja się zmienia. Zadanie jest jedno: '
        'przy każdej decyzji pytać, z czego to wynika i gdzie jest zapisane.',
        '**Sprawdź obsadę nauczyciela współorganizującego** dla uczniów z autyzmem, w tym z zespołem Aspergera, '
        'oraz z niepełnosprawnościami sprzężonymi — [[tu zatrudnienie jest obowiązkiem (§ 7 ust. 2). Przy innych '
        'niepełnosprawnościach, niedostosowaniu społecznym i zagrożeniu niedostosowaniem wystąp do organu '
        'prowadzącego o zgodę (§ 7 ust. 3) — bez niej nie ma podstawy do zatrudnienia.]]',
        '**Przygotuj zarządzenie dyrektora** porządkujące obieg dokumentacji ucznia: kto zakłada teczkę, gdzie jest '
        'przechowywana, kto ma do niej dostęp, w jakim terminie wpływają arkusze. Wzór — załącznik D.',
        '**Nie przepisuj dokumentów już sporządzonych.** Program i ocenę aktualizuj przy najbliższej wielospecjalistycznej '
        'ocenie; dotychczasowe zapisy zostają w teczce jako historia wsparcia.',
    ])


# ================================================================ CZĘŚĆ 2

def czesc_2(doc):
    page_break(doc)
    band(doc, 'CZĘŚĆ 2',
         'Dlaczego szkoła musi zmienić dokumentację kształcenia specjalnego — uzasadnienie i zakres zmian',
         '12:20  ·  S2.mp4')

    accent_box(doc, [
        'Ocena funkcjonalna i opinia szkoły o funkcjonowaniu ucznia — § 7 ust. 2–3, 6–7 i § 8 rozporządzenia ME '
        'z 2 marca 2026 r. (Dz.U. 2026 poz. 428): od 1.09.2026 opis funkcjonowania odnosi się do aktywności '
        'i uczestniczenia w rozumieniu ICF, a opinię wydaje się w terminie 10 dni. [[Publikator zweryfikowany; '
        'akt obowiązuje od 14.04.2026 i uchylił rozporządzenie z 7 września 2017 r. (t.j. Dz.U. 2023 poz. 2061).]]',
        'WOPFU jako podstawa programu — § 6 ust. 4 rozporządzenia MEN z 9 sierpnia 2017 r. (t.j. Dz.U. 2020 poz. 1309): '
        'program opracowuje się po dokonaniu oceny, uwzględniając diagnozę i wnioski z oceny oraz zalecenia z orzeczenia.',
        'Ocena efektywności udzielanej pomocy — § 6 rozporządzenia o kształceniu specjalnym oraz przepisy '
        'rozporządzenia o pomocy psychologiczno-pedagogicznej (t.j. Dz.U. 2023 poz. 1798).',
        'Prawa rodziców — § 6 rozporządzenia: udział w spotkaniach zespołu, otrzymanie kopii programu i kopii '
        'wielospecjalistycznej oceny.',
        'Minimalizacja danych — art. 5 ust. 1 lit. c RODO; dane o zdrowiu — art. 9 RODO.',
        'Wydatkowanie środków na kształcenie specjalne — [[art. 8 ust. 1 ustawy z 27 października 2017 r. '
        'o finansowaniu zadań oświatowych (t.j. Dz.U. 2026 poz. 650); obowiązek spoczywa na jednostce samorządu '
        'terytorialnego]].',
        'Dostępność — ustawa z 19 lipca 2019 r. o zapewnianiu dostępności osobom ze szczególnymi potrzebami '
        '(t.j. Dz.U. 2024 poz. 1411).',
        'Nadzór pedagogiczny — [[art. 55 oraz art. 68 ust. 1 ustawy Prawo oświatowe (t.j. Dz.U. 2026 poz. 820) '
        'i rozporządzenie MEN z 25 sierpnia 2017 r. w sprawie nadzoru pedagogicznego (t.j. Dz.U. 2024 poz. 15)]].',
    ], title='PODSTAWA PRAWNA CZĘŚCI')

    h3(doc, 'Transkrypcja narracji')
    narration(doc, [
        'Szanowni Państwo, ten moduł jest inny niż pozostałe. Nie omawiamy w nim druku. Odpowiadamy na pytanie, '
        'które usłyszałam na każdej radzie pedagogicznej, na której byłam.',
        'Pytanie brzmi tak. Prowadzimy dokumentację od lat. Mamy oceny wielospecjalistyczne, mamy programy, mamy teczki. '
        'Kuratorium nigdy nam nic nie zarzuciło. Po co to zmieniać?',
        'To jest dobre pytanie i zasługuje na uczciwą odpowiedź. Zacznę od tego, czego nie powiem. Nie powiem, '
        'że dotychczasowa dokumentacja była zła. Nie powiem, że pracowaliśmy źle. I nie poproszę nikogo, żeby przepisywał '
        'dokumenty, które już powstały.',
        'Powiem coś innego. Od pierwszego września zmienia się nie objętość dokumentacji, lecz jej język i jej funkcja. '
        'A dokument napisany w starym języku przestaje działać w nowym obiegu. Nie dlatego, że jest brzydki. Dlatego, '
        'że nie da się z niego wyjąć informacji, o którą pyta poradnia.',
        'Pokażę to na jednym przykładzie. Weźmy zdanie, które znajdziemy w setkach szkolnych ocen wielospecjalistycznych. '
        'Uczennica ma trudności w koncentracji uwagi wynikające z zaburzeń rozwojowych, wymaga stałej pomocy nauczyciela.',
        'Zapytajmy o to zdanie tak, jak zapyta zespół orzekający. W jakich sytuacjach uczennica traci uwagę, a w jakich '
        'ją utrzymuje? Jak długo pracuje bez podpowiedzi przy zadaniu na kartach, a jak długo przy tablicy? Co dokładnie '
        'znaczy stała pomoc: obecność dorosłego obok, podpowiedź słowna, wspólne wykonanie? Ile razy dziennie? '
        'Co się zmieniło od poprzedniej oceny?',
        'Nasze zdanie nie odpowiada na żadne z tych pytań. Nie dlatego, że nauczycielka nie wie. Ona wie doskonale. '
        'Dlatego, że druk nie miał miejsca, w którym ta wiedza mogłaby zostać zapisana.',
        'A teraz to samo w języku funkcjonalnym. Zofia pracuje samodzielnie przy zadaniu na karcie pracy przez około '
        'osiem minut, jeśli siedzi w pierwszej ławce i zadanie jest podzielone na trzy kroki oznaczone piktogramami. '
        'Przy pracy z tablicą utrzymuje uwagę około dwóch minut i wymaga wtedy podpowiedzi słownej średnio cztery razy '
        'na lekcji. W klasie o podwyższonym hałasie czas ten skraca się do minuty. Od września skróciła się liczba '
        'podpowiedzi z sześciu do czterech.',
        'To jest ta sama uczennica i ta sama wiedza tej samej nauczycielki. Różnica polega wyłącznie na tym, że drugi '
        'zapis odpowiada na cztery pytania: co uczeń robi, w jakich warunkach, przy jakim wsparciu i jak często. '
        'I dopiero z drugiego zapisu da się zbudować cel, dostosowanie i ocenę efektywności.',
        'Przejdźmy teraz do powodów. Jest ich dziesięć i każdy ma swoją podstawę.',
        'Powód pierwszy. Zmienił się adresat naszej dokumentacji. Do tej pory pisaliśmy ją głównie dla siebie i dla '
        'kontroli. Od pierwszego września dwa tysiące dwudziestego szóstego roku poradnia opiera orzeczenie na ocenie '
        'funkcjonalnej, a paragraf siódmy ustęp siódmy rozporządzenia o orzekaniu wymaga, aby opis odnosił się '
        'do aktywności i uczestniczenia w rozumieniu Międzynarodowej Klasyfikacji Funkcjonowania. Jeżeli nasza '
        'dokumentacja mówi innym językiem, poradnia jej po prostu nie użyje. Orzeczenie powstanie bez danych ze szkoły, '
        'czyli bez danych od jedynych dorosłych, którzy widzą ucznia codziennie przez pięć godzin.',
        'Powód drugi. Termin dziesięciu dni. Opinię o funkcjonowaniu ucznia wydaje dyrektor w terminie dziesięciu dni '
        'od dnia otrzymania prośby. W szkole podstawowej ucznia klasy siódmej uczy nawet dwanaścioro nauczycieli. '
        'Zebranie od nich rzetelnych informacji, uzgodnienie ich i napisanie opinii w dziesięć dni jest wykonalne '
        'tylko wtedy, gdy dane już istnieją. Jeżeli obserwację zaczynamy po wpłynięciu prośby, opinia powstanie '
        'z pamięci albo powstanie po terminie. Obie możliwości są złe.',
        'Powód trzeci. Ocena wielospecjalistyczna przestaje być wypracowaniem, a staje się dokumentem scalającym. '
        'Paragraf szósty ustęp czwarty mówi, że program opracowuje się po dokonaniu oceny, uwzględniając diagnozę '
        'i wnioski z tej oceny. To znaczy, że każdy wniosek w ocenie musi mieć źródło, a każde zalecenie w programie '
        'musi mieć wniosek. Jeżeli w ocenie nie widać, skąd wzięło się zdanie, nie da się obronić zalecenia, '
        'które z niego wypływa. To jest najczęstsze pytanie w nadzorze: na jakiej podstawie zespół to stwierdził?',
        'Powód czwarty. Bez celu mierzalnego nie da się przeprowadzić oceny efektywności, a ocena efektywności '
        'jest obowiązkiem. Program, w którym stoi zapis rozwijanie kompetencji społecznych, jest programem, którego '
        'nie da się zewaluować. Po pół roku zespół napisze: cel realizowany częściowo. To zdanie nic nie znaczy '
        'i niczego nie zmienia. Cel z liczbą pozwala napisać: kryterium było trzy z pięciu dni, osiągnięto dwa, '
        'modyfikujemy metodę. I to jest właśnie ocena efektywności w rozumieniu rozporządzenia.',
        'Powód piąty. Zalecenia z orzeczenia muszą mieć wskazany sposób realizacji. Samo przepisanie zalecenia poradni '
        'do programu nie wystarcza. Przy każdym zaleceniu zapisujemy, w jakiej formie je realizujemy, kto je realizuje, '
        'w jakim wymiarze godzin i od kiedy. Brak tej kolumny to najczęstsze uchybienie stwierdzane w kontrolach '
        'dokumentacji kształcenia specjalnego.',
        '[[Powód szósty, o którym mówi się najrzadziej. Pieniądze. Artykuł ósmy ustęp pierwszy ustawy '
        'o finansowaniu zadań oświatowych nakłada obowiązek na jednostkę samorządu terytorialnego: musi ona '
        'przeznaczyć na zadania wymagające stosowania specjalnej organizacji nauki i metod pracy środki w wysokości '
        'nie niższej niż kwota wynikająca z podziału części oświatowej subwencji ogólnej. Powiedzmy wyraźnie: '
        'adresatem tego przepisu jest organ prowadzący, a nie szkoła, i nie wynika z niego żaden bezpośredni '
        'obowiązek dokumentacyjny dla nas. Nasz wniosek jest pośredni, ale praktyczny. Skoro organ prowadzący ma '
        'zaplanować i rozliczyć te środki, potrzebuje wiedzieć, jakie zajęcia, w jakim wymiarze i przez kogo są '
        'realizowane. Tę informację niesie program ucznia i arkusz organizacyjny. Program bez wymiaru godzin, form '
        'i osób realizujących nie daje jej wcale.]]',
        'Powód siódmy. Dostosowania muszą mieć uzasadnienie w opisie bariery. Ustawa o zapewnianiu dostępności '
        'i model biopsychospołeczny mówią to samo: trudność powstaje na styku możliwości ucznia i wymagań otoczenia. '
        'Jeżeli w ocenie nie opisaliśmy, że hałas w sali skraca czas pracy z ośmiu minut do jednej, to dostosowanie '
        'w postaci miejsca w pierwszej ławce i słuchawek wyciszających wisi w próżni. Nikt nie wie, po co je wpisano '
        'i po czym poznamy, że pomogło.',
        'Powód ósmy. Prawa rodziców i głos ucznia. Rodzice mają prawo uczestniczyć w spotkaniach zespołu oraz '
        'otrzymać kopię programu i kopię oceny. To prawo istnieje od lat, a mimo to w wielu teczkach nie ma śladu '
        'jego realizacji. Zapis o przekazaniu kopii, z datą i podpisem, jest elementem dokumentacji, nie uprzejmością. '
        'Do tego dochodzi element nowy, który wprowadzamy ponad wymóg rozporządzenia: sekcja Mój głos, '
        'czyli perspektywa ucznia zapisana jego słowami albo przez wskazanie.',
        'Powód dziewiąty, specyficznie szkolny. Ciągłość między pierwszym a drugim etapem edukacyjnym. W klasie trzeciej '
        'ucznia zna jedna wychowawczyni, która widzi go codziennie. W klasie czwartej zna go dziesięcioro nauczycieli, '
        'z których każdy widzi go dwie godziny w tygodniu. Jeżeli wiedza o uczniu istnieje tylko w głowie wychowawczyni, '
        'we wrześniu klasy czwartej przestaje istnieć. Dlatego wprowadzamy kartę przekazania i wymóg, aby dostosowania '
        'były zapisane przedmiotowo, a nie ogólnie.',
        'Powód dziesiąty. Egzamin ósmoklasisty. Dostosowanie warunków i form egzaminu przysługuje na podstawie '
        'orzeczenia, opinii poradni albo pozytywnej opinii rady pedagogicznej, a rada opiera się na dokumentacji. '
        'Uczeń, którego trudności nie zostały opisane w dokumentacji w listopadzie, nie otrzyma dostosowania w maju. '
        'Tego się nie da naprawić po terminie.',
        'To było dziesięć powodów. Powiem teraz jedno zdanie, które chciałabym, żeby Państwo zapamiętali.',
        'Nie zmieniamy dokumentacji dlatego, że ktoś nam kazał. Zmieniamy ją dlatego, że stara dokumentacja przestała '
        'odpowiadać na pytania, które nam teraz zadają: poradnia w terminie dziesięciu dni, rodzic na zespole, '
        'organ prowadzący przy rozliczeniu godzin i nauczyciel klasy czwartej we wrześniu.',
        'Przejdźmy do części praktycznej. Co dokładnie zmieniamy?',
        'Zmiana pierwsza dotyczy narzędzia obserwacji. Zastępujemy opisowe karty obserwacji kwestionariuszem '
        'z kodami klasyfikacji i skalą. W szkole podstawowej używamy Kwestionariusza Szkolnej Oceny Funkcjonalnej, '
        'w wersjach dla klas pierwszej do trzeciej, czwartej do szóstej i siódmej do ósmej.',
        'Zmiana druga dotyczy oceny wielospecjalistycznej. Przestaje być wypracowaniem pisanym od nowa, a staje się '
        'kartą scalającą z wpisanymi źródłami. Każda sekcja oceny wskazuje druk, z którego pochodzi.',
        'Zmiana trzecia dotyczy programu. Zalecenia z orzeczenia i z oceny wpisujemy w dwóch kolumnach: treść zalecenia '
        'i sposób realizacji, z formą, osobą, wymiarem i datą rozpoczęcia.',
        'Zmiana czwarta dotyczy celów. Każdy cel otrzymuje kryterium liczbowe i datę pomiaru, a wskaźnik z celu '
        'staje się gotowym narzędziem ewaluacji.',
        'Zmiana piąta dotyczy dostosowań. Zapisujemy je przedmiotowo, a nie ogólnie, i przy każdym wskazujemy barierę '
        'opisaną w ocenie, z której wynika.',
        'Zmiana szósta dotyczy ucznia bez orzeczenia. Wprowadzamy plan wsparcia edukacyjnego, żeby pomoc '
        'psychologiczno-pedagogiczna miała swój dokument, tak jak kształcenie specjalne ma program.',
        'Zmiana siódma dotyczy danych. Usuwamy z druków pochodnych numer PESEL, adres zamieszkania, dane o miejscu '
        'pracy rodziców i kopie zaświadczeń lekarskich, które nie są potrzebne do organizacji kształcenia. '
        'Druki dziedziczą dane z metryczki, zamiast je powielać.',
        'Zmiana ósma dotyczy śladu współpracy. Rejestr kontaktów z rodzicami, potwierdzenie przekazania kopii oceny '
        'i programu oraz zapis udziału rodzica w zespole stają się stałym elementem teczki.',
        'A teraz równie ważne: czego nie zmieniamy.',
        'Nie przepisujemy dokumentów już sporządzonych. Nie unieważniamy dotychczasowych ocen i programów. '
        'Nie tworzymy dokumentów, których nie wymaga przepis ani nasza procedura. I nie zwiększamy liczby druków. '
        'Wprowadzamy druki po to, żeby te same dane wpisywać raz, a nie w pięciu miejscach.',
        'Program i ocenę aktualizujemy przy najbliższej wielospecjalistycznej ocenie, czyli we wrześniu, '
        'a dotychczasowe zapisy zostają w teczce jako historia wsparcia. Historia wsparcia to nie balast. '
        'To dowód, że pracowaliśmy, i punkt odniesienia dla oceny efektywności.',
        'Podsumowując moduł drugi. Zmieniamy język, nie objętość. Zmieniamy strukturę, nie liczbę druków. '
        'I robimy to po to, żeby dziesięć dni na opinię było terminem realnym, a nie źródłem stresu. '
        'W module trzecim zobaczymy cały obieg dokumentów z lotu ptaka.',
    ])

    h2(doc, 'Dziesięć powodów zmiany — tabela argumentacyjna dla dyrektora i rady pedagogicznej')
    table(doc,
          ['#', 'Powód', 'Podstawa', 'Co się dzieje, jeśli nie zmienimy'],
          [['1', 'Zmienił się adresat: poradnia opiera orzeczenie na ocenie funkcjonalnej w języku ICF.',
            '§ 7 ust. 7 rozp. Dz.U. 2026 poz. 428\n[[wchodzi 1.09.2026]]',
            'Dane ze szkoły nie zostaną wykorzystane; orzeczenie powstaje bez perspektywy nauczycieli.'],
           ['2', 'Opinia o funkcjonowaniu ucznia w terminie 10 dni od wpływu prośby do dyrektora.',
            '§ 7 ust. 2–3 rozp. Dz.U. 2026 poz. 428\n[[zweryfikowane]]',
            'Opinia pisana z pamięci albo po terminie; uchybienie formalne i niska jakość danych.'],
           ['3', 'WOPFU jest podstawą programu — każdy wniosek musi mieć źródło.',
            '§ 6 ust. 4 rozp. t.j. Dz.U. 2020 poz. 1309',
            'Zalecenia bez uzasadnienia; pytanie nadzoru „na jakiej podstawie zespół to stwierdził” zostaje bez odpowiedzi.'],
           ['4', 'Ocena efektywności wymaga celu z kryterium.',
            '§ 6 rozp. t.j. Dz.U. 2020 poz. 1309; rozp. o pomocy pp t.j. Dz.U. 2023 poz. 1798',
            'Ewaluacja sprowadza się do formuły „cel realizowany częściowo”, która niczego nie zmienia.'],
           ['5', 'Przy każdym zaleceniu wymagany sposób realizacji: forma, osoba, wymiar, data.',
            '§ 6 ust. 1 i 4 rozp. t.j. Dz.U. 2020 poz. 1309',
            'Najczęstsze uchybienie w kontroli dokumentacji kształcenia specjalnego.'],
           ['6', '[[Dokumentacja ucznia jest źródłem danych, na których organ prowadzący planuje i rozlicza środki '
            'na specjalną organizację nauki (obowiązek JST, nie szkoły).]]',
            '[[art. 8 ust. 1 ustawy o finansowaniu zadań oświatowych, t.j. Dz.U. 2026 poz. 650]]',
            'Organ prowadzący nie ma na czym oprzeć planowania i rozliczenia; ryzyko przy kontroli.'],
           ['7', 'Dostosowanie musi wynikać z opisanej bariery środowiskowej.',
            'ustawa o dostępności (t.j. Dz.U. 2024 poz. 1411); model ICF (WHO 2001)',
            'Dostosowania nieuzasadnione i nieweryfikowalne; nie wiadomo, po czym poznać, że pomogły.'],
           ['8', 'Prawa rodziców: udział w zespole, kopia oceny i programu; głos ucznia.',
            '§ 6 rozp. t.j. Dz.U. 2020 poz. 1309',
            'Brak śladu realizacji prawa; ryzyko skargi i sporu o przebieg współpracy.'],
           ['9', 'Ciągłość między I a II etapem edukacyjnym (klasa III → klasa IV).',
            'decyzja rady pedagogicznej wpisana do procedury',
            'Wiedza o uczniu ginie we wrześniu klasy IV; dostosowania trzeba budować od zera.'],
           ['10', 'Dostosowanie warunków egzaminu ósmoklasisty wymaga podstawy w dokumentacji.',
            'komunikat dyrektora CKE na dany rok ⚑',
            'Uczeń nie otrzyma dostosowania na egzaminie; po terminie nie da się tego naprawić.']],
          widths=[0.9, 5.4, 4.3, 5.4], size=8)

    h2(doc, 'Co dokładnie zmienić — mapa „było → ma być”')
    table(doc,
          ['Element dokumentacji', 'Było (praktyka dotychczasowa)', 'Ma być od 1.09.2026', 'Dlaczego'],
          [['Podstawa prawna w stopce druku', 'Publikator pierwotny (Dz.U. 2017 poz. 1578).',
            'Obowiązujący tekst jednolity (t.j. Dz.U. 2020 poz. 1309) i data sprawdzenia w ISAP.',
            'Druk z nieaktualnym publikatorem jest drukiem wadliwym.'],
           ['Narzędzie obserwacji', 'Opisowa karta obserwacji wychowawcy, bez skali.',
            'KSzOF: 52 twierdzenia, 9 obszarów ICF, skala 1–5, kody d110–d920, normy stenowe.',
            'Wynik porównywalny w czasie i zrozumiały dla poradni.'],
           ['Kto wypełnia arkusz', 'Wyłącznie wychowawca.',
            'Niezależnie wychowawca i rodzic; w II etapie dodatkowo nauczyciel przedmiotu i specjalista.',
            'Rozbieżność ocen jest informacją o środowisku, nie błędem.'],
           ['Termin obserwacji', 'Kwiecień, przed zespołem.',
            'Wrzesień (pomiar bazowy) i maj (pomiar kontrolny).',
            'Dane muszą istnieć, zanim wpłynie prośba z poradni z terminem 10 dni.'],
           ['WOPF', 'Wypracowanie opisowe pisane od nowa przy każdej ocenie.',
            'Karta scalająca z jawnymi źródłami: sekcja V z KSzOF, VI z ABC/FBA, VII z ToM, VIII z karty mowy, IX z profilu sensorycznego.',
            '§ 6 ust. 4 — program opracowuje się po ocenie i na jej podstawie.'],
           ['Opis trudności', 'Etykieta i rozpoznanie („zaburzenia koncentracji”).',
            'Zachowanie + warunki + wsparcie + częstotliwość.',
            'Język funkcjonalny wymagany przez § 7 ust. 7 rozp. o orzekaniu.'],
           ['Bariery i ułatwienia', 'Sekcja pomijana albo wypełniana ogólnikami.',
            'Konkretne czynniki środowiskowe z lekcji, przerwy, świetlicy i drogi do szkoły.',
            'Dostosowanie bez opisanej bariery jest nieuzasadnione.'],
           ['Zalecenia w IPET', 'Przepisane z orzeczenia, jedna kolumna.',
            'Dwie kolumny: treść zalecenia oraz sposób realizacji (forma, kto, wymiar, od kiedy).',
            'Wymóg rozporządzenia i dowód wydatkowania środków.'],
           ['Cele', '„Rozwijanie…”, „doskonalenie…”, „usprawnianie…”.',
            'Cel z zachowaniem, sytuacją, liczbą prób, poziomem wsparcia, datą i sposobem pomiaru.',
            'Bez kryterium nie da się dokonać wymaganej oceny efektywności.'],
           ['Dostosowania', 'Lista ogólna dla całego programu.',
            'Zapis przedmiotowy: co zmieniamy na języku polskim, co na matematyce, co na wychowaniu fizycznym.',
            'W II etapie ucznia uczy kilkunastu nauczycieli — ogólna lista nie działa.'],
           ['Uczeń bez orzeczenia', 'Notatka w dzienniku, brak dokumentu wiodącego.',
            'Plan wsparcia edukacyjnego ucznia (PWES) + karta oceny efektywności.',
            '[[Pomoc pp podlega ocenie efektywności. Sam PWES to narzędzie wewnętrzne szkoły — rozporządzenie '
            'nie przewiduje takiego druku.]]'],
           ['Ocena zachowania', 'Bez odniesienia do orzeczenia lub opinii.',
            'Zapis o uwzględnieniu wpływu zaburzeń, z odesłaniem do dokumentu.',
            '[[Rozp. o ocenianiu — t.j. Dz.U. 2023 poz. 2572, z późn. zm.]]'],
           ['Dane osobowe w drukach', 'PESEL, adres, miejsce pracy rodziców powielane w każdym druku.',
            'Dane wyłącznie w metryczce i w dokumentacji przebiegu nauczania; druki dziedziczą.',
            'art. 5 ust. 1 lit. c RODO — minimalizacja danych.'],
           ['Współpraca z rodzicami', 'Ustalenia ustne.',
            'Rejestr kontaktów + potwierdzenie przekazania kopii oceny i programu, z datą i podpisem.',
            '§ 6 rozporządzenia — prawa rodziców muszą mieć ślad.'],
           ['Głos ucznia', 'Nieobecny.',
            'Sekcja „Mój głos” — słowami ucznia albo przez wskazanie, piktogram, AAC.',
            'Element ponad wymóg rozporządzenia; podnosi trafność celów.'],
           ['Przejście III → IV', 'Rozmowa wychowawców na korytarzu.',
            'Karta przekazania informacji o uczniu, przekazywana w czerwcu.',
            'Zapewnia ciągłość wsparcia na starcie II etapu.'],
           ['Egzamin ósmoklasisty', 'Decyzja podejmowana wiosną, „bo uczeń ma opinię”.',
            'Karta dostosowań egzaminacyjnych zakładana w listopadzie, ze wskazaniem podstawy.',
            'Dostosowanie musi mieć podstawę w dokumentacji przed terminem z komunikatu CKE ⚑.']],
          widths=[3.3, 4.3, 4.6, 3.8], size=7.5)

    h2(doc, 'Czego NIE zmieniamy — cztery zdania, które chronią zespół przed przepisywaniem dokumentacji')
    accent_box(doc, [
        'Nie przepisujemy dokumentów już sporządzonych. Ocena i program obowiązują do najbliższej '
        'wielospecjalistycznej oceny i dopiero wtedy powstają w nowym układzie.',
        'Nie unieważniamy dotychczasowych zapisów. Zostają w teczce jako historia wsparcia i są punktem odniesienia '
        'dla oceny efektywności — bez nich nie da się wykazać postępu.',
        'Nie tworzymy dokumentów, których nie wymaga ani przepis, ani nasza procedura. Każdy druk w załączniku C '
        'ma wskazane, z czego wynika i kto go wypełnia.',
        'Nie zwiększamy liczby druków. Zwiększamy liczbę druków, które dziedziczą dane. Ta sama informacja '
        'wpisywana jest raz i wędruje dalej.',
    ], title='ZASADA CIĄGŁOŚCI')

    h2(doc, 'Ryzyka zaniechania — co realnie grozi szkole, która niczego nie zmieni')
    numbered(doc, [
        '**Uchybienie w nadzorze pedagogicznym.** Kontrola dokumentacji kształcenia specjalnego sprawdza trzy rzeczy: '
        'czy ocena poprzedza program, czy przy zaleceniach wskazano sposób realizacji i czy dokonano oceny efektywności. '
        'Dokumentacja opisowa bez źródeł nie przechodzi żadnego z tych testów.',
        '**Niedotrzymany termin ustawowy.** Dziesięć dni na opinię o funkcjonowaniu ucznia liczy się od wpływu prośby '
        'do dyrektora, a nie od momentu, w którym zespół znajdzie czas.',
        '**Spór z rodzicem bez dowodów.** Brak rejestru kontaktów i potwierdzeń przekazania kopii oznacza, '
        'że w sporze szkoła nie ma czym wykazać, że informowała, konsultowała i ustalała.',
        '**Ryzyko przy rozliczeniu środków.** Bez wymiaru godzin, form i osób realizujących w programie nie da się '
        'wykazać, że środki naliczone na kształcenie specjalne posłużyły organizacji tego kształcenia.',
        '**Naruszenie zasady minimalizacji danych.** Powielanie numeru PESEL, adresu i danych medycznych w kolejnych '
        'drukach to przetwarzanie danych ponad cel — przy danych o zdrowiu mówimy o danych szczególnej kategorii.',
        '**Utrata dostosowań na egzaminie ósmoklasisty.** Uczeń, którego trudności nie zostały opisane w dokumentacji '
        'przed terminem wskazanym w komunikacie CKE, nie otrzyma dostosowania warunków egzaminu. ⚑',
        '**Koszt ukryty: podwójna praca.** Najdroższa jest dokumentacja, w której te same dane wpisuje się '
        'po raz czwarty, bo żaden druk nie dziedziczy poprzedniego. Zmiana, o której mówimy, ten koszt zdejmuje.',
    ])


# ================================================================ CZĘŚĆ 3

def czesc_3(doc):
    page_break(doc)
    band(doc, 'CZĘŚĆ 3', 'Obieg dokumentów w szkole — jak jeden wynika z drugiego', '6:10  ·  S3.mp4')

    accent_box(doc, [
        'Kształcenie specjalne — § 6 rozporządzenia MEN z 9 sierpnia 2017 r. (t.j. Dz.U. 2020 poz. 1309): '
        'WOPFU poprzedza IPET i jest jego podstawą; ocena co najmniej dwa razy w roku szkolnym.',
        'Pomoc psychologiczno-pedagogiczna — rozpoznawanie potrzeb przez nauczycieli, formy pomocy w szkole '
        'i ocena efektywności (t.j. Dz.U. 2023 poz. 1798).',
        'Opinia o funkcjonowaniu ucznia dla poradni — § 7 ust. 2–3 rozporządzenia o orzekaniu (Dz.U. 2026 poz. 428).',
        'Dokumentacja badań i czynności uzupełniających — rozporządzenie o dokumentacji przebiegu nauczania '
        '(t.j. Dz.U. 2024 poz. 50).',
    ], title='PODSTAWA PRAWNA CZĘŚCI')

    h3(doc, 'Transkrypcja narracji')
    narration(doc, [
        'Szanowni Państwo, w trzecim module przyjrzymy się całej dokumentacji z lotu ptaka.',
        'Po tym module będą Państwo wiedzieć trzy rzeczy. Jakie dokumenty tworzymy w ciągu roku szkolnego. '
        'W jakiej kolejności one powstają. I dlaczego żaden z nich nie powstaje osobno.',
        'Zacznijmy od najważniejszego zdania tego szkolenia. Dokumentacja ucznia to jeden obieg, w którym każdy dokument '
        'bierze dane z poprzedniego. Jeśli którykolwiek etap pominiemy, następny trzeba wypełniać z pamięci. '
        'A dokumentacja wypełniana z pamięci nie służy ani uczniowi, ani nam.',
        'W szkole obieg ma siedem przystanków i jedno rozgałęzienie. Omówię je po kolei.',
        'Przystanek pierwszy to metryczka ucznia i teczka.',
        'Metryczka gromadzi dane formalne, kontakty do rodziców i informacje o zdrowiu istotne dla funkcjonowania '
        'w szkole. Jej najważniejsza funkcja jest jednak inna. Metryczka odpowiada na pytanie, czy uczeń posiada '
        'orzeczenie, opinię lub inną formę wsparcia, i od kiedy. Data wpływu orzeczenia do szkoły uruchamia '
        'trzydziestodniowy termin na opracowanie programu. Metryczka przekazuje dalej także sygnały zdrowotne, '
        'na przykład informację o nadwrażliwości sensorycznej albo o chorobie przewlekłej. Do tych sygnałów wrócimy '
        'przy decyzji o obserwacji pogłębionej.',
        'Przystanek drugi to Kwestionariusz Szkolnej Oceny Funkcjonalnej, w skrócie KSzOF.',
        'Kwestionariusz jest narzędziem przesiewowym. Obejmuje wszystkich uczniów w oddziale, nie tylko tych, '
        'którzy budzą nasz niepokój. Wypełniają go niezależnie wychowawca i rodzic, a w drugim etapie edukacyjnym '
        'także nauczyciel przedmiotu i specjalista — każdy na podstawie kilku tygodni obserwacji. Wynikiem jest profil '
        'ucznia w dziewięciu obszarach funkcjonowania, wynik ogólny w stenach oraz lista zachowań, które zespół omawia '
        'w pierwszej kolejności. Kwestionariusz odpowiada na pytanie: gdzie? Gdzie uczeń radzi sobie dobrze, '
        'a gdzie potrzebuje naszej pomocy.',
        'Przystanek trzeci to obserwacja pogłębiona.',
        'Uruchamiamy ją wyłącznie wtedy, gdy wynik kwestionariusza wskaże konkretny kierunek. Mamy do dyspozycji '
        'cztery narzędzia. Analizę zachowania w modelu ABC wraz z arkuszem analizy funkcjonalnej, profil sensoryczny, '
        'kartę obserwacji rozwoju mowy i komunikacji oraz obserwację poznania społecznego i teorii umysłu. '
        'Obserwacja pogłębiona odpowiada na pytanie: dlaczego? Dlaczego uczeń zachowuje się w określony sposób '
        'i co w środowisku szkolnym mu pomaga albo przeszkadza.',
        'Przystanek czwarty to wielospecjalistyczna ocena poziomu funkcjonowania, w naszym druku oznaczona '
        'jako WOPF dla szkoły podstawowej.',
        'Ocena scala wszystko, co zebraliśmy wcześniej. Profil z kwestionariusza, wnioski z obserwacji pogłębionej, '
        'treść orzeczenia lub opinii poradni, informacje od rodziców, głos ucznia oraz efekty dotychczasowego wsparcia. '
        'Ocenę piszemy językiem funkcjonalnym: co uczeń robi, w jakich warunkach i przy jakim wsparciu. '
        'Opisujemy mocne strony, trudności oraz bariery i ułatwienia w środowisku.',
        'I tutaj jest rozgałęzienie, o którym mówiłam w module pierwszym. Sekcja pierwsza a druku oceny każe wybrać '
        'ścieżkę. Uczeń z orzeczeniem o potrzebie kształcenia specjalnego idzie do indywidualnego programu '
        'edukacyjno-terapeutycznego. Uczeń bez orzeczenia — do planu wsparcia edukacyjnego w ramach pomocy '
        'psychologiczno-pedagogicznej. Jeden druk, dwie ścieżki.',
        'Przystanek piąty, ścieżka pierwsza, to indywidualny program edukacyjno-terapeutyczny.',
        'Program wynika z oceny zdanie po zdaniu. Jeśli czegoś nie ma w ocenie, nie może pojawić się w programie. '
        'Program zawiera zalecenia z orzeczenia wraz ze sposobem realizacji, cele w postaci mierzalnej, dostosowania '
        'zapisane przedmiotowo, zintegrowane działania nauczycieli i specjalistów, formy i wymiar wsparcia '
        'oraz zakres współpracy z rodzicami.',
        'Przystanek piąty, ścieżka druga, to plan wsparcia edukacyjnego ucznia. Powstaje dla ucznia z opinią poradni '
        'albo rozpoznanego przez nauczycieli. Ma tę samą logikę co program: potrzeba, cel z kryterium, forma pomocy, '
        'osoba, wymiar i termin oceny efektywności. Różni się podstawą prawną i tym, że nie wymaga orzeczenia.',
        'Przystanek szósty to ewaluacja.',
        'Ewaluacja nie wymaga tworzenia nowych narzędzi. Wskaźnik został zapisany wcześniej, w celu. Zespół odnotowuje '
        'wartość osiągniętą i podejmuje jedną z czterech decyzji. Zamyka cel, kontynuuje pracę, modyfikuje program '
        'albo spotyka się z rodzicami i rozważa wystąpienie do poradni. Decyzja wraca do przystanku czwartego '
        'i obieg rozpoczyna się od nowa.',
        'Przystanek siódmy jest szkolny i sezonowy. To dokumenty przejścia i egzaminu. W czerwcu klasy trzeciej — '
        'karta przekazania informacji o uczniu do drugiego etapu edukacyjnego. W listopadzie klasy ósmej — karta '
        'dostosowań warunków egzaminu ósmoklasisty, przygotowana na podstawie dokumentacji, którą już mamy.',
        'Obok obiegu stoi jeszcze jeden dokument. Opinia o funkcjonowaniu ucznia dla poradni '
        'psychologiczno-pedagogicznej.',
        'Wydajemy ją na prośbę przewodniczącego zespołu orzekającego, w terminie dziesięciu dni od dnia, w którym '
        'dyrektor otrzymał prośbę. Nie piszemy jej od zera. Składamy ją z tego, co już mamy. Mocne strony bierzemy '
        'z kwestionariusza. Opis trudności z obserwacji pogłębionej. Efekty wsparcia z karty ewaluacji. '
        'Przebieg współpracy z rodzicami z rejestru w metryczce. Dzięki temu dziesięć dni to termin realny, '
        'a nie źródło stresu.',
        'Podsumowując. Siedem przystanków: metryczka, kwestionariusz, obserwacja pogłębiona, ocena, program albo plan '
        'wsparcia, ewaluacja, dokumenty przejścia. Każdy z nich omówimy w kolejnych modułach, zawsze w tym samym '
        'porządku. Najpierw po co dokument powstaje. Potem jak go stworzyć. Na końcu jak go wypełnić, krok po kroku, '
        'na Państwa drukach. Zapraszam do modułu czwartego.',
    ])

    h2(doc, 'Obieg dokumentów — jeden druk zasila następny')
    table(doc,
          ['Przystanek', 'Dokument', 'Pytanie, na które odpowiada', 'Zasila'],
          [['1', 'Metryczka i teczka ucznia', 'Kogo wezwać? Co podać? Od kiedy liczyć termin?',
            'Wszystkie kolejne druki (dziedziczenie danych).'],
           ['2', 'KSzOF — kwestionariusz obserwacji', 'GDZIE uczeń radzi sobie dobrze, a gdzie potrzebuje wsparcia?',
            'Sekcja V oceny; decyzja o obserwacji pogłębionej.'],
           ['3', 'Obserwacja pogłębiona (ABC/FBA, sensoryka, mowa, ToM)', 'DLACZEGO zachowanie się powtarza?',
            'Sekcje VI–IX oceny; opis barier.'],
           ['4', 'WOPF-SP — ocena scalająca', 'Jaki jest całościowy obraz funkcjonowania?',
            'IPET (z orzeczeniem) albo PWES (bez orzeczenia).'],
           ['5a', 'IPET — program edukacyjno-terapeutyczny', 'Co robimy, kto, jak często i do kiedy?',
            'Karty celów SMART; arkusz organizacyjny.'],
           ['5b', 'PWES — plan wsparcia edukacyjnego', 'Jak organizujemy pomoc pp bez orzeczenia?',
            'Karta oceny efektywności pomocy.'],
           ['6', 'Karta ewaluacji celu', 'Czy zadziałało i co dalej?',
            'Kolejna ocena wielospecjalistyczna.'],
           ['7', 'Karta przekazania III→IV; karta dostosowań egzaminacyjnych', 'Jak zachować ciągłość i prawa ucznia?',
            'Dokumentacja II etapu; procedura egzaminacyjna.'],
           ['obok', 'Opinia o funkcjonowaniu ucznia dla poradni', 'Jak uczeń funkcjonuje w szkole — dla zespołu orzekającego?',
            'Składana z druków 1–6, termin 10 dni.']],
          widths=[1.5, 4.4, 5.2, 5.0], size=8)

    h2(doc, 'Jak ustawić obieg dokumentów w szkole')
    steps(doc, [
        '**Załóż teczkę ucznia** i umieść w niej metryczkę jako pierwszy dokument. To ona odpowiada na pytania: '
        'kogo wezwać, co podać, od kiedy liczyć termin.',
        '**We wrześniu wypełnij KSzOF dla wszystkich uczniów w oddziale.** Wynik wskazuje obszary, a nie diagnozę.',
        '**Uruchom obserwację pogłębioną tylko wtedy, gdy zadziała reguła przekierowania.** Wybór narzędzia '
        'zapisz w karcie decyzyjnej — także wtedy, gdy decyzja brzmi „nie uruchamiamy”.',
        '**Zbierz dane w ocenie wielospecjalistycznej.** Każdy blok oceny ma mieć źródło w druku, który już powstał; '
        'w sekcji I a wybierz ścieżkę: IPET albo PWES.',
        '**Opracuj IPET wprost z oceny:** zalecenia z orzeczenia i z oceny, przy każdym zapis realizacji — forma, kto, '
        'wymiar, od kiedy.',
        '**Dla ucznia bez orzeczenia sporządź PWES** i wyznacz termin oceny efektywności.',
        '**Zaplanuj ewaluację w kalendarzu roku** i po każdym pomiarze podejmij jedną z czterech decyzji: '
        'zamykamy cel, kontynuujemy, modyfikujemy metodę albo występujemy do poradni.',
        '**W czerwcu klasy III wypełnij kartę przekazania,** a w listopadzie klasy VIII — kartę dostosowań '
        'egzaminacyjnych. Oba druki składasz z danych, które już masz.',
    ])


# ================================================================ CZĘŚĆ 4

def czesc_4(doc):
    page_break(doc)
    band(doc, 'CZĘŚĆ 4', 'Metryczka i teczka ucznia — pierwszy dokument września', '6:30  ·  S4.mp4')

    accent_box(doc, [
        'Dokumentacja przebiegu nauczania — rozporządzenie MEN z 25 sierpnia 2017 r. (t.j. Dz.U. 2024 poz. 50): '
        'księga uczniów, dzienniki, arkusze ocen, dokumentacja badań i czynności uzupełniających.',
        'Metryczka jest narzędziem wewnętrznym szkoły — wprowadza ją zarządzenie dyrektora; porządkuje dane '
        'gromadzone z innych tytułów i nie zastępuje dokumentacji przebiegu nauczania.',
        'RODO — art. 5 ust. 1 lit. c (minimalizacja danych) i art. 9 (dane o zdrowiu jako dane szczególnej kategorii).',
        'Termin IPET — § 6 ust. 1 rozporządzenia o kształceniu specjalnym: 30 dni od dnia złożenia w szkole '
        'orzeczenia; datę wpływu odnotowuje się w metryczce.',
    ], title='PODSTAWA PRAWNA CZĘŚCI')

    h3(doc, 'Transkrypcja narracji')
    narration(doc, [
        'Szanowni Państwo, moduł czwarty poświęcamy metryczce ucznia. To pierwszy dokument, który wypełniamy we wrześniu.',
        'Po tym module będą Państwo potrafili uzasadnić, po co prowadzimy metryczkę, wypełnić ją bez zbierania danych '
        'nadmiarowych oraz odczytać z niej w kilkanaście sekund trzy najważniejsze informacje.',
        'Metryczka ucznia to karta danych, która gromadzi w jednym miejscu to, co i tak musimy posiadać. Dane '
        'identyfikacyjne w zakresie minimalnym, kontakty do rodziców, informacje o zdrowiu istotne dla funkcjonowania '
        'w szkole, podstawę objęcia wsparciem oraz spis dokumentacji ucznia. Jest naszym narzędziem wewnętrznym, '
        'wprowadzonym zarządzeniem dyrektora. Powstała po to, żeby oszczędzać czas i chronić ucznia.',
        'Zasadność metryczki przedstawiamy w pięciu punktach.',
        'Po pierwsze, jedno miejsce zamiast siedmiu. Nauczyciel na zastępstwie nie szuka po segregatorach numeru telefonu '
        'do rodzica ani informacji o alergii. Otwiera jedną kartę.',
        'Po drugie, bezpieczeństwo ucznia tu i teraz. Kontakt w nagłych wypadkach, choroby przewlekłe, leki, dieta '
        'i procedura postępowania. To informacje, które muszą być dostępne natychmiast. Metryczka jest dokumentem '
        'operacyjnym, a nie archiwalnym.',
        'Po trzecie, punkt wyjścia dla wsparcia. Sekcja o podstawie objęcia wsparciem pokazuje, czy uczeń ma orzeczenie, '
        'opinię, czy jest objęty pomocą na podstawie rozpoznania nauczycieli, i od kiedy. Data wpływu orzeczenia '
        'do szkoły uruchamia trzydziestodniowy termin na program. To jedyne miejsce, w którym ta data jest zapisana.',
        'Po czwarte, udokumentowana współpraca z rodzicami. Rejestr kontaktów i ustaleń potwierdza, że szkoła '
        'informowała, konsultowała i ustalała. Bez tego rejestru w sporze mamy tylko swoje wspomnienia.',
        'Po piąte, zgodność z przepisami o ochronie danych. Klauzula informacyjna jest częścią dokumentu i jest '
        'podpisana przez rodzica, a zakres danych ograniczamy do niezbędnego.',
        'Przejdźmy teraz do wypełniania. Posłużymy się przykładem uczennicy z dokumentacji wzorcowej: '
        'Zofia Lewandowska, klasa trzecia A.',
        'Sekcja pierwsza to dane ucznia w zakresie minimalnym. Imię i nazwisko, klasa lub oddział, etap edukacyjny, '
        'rok szkolny, wychowawca. Numeru PESEL w metryczce nie wpisujemy. Jest w księdze uczniów i w dokumentacji '
        'przebiegu nauczania, i nie powielamy go w kolejnym dokumencie, zgodnie z zasadą minimalizacji danych. '
        'Nie wpisujemy też adresu zamieszkania, miejsca urodzenia ani danych o miejscu pracy rodziców. '
        'To najczęstszy nadmiar w szkolnych drukach.',
        'Sekcja druga to podstawa objęcia kształceniem specjalnym albo pomocą psychologiczno-pedagogiczną. '
        'Przy każdej formie wsparcia wpisujemy numer i datę dokumentu, a nie tylko zaznaczenie. W naszym przykładzie: '
        'orzeczenie numer PPP kropka cztery dwa dwa trzy kropka osiemnaście kropka dwa tysiące dwadzieścia sześć, '
        'z dnia dwunastego czerwca dwa tysiące dwudziestego szóstego roku, wydane przez poradnię w Koszalinie '
        'ze względu na niepełnosprawność sprzężoną. Obok — i to jest pole, o którym najczęściej się zapomina — '
        'data wpływu orzeczenia do szkoły. To od niej liczymy trzydzieści dni.',
        'Sekcja trzecia to wybór ścieżki wsparcia. Jedno pole, dwie możliwości: kształcenie specjalne prowadzące '
        'do programu albo pomoc psychologiczno-pedagogiczna prowadząca do planu wsparcia. Zaznaczenie tego pola '
        'we wrześniu oszczędza zespołowi godziny dyskusji w styczniu.',
        'Sekcja czwarta to kontakty. Rodzice, preferowana forma kontaktu, kolejność powiadamiania w nagłych wypadkach. '
        'W szkole dochodzi jeszcze jedno pole, którego nie ma w przedszkolu: sposób powrotu ucznia do domu '
        'i ewentualne upoważnienia, jeśli uczeń nie wraca samodzielnie.',
        'Sekcja piąta dotyczy zdrowia i funkcjonowania w szkole. Tutaj warto zatrzymać się dłużej. Wpisujemy choroby '
        'przewlekłe i związane z nimi ostrzeżenia, na przykład astmę, padaczkę albo cukrzycę, oraz przyjmowane leki, '
        'bo to wpływa na codzienne funkcjonowanie ucznia. W polu zaleceń dotyczących postępowania zapisujemy konkretnie, '
        'kto podaje lek, na jakiej podstawie, gdzie lek jest przechowywany i kogo powiadamiamy w jakiej kolejności. '
        'Zaznaczamy również sygnały, które uruchomią obserwację pogłębioną — na przykład nadwrażliwość sensoryczną. '
        'Ten sygnał wróci do nas w module szóstym.',
        'Nie kopiujemy do metryczki dokumentacji medycznej. Wpisujemy wyłącznie tę informację, która jest niezbędna '
        'do organizacji kształcenia i zapewnienia bezpieczeństwa. Zaświadczenie lekarskie zostaje w teczce, '
        'a nie w treści druku.',
        'Sekcja szósta to źródła informacji o uczniu, czyli audyt. Wypisujemy w niej, jakie dokumenty o uczniu '
        'już posiadamy: orzeczenie, opinię, wcześniejsze oceny, informację z przedszkola, karty specjalistów. '
        'Ten spis pokazuje, czego nie musimy zbierać po raz drugi.',
        'Sekcja siódma to rejestr kontaktów z rodzicami. Data, forma kontaktu, temat, ustalenia i podpis. '
        'Notujemy rozmowę, telefon, spotkanie zespołu i przekazanie kopii dokumentu. Ten rejestr zasili później '
        'opinię dla poradni i będzie dowodem w każdej sytuacji spornej.',
        'Sekcja ósma to wykaz dokumentacji ucznia, czyli spis zawartości teczki, z datami. Dzięki niemu w każdej chwili '
        'wiemy, co w teczce jest i czego brakuje.',
        'Trzy dobre praktyki na koniec.',
        'Zbieramy tylko te dane, które są potrzebne do realizacji zadań szkoły. Aktualizujemy metryczkę przy każdej '
        'zmianie zgłoszonej przez rodzica, z datą. I przechowujemy ją w miejscu wskazanym zarządzeniem dyrektora, '
        'bo zawiera dane o zdrowiu, a więc dane szczególnej kategorii.',
        'Podsumowując. Metryczka jest gotowa wtedy, gdy w kilkanaście sekund odczytamy z niej odpowiedzi na trzy pytania. '
        'Kogo wezwać? Co podać? Od kiedy liczyć termin? W module piątym przechodzimy do kwestionariusza, '
        'czyli do obserwacji ucznia.',
    ])

    h2(doc, 'Czego pilnujemy w metryczce ucznia')
    table(doc,
          ['Element', 'Co wpisujemy', 'Skąd bierzemy dane', 'Czego NIE wpisujemy'],
          [['Dane ucznia', 'imię i nazwisko, klasa/oddział, etap edukacyjny, rok szkolny, wychowawca',
            'dokumenty rekrutacyjne, ewidencja szkoły',
            'PESEL, adres zamieszkania, miejsce urodzenia, obywatelstwo'],
           ['Podstawa wsparcia', 'numer i data orzeczenia lub opinii, poradnia wydająca, podstawa wydania orzeczenia, '
            'DATA WPŁYWU do szkoły',
            'orzeczenie, opinia, decyzja dyrektora',
            'kopii treści diagnozy medycznej w treści druku'],
           ['Ścieżka wsparcia', 'kształcenie specjalne (→ IPET) albo pomoc pp (→ PWES)',
            'orzeczenie / opinia / rozpoznanie nauczycieli', '—'],
           ['Rodzice i powrót do domu', 'kontakty, preferowana forma kontaktu, kolejność powiadamiania, '
            'sposób powrotu ucznia do domu',
            'oświadczenia rodziców, wyłącznie na piśmie',
            'danych o miejscu pracy rodziców'],
           ['Zdrowie', 'choroby przewlekłe, leki, dieta, ostrzeżenia, procedura postępowania (kto, na jakiej podstawie, gdzie, kogo)',
            'zaświadczenia lekarskie, zgody rodziców',
            'kopii dokumentacji medycznej, rozpoznań nieistotnych dla nauki'],
           ['Źródła informacji', 'wykaz posiadanych dokumentów o uczniu z datami',
            'teczka ucznia, informacja z przedszkola / poprzedniej szkoły', '—'],
           ['Rejestr kontaktów', 'data, forma, temat, ustalenia, podpis; potwierdzenie przekazania kopii oceny i programu',
            'bieżące notatki wychowawcy i specjalistów', '—'],
           ['Wykaz dokumentacji', 'spis zawartości teczki z datami',
            'bieżąca aktualizacja przy każdym nowym druku', '—']],
          widths=[3.0, 5.4, 4.2, 3.4], size=8)

    h2(doc, 'Jak przygotować metryczkę — krok po kroku')
    steps(doc, [
        '**Sekcja I.** Wpisz imię i nazwisko, klasę, etap edukacyjny, rok szkolny i wychowawcę. Numeru PESEL '
        'i adresu nie powielamy — są w księdze uczniów i w dokumentacji przebiegu nauczania.',
        '**Sekcja II.** Przy każdej formie wsparcia wpisz numer i datę dokumentu, poradnię wydającą oraz podstawę '
        'wydania orzeczenia. Osobno odnotuj DATĘ WPŁYWU orzeczenia do szkoły — od niej biegnie 30 dni na IPET.',
        '**Sekcja III.** Zaznacz ścieżkę: kształcenie specjalne (IPET) albo pomoc psychologiczno-pedagogiczna (PWES). '
        'Ta decyzja porządkuje całą dalszą dokumentację.',
        '**Sekcja IV.** Wypełnij razem z rodzicami: dane kontaktowe, preferowaną formę kontaktu, kolejność '
        'powiadamiania w nagłych wypadkach oraz sposób powrotu ucznia do domu.',
        '**Sekcja V.** Zaznacz występujące pozycje i opisz szczegóły: choroby przewlekłe, leki, dieta, ostrzeżenia. '
        'W polu zaleceń zapisz konkretnie: kto podaje lek, na jakiej podstawie, gdzie lek jest przechowywany '
        'i kogo powiadamiamy w jakiej kolejności. Zaznacz sygnały uruchamiające obserwację pogłębioną.',
        '**Sekcja VI.** Wypisz dokumenty o uczniu, które już posiadasz. To lista tego, czego nie trzeba zbierać drugi raz.',
        '**Sekcja VII.** Notuj każdą rozmowę i ustalenie z rodzicami: data, forma, temat, ustalenia, podpis. '
        'Osobno odnotuj przekazanie kopii oceny i programu — to realizacja prawa rodziców.',
        '**Sekcja VIII.** Prowadź wykaz dokumentacji ucznia z datami. Aktualizuj metryczkę przy każdej zmianie '
        'zgłoszonej przez rodzica i przechowuj ją w miejscu wskazanym zarządzeniem dyrektora.',
    ])


# ================================================================ CZĘŚĆ 5

def czesc_5(doc):
    page_break(doc)
    band(doc, 'CZĘŚĆ 5', 'KSzOF — budowa narzędzia, skala, steny, liczenie wyniku, odczyt profilu',
         '11:20  ·  S5.mp4')

    accent_box(doc, [
        'Rozpoznawanie indywidualnych potrzeb rozwojowych i edukacyjnych oraz możliwości psychofizycznych uczniów — '
        'rozporządzenie o pomocy psychologiczno-pedagogicznej (t.j. Dz.U. 2023 poz. 1798).',
        'Obserwacja pedagogiczna i arkusze obserwacji jako dokumentacja badań i czynności uzupełniających — '
        'rozporządzenie o dokumentacji przebiegu nauczania (t.j. Dz.U. 2024 poz. 50).',
        'ICF (WHO 2001) — model biopsychospołeczny; rozporządzenie o orzekaniu (Dz.U. 2026 poz. 428) nakazuje '
        'opisywać funkcjonowanie w kategoriach aktywności i uczestniczenia — [[§ 7 ust. 7, obowiązuje od 1.09.2026]].',
        'Narzędzie źródłowe: Z. Gajdzica, E. Widawska, S. Byra i in. (2024), Kwestionariusz Szkolnej Oceny '
        'Funkcjonalnej (KSzOF-I-III) — 52 twierdzenia, 9 obszarów ICF, skala 1–5, normy stenowe odrębne '
        'dla nauczyciela i rodzica.',
        'Progi poziomów wsparcia i zasada przeliczania wyniku obszaru na skalę 0–20 są decyzją rady pedagogicznej '
        'wpisaną do procedury szkoły — nie wynikają wprost z przepisu.',
    ], title='PODSTAWA PRAWNA CZĘŚCI')

    h3(doc, 'Transkrypcja narracji')
    narration(doc, [
        'Szanowni Państwo, moduł piąty poświęcamy Kwestionariuszowi Szkolnej Oceny Funkcjonalnej. To najdłuższy moduł '
        'szkolenia, bo kwestionariusz jest sercem całej dokumentacji.',
        'Po tym module będą Państwo znać budowę arkusza, zasady rzetelnej obserwacji, sposób obliczania wyniku, '
        'odczytu stenów oraz odczytu profilu ucznia.',
        'Kwestionariusz jest narzędziem kryterialnym. Opisuje funkcjonowanie ucznia w dziewięciu obszarach '
        'Międzynarodowej Klasyfikacji Funkcjonowania, w codziennych sytuacjach szkolnych i domowych. Przy każdym '
        'twierdzeniu stoi kod klasyfikacji. Dzięki temu nasz opis ucznia mówi tym samym językiem, co dokumentacja '
        'poradni.',
        'Zatrzymajmy się przy klasyfikacji ICF, bo to ona porządkuje całą nową dokumentację. ICF to Międzynarodowa '
        'Klasyfikacja Funkcjonowania, Niepełnosprawności i Zdrowia, opracowana przez Światową Organizację Zdrowia '
        'w dwa tysiące pierwszym roku. Nie opisuje choroby ani rozpoznania. Opisuje, jak człowiek funkcjonuje: '
        'co robi, w czym uczestniczy i co w jego otoczeniu mu pomaga albo przeszkadza.',
        'ICF opiera się na modelu biopsychospołecznym. Funkcjonowanie ucznia jest wypadkową trzech rzeczy: stanu zdrowia '
        'i funkcji ciała, aktywności i uczestniczenia oraz czynników środowiskowych i osobowych. Dwoje uczniów '
        'z tym samym rozpoznaniem może funkcjonować zupełnie inaczej, bo inna jest klasa, inny nauczyciel, '
        'inny poziom hałasu i inne wsparcie w domu. Bez czynników środowiskowych opis ucznia jest niepełny, '
        'a bez opisu barier nie da się zaplanować dostosowań.',
        'Dziewięć obszarów kwestionariusza to dziewięć rozdziałów aktywności i uczestniczenia z klasyfikacji ICF. '
        'Uczenie się i stosowanie wiedzy, kod d sto dziesięć i dalsze. Ogólne zadania i obowiązki, d dwieście dziesięć. '
        'Porozumiewanie się, d trzysta dziesięć. Motoryka i poruszanie się, d czterysta czterdzieści. Dbanie o siebie '
        'i samoobsługa, d pięćset dziesięć. Życie domowe, d sześćset czterdzieści. Wzajemne kontakty i związki, '
        'd siedemset dziesięć. Edukacja szkolna, d osiemset dwadzieścia. I życie w społeczności lokalnej, '
        'd dziewięćset dwadzieścia.',
        'Warto od razu powiedzieć, czym kwestionariusz nie jest. Nie jest diagnozą i nie zastępuje badania '
        'psychologicznego, logopedycznego ani lekarskiego. Jest uporządkowanym zapisem naszej obserwacji '
        'i punktem wyjścia do decyzji zespołu.',
        'Arkusz ma trzy wersje, dopasowane do etapu edukacyjnego. Wersja dla klas pierwszej do trzeciej zawiera '
        'pięćdziesiąt dwa twierdzenia. Wersje dla klas czwartej do szóstej i siódmej do ósmej mają ten sam układ '
        'dziewięciu obszarów, ale twierdzenia sformułowane odpowiednio do wieku i do wymagań drugiego etapu. '
        'O wyborze wersji decyduje wiek rozwojowy ucznia, a nie metrykalny. Ucznia klasy siódmej z niepełnosprawnością '
        'intelektualną w stopniu umiarkowanym obserwujemy wersją dla klas pierwszej do trzeciej, bo tylko ona '
        'da nam użyteczną informację.',
        'Skala ma pięć wartości i zaznaczamy zawsze jedną. Jeden oznacza niewielki stopień występowania zachowania, '
        'dwa mały, trzy umiarkowany, cztery duży, a pięć bardzo duży. Pięć to mocna strona ucznia i tak ją zapisujemy '
        'w ocenie.',
        'Do skali dodajemy w naszej szkole szóstą możliwość: literę N, brak możliwości obserwacji. Pozycję zostawiamy '
        'wtedy pustą. To pełnoprawna i uczciwa odpowiedź. Rodzic nie widzi ucznia na lekcji, a wychowawca nie widzi, '
        'czy uczeń pomaga w pracach domowych. Nie zgadujemy. Litera N nie obniża wyniku obszaru, bo nie wlicza się '
        'do średniej.',
        'Litera N ma jednak jedno ograniczenie i o nim trzeba pamiętać. Normy stenowe zbudowano dla arkusza wypełnionego '
        'w całości. Dlatego, jeżeli w arkuszu jest więcej niż pięć pozycji N, wyniku ogólnego nie przeliczamy na steny. '
        'Odczytujemy wtedy sam profil obszarowy i zaznaczamy w druku, że arkusz jest niepełny. To jest uczciwsze '
        'niż wynik stenowy policzony z połowy danych.',
        'Osiem zasad rzetelnej obserwacji.',
        'Wypełniamy cały arkusz, wszystkie dziewięć obszarów, nie dzieląc ich między oceniających. Oceniamy '
        'na podstawie dwóch do czterech tygodni obserwacji, a nie jednego dnia. Wypełniamy samodzielnie, '
        'bez konsultowania ocen przed spotkaniem zespołu. Oceniamy to, co uczeń robi, a nie to, co potrafiłby zrobić. '
        'Odnosimy się do oczekiwań rozwojowych dla wieku ucznia. Zapisujemy obserwacje jakościowe, czyli konkretne '
        'przykłady zachowań, zwłaszcza przy ocenach skrajnych. W drugim etapie edukacyjnym zbieramy oceny od co najmniej '
        'trzech nauczycieli przedmiotów, bo żaden z nas nie widzi ucznia przez cały dzień. I zwracamy arkusz '
        'koordynatorowi w umówionym terminie, żeby omówić wyniki wspólnie z rodzicami.',
        'Przejdźmy do wypełniania. Posłużymy się przykładem uczennicy klasy trzeciej A.',
        'Obszar pierwszy, uczenie się i stosowanie wiedzy. Twierdzenie drugie: uważnie słucha wypowiedzi nauczyciela, '
        'kolegów, innych. Przez trzy tygodnie obserwowałam, że słucha przy zadaniach na karcie pracy, a traci uwagę '
        'przy pracy z tablicą. Zaznaczam dwa i dopisuję obserwację jakościową. Twierdzenie dwunaste: samodzielnie pisze '
        'zdania i teksty. Przepisuje z tablicy z licznymi opuszczeniami, samodzielnie pisze pojedyncze zdania '
        'z pomocą. Zaznaczam jeden.',
        'Obszar szósty, życie domowe, ma tylko dwa twierdzenia i dotyczy sytuacji, których jako wychowawczyni nie widzę. '
        'Zaznaczam N w obu i pozostawiam ten obszar arkuszowi rodzica. To jest właśnie sens dwóch niezależnych arkuszy.',
        'Teraz obliczamy wynik. Robimy to w trzech krokach.',
        'Krok pierwszy: wynik ogólny. Sumujemy wszystkie punkty. Maksimum wynosi dwieście sześćdziesiąt, '
        'bo pięćdziesiąt dwa twierdzenia razy pięć punktów. W naszym przykładzie suma wynosi sto czternaście punktów.',
        'Krok drugi: odczyt stena z tabeli norm. Uwaga, tabela ma dwie kolumny: osobne przedziały dla nauczyciela '
        'i osobne dla rodzica. To nie jest pomyłka w druku. Nauczyciele i rodzice systematycznie różnią się w ocenie '
        'tych samych zachowań, więc normy to uwzględniają. Sto czternaście punktów w kolumnie nauczyciela to sten '
        'czwarty, czyli wynik niski — sygnał trudności.',
        'Krok trzeci: profil obszarowy. Tu potrzebujemy wspólnej skali, bo obszary mają różną liczbę twierdzeń. '
        'Obszar pierwszy ma piętnaście twierdzeń, a obszar dziewiąty tylko dwa. Porównywanie surowych sum '
        'nie miałoby sensu. Dlatego każdy obszar przeliczamy na skalę od zera do dwudziestu punktów według '
        'jednego wzoru: suma punktów obszaru dzielona przez maksimum obszaru, razy dwadzieścia, zaokrąglone '
        'do pełnego punktu.',
        'Policzmy razem. W obszarze pierwszym uczennica uzyskała dwadzieścia trzy punkty na siedemdziesiąt pięć '
        'możliwych. Dwadzieścia trzy dzielone przez siedemdziesiąt pięć to zero i trzydzieści setnych. Razy dwadzieścia '
        'daje sześć. Zapisujemy: sześć na dwadzieścia. W obszarze trzecim, porozumiewanie się, dwadzieścia dwa punkty '
        'na czterdzieści. To zero i pięćdziesiąt pięć setnych, razy dwadzieścia daje jedenaście.',
        'Progi poziomów wsparcia na tej skali są cztery. Od osiemnastu do dwudziestu to zasób, mocna strona ucznia. '
        'Od czternastu do siedemnastu to poziom pierwszy, wsparcie minimalne. Od dziewięciu do trzynastu to poziom '
        'drugi, wsparcie umiarkowane. Osiem i mniej to poziom trzeci, wsparcie znaczne, wymagające oceny '
        'wielospecjalistycznej i konsultacji z poradnią.',
        'W naszym przykładzie obszar uczenia się, obszar zadań i obowiązków, obszar życia domowego oraz obszar edukacji '
        'szkolnej mieszczą się w poziomie trzecim. Pozostałe pięć obszarów w poziomie drugim. Rekomendowany ogólny '
        'poziom wsparcia to poziom drugi, bo odpowiada przeważającym obszarom, ale w czterech obszarach uczennica '
        'wymaga wsparcia znacznego. Ten zapis wędruje wprost do programu.',
        'I reguła nadrzędna, którą proszę zapamiętać. Każde pojedyncze twierdzenie ocenione na jeden lub dwa podlega '
        'analizie zespołu, niezależnie od wyniku obszaru i niezależnie od stena. Uczeń może mieć w obszarze '
        'porozumiewania się wynik czternaście na dwadzieścia, a jednocześnie jedynkę przy twierdzeniu o proszeniu '
        'o pomoc w sytuacji trudnej. Średnia to zamaskuje. Reguła nadrzędna nie pozwoli tego przeoczyć.',
        'Profil odczytujemy w trzech krokach. Najpierw kolor: zielony to zasób i poziom pierwszy, żółty to poziom drugi, '
        'czerwony to poziom trzeci. Potem kształt. Profil płaski i niski wskazuje na trudność globalną i kieruje nas '
        'ku ocenie wielospecjalistycznej. Profil z jednym głębokim wcięciem wskazuje na trudność wybiórczą i konkretny '
        'moduł pogłębiony. Na końcu wracamy do pojedynczych twierdzeń ocenionych nisko, także w obszarach zielonych.',
        'Na zakończenie sprawa arkuszy równoległych. Ten sam kwestionariusz wypełniają niezależnie wychowawca i rodzic, '
        'a w drugim etapie edukacyjnym dodatkowo nauczyciel przedmiotu i specjalista. Na spotkaniu zespołu kładziemy '
        'profile obok siebie i nie uśredniamy ich. Rozbieżność jest cenną informacją o tym, że uczeń funkcjonuje '
        'inaczej w różnych środowiskach. Jeśli w domu jest wyżej niż w szkole, szukamy barier w klasie. Jeśli w szkole '
        'wyżej niż w domu, dzielimy się z rodzicami sprawdzonymi rutynami. A jeśli różnią się dwaj nauczyciele '
        'przedmiotów, pytamy, czym różnią się ich lekcje.',
        'Podsumowując. Kwestionariusz wypełniamy dwa razy w roku. We wrześniu jako pomiar bazowy, dla wszystkich uczniów '
        'w oddziale. Uczeń z orzeczeniem musi mieć program do trzydziestego września, a program opiera się na ocenie '
        'funkcjonowania. I w maju jako pomiar kontrolny, na tym samym arkuszu, innym kolorem. Wtedy widać drogę, '
        'którą uczeń przeszedł. W module szóstym omówimy, kiedy i jak uruchamiamy obserwację pogłębioną.',
    ])

    h2(doc, 'Budowa arkusza — dziewięć obszarów, 52 twierdzenia, 260 punktów')
    table(doc,
          ['Obszar ICF', 'Nazwa', 'Twierdzenia', 'Maks. pkt', 'Kody ICF'],
          [['I', 'Uczenie się i stosowanie wiedzy', '1–15 (15)', '75', 'd110–d177'],
           ['II', 'Ogólne zadania i obowiązki', '16–21 (6)', '30', 'd210–d240'],
           ['III', 'Porozumiewanie się', '22–29 (8)', '40', 'd310–d350'],
           ['IV', 'Motoryka i poruszanie się', '30–33 (4)', '20', 'd440–d450'],
           ['V', 'Dbanie o siebie i samoobsługa', '34–37 (4)', '20', 'd510–d570'],
           ['VI', 'Życie domowe (obszar rodzicielski)', '38–39 (2)', '10', 'd640'],
           ['VII', 'Wzajemne kontakty i związki', '40–47 (8)', '40', 'd710–d760'],
           ['VIII', 'Edukacja szkolna', '48–50 (3)', '15', 'd820'],
           ['IX', 'Życie w społeczności lokalnej', '51–52 (2)', '10', 'd920'],
           ['RAZEM', '—', '52', '260', 'd1–d9']],
          widths=[1.8, 6.4, 2.6, 2.0, 4.2], size=8.5)

    h2(doc, 'Tabela norm stenowych (klasy I–III) — odrębne przedziały dla nauczyciela i rodzica')
    table(doc,
          ['Opis wyniku', 'Sten', 'Nauczyciel', 'Rodzic', 'Odczyt'],
          [['WYNIK WYSOKI\nbardzo dobre funkcjonowanie', '10\n9\n8', '238–260\n218–237\n197–217',
            '233–260\n214–232\n194–213', 'Zasób — opisujemy w mocnych stronach.'],
           ['WYNIK PRZECIĘTNY\ntypowe funkcjonowanie', '7\n6\n5', '175–196\n154–174\n135–153',
            '172–193\n151–171\n132–150', 'Standardowy poziom rozwojowy.'],
           ['WYNIK NISKI\nsygnał trudności', '4\n3\n2\n1', '113–134\n91–112\n72–90\n52–71',
            '110–131\n90–109\n70–89\n52–69', 'Poziom I, II lub III wsparcia — patrz profil obszarowy.']],
          widths=[4.6, 1.6, 3.2, 3.2, 4.4], size=8.5)

    h2(doc, 'Przeliczanie obszaru na skalę 0–20 i progi poziomów wsparcia')
    accent_box(doc, [
        'WZÓR: wynik obszaru w skali 0–20  =  (suma punktów obszaru ÷ maksimum obszaru) × 20, '
        'zaokrąglone do pełnego punktu. Obszary mają różną liczbę twierdzeń, więc bez przeliczenia nie da się '
        'ich porównać na jednym profilu.',
        'PROGI: 18–20 pkt = zasób · 14–17 pkt = Poziom I (wsparcie minimalne) · 9–13 pkt = Poziom II '
        '(wsparcie umiarkowane) · 0–8 pkt = Poziom III (wsparcie znaczne).',
        'REGUŁA NADRZĘDNA: każde pojedyncze twierdzenie ocenione na 1 lub 2 podlega analizie zespołu — '
        'niezależnie od wyniku obszaru i niezależnie od stena.',
        'ARKUSZ NIEPEŁNY: przy więcej niż pięciu pozycjach N wyniku ogólnego nie przeliczamy na steny; '
        'odczytujemy sam profil obszarowy i zaznaczamy w druku, że arkusz jest niepełny.',
    ], title='ZASADY LICZENIA — DECYZJA RADY PEDAGOGICZNEJ WPISANA DO PROCEDURY')

    h2(doc, 'Przykład obliczeniowy — Zofia Lewandowska, klasa III A, arkusz wychowawcy, wrzesień 2026')
    table(doc,
          ['Obszar', 'Suma / maks.', 'Skala 0–20', 'Sten obszaru*', 'Poziom', 'Powiązana sfera IPET'],
          [['I. Uczenie się i stosowanie wiedzy', '23 / 75', '6', '3', 'III', 'Sfera 1 — poznawcze'],
           ['II. Ogólne zadania i obowiązki', '10 / 30', '7', '4', 'III', 'Sfera 6 — samodzielność'],
           ['III. Porozumiewanie się', '22 / 40', '11', '5', 'II', 'Sfera 4 — komunikacja'],
           ['IV. Motoryka i poruszanie się', '12 / 20', '12', '5', 'II', 'Sfera 3 — motoryczne'],
           ['V. Dbanie o siebie i samoobsługa', '10 / 20', '10', '5', 'II', 'Sfera 5 — samoobsługa'],
           ['VI. Życie domowe (arkusz rodzica)', '4 / 10', '8', '4', 'III', 'Sfera 6 — samodzielność'],
           ['VII. Wzajemne kontakty i związki', '22 / 40', '11', '5', 'II', 'Sfera 2 — społeczno-emocjonalne'],
           ['VIII. Edukacja szkolna', '5 / 15', '7', '4', 'III', 'Sfera 1 — poznawcze'],
           ['IX. Życie w społeczności lokalnej', '6 / 10', '12', '6', 'II', 'Sfera 2 — społeczno-emocjonalne'],
           ['WYNIK OGÓLNY', '114 / 260', '—', 'sten 4 — wynik niski', 'ogólnie II', 'IPET/2026-2027/III A/07']],
          widths=[5.4, 2.4, 1.9, 3.0, 1.6, 4.3], size=8)
    para(doc, '* Sten obszarowy jest odczytem orientacyjnym, pomocnym przy rozmowie z rodzicem. Wiążący dla kwalifikacji '
              'do poziomu wsparcia jest wynik w skali 0–20 i próg z procedury szkoły. Wynik ogólny w stenach '
              'odczytujemy wyłącznie z tabeli norm.',
         size=8, italic=True, color=GREY, after=8)

    h2(doc, 'Jak przygotować i policzyć kwestionariusz — krok po kroku')
    steps(doc, [
        '**Wybierz wersję arkusza** odpowiednią do etapu i wieku rozwojowego ucznia: I–III, IV–VI albo VII–VIII. '
        'O wyborze decyduje wiek rozwojowy, nie metrykalny.',
        '**Wypełnij metryczkę arkusza:** uczeń, klasa, okres obserwacji, osoba wypełniająca i jej rola. '
        'Arkusz wypełniają niezależnie wychowawca i rodzic; w II etapie dodatkowo nauczyciel przedmiotu i specjalista.',
        '**Obserwuj w naturalnych sytuacjach** przez dwa do czterech tygodni, zanim zaznaczysz odpowiedzi. '
        'Oceniaj zachowanie, nie wrażenie.',
        '**Zaznacz jedną wartość skali 1–5 przy każdym twierdzeniu.** Odpowiedź N — nie obserwowano — jest '
        'pełnoprawna i nie obniża wyniku, ale przy więcej niż pięciu N nie przeliczamy wyniku na steny.',
        '**Policz wynik ogólny:** suma wszystkich punktów, maksimum 260. Odczytaj sten z właściwej kolumny — '
        'osobnej dla nauczyciela, osobnej dla rodzica.',
        '**Przelicz każdy obszar na skalę 0–20** według wzoru (suma ÷ maksimum) × 20 i przypisz poziom wsparcia '
        'według progów z procedury szkoły.',
        '**Zastosuj regułę nadrzędną:** wypisz osobno każde twierdzenie ocenione na 1 lub 2 — niezależnie od wyniku obszaru.',
        '**Nanieś wyniki na profil dziewięciu obszarów.** Zielony to zasób i Poziom I, żółty to Poziom II, '
        'czerwony to Poziom III.',
        '**Połóż profile obok siebie i nie uśredniaj ich.** Rozbieżność wychowawca–rodzic lub nauczyciel–nauczyciel '
        'jest informacją o środowisku, nie błędem pomiaru.',
        '**Powtórz kwestionariusz w maju** na tym samym arkuszu, innym kolorem. Wtedy widać drogę, którą uczeń przeszedł.',
    ])


# ================================================================ CZĘŚĆ 6

def czesc_6(doc):
    page_break(doc)
    band(doc, 'CZĘŚĆ 6', 'Obserwacja pogłębiona — ABC i FBA, profil sensoryczny, teoria umysłu, karta mowy',
         '10:30  ·  S6.mp4')

    accent_box(doc, [
        'Rozpoznawanie indywidualnych potrzeb i możliwości psychofizycznych ucznia oraz ocena efektywności udzielanej '
        'pomocy — rozporządzenie o pomocy psychologiczno-pedagogicznej (t.j. Dz.U. 2023 poz. 1798).',
        'Wystąpienie do poradni, gdy mimo udzielanej pomocy nie ma poprawy — za zgodą rodziców, wniosek dyrektora.',
        'Arkusze obserwacji pogłębionej jako dokumentacja badań i czynności uzupełniających — rozporządzenie '
        'o dokumentacji przebiegu nauczania (t.j. Dz.U. 2024 poz. 50).',
        'Reguły przekierowania i karta decyzyjna nie wynikają wprost z przepisu — to decyzja rady pedagogicznej '
        'wpisana do procedury szkoły.',
        'Granica kompetencji: nauczyciel opisuje obserwowane zachowanie; rozpoznanie i kwalifikacja do terapii '
        'należą do specjalisty (logopedy, psychologa, terapeuty integracji sensorycznej, lekarza).',
    ], title='PODSTAWA PRAWNA CZĘŚCI')

    h3(doc, 'Transkrypcja narracji')
    narration(doc, [
        'Szanowni Państwo, moduł szósty poświęcamy obserwacji pogłębionej. Kwestionariusz powiedział nam, gdzie uczeń '
        'potrzebuje wsparcia. Obserwacja pogłębiona odpowiada na pytanie, dlaczego.',
        'Po tym module będą Państwo wiedzieć, kiedy uruchamiamy obserwację pogłębioną, które z czterech narzędzi wybrać '
        'i gdzie przebiega granica kompetencji nauczyciela.',
        'Przesiew obejmuje wszystkich uczniów. Obserwacja pogłębiona obejmuje pojedynczego ucznia. Uruchamiamy ją wtedy, '
        'gdy wynik kwestionariusza wskaże kierunek. To kilkanaście godzin pracy zespołu, dlatego uruchamiamy ją '
        'z przesłanką, a nie na wszelki wypadek.',
        'Siedem reguł przekierowania. Wystarczy jedna, aby zespół usiadł nad kartą decyzyjną.',
        'Reguła pierwsza: wynik któregokolwiek obszaru wynoszący osiem punktów lub mniej w skali od zera do dwudziestu, '
        'czyli poziom trzeci. Reguła druga: dwa lub więcej twierdzeń ocenionych na jeden lub dwa w tym samym obszarze. '
        'Reguła trzecia: rozbieżność między oceniającymi wynosząca dwa steny lub więcej w wyniku ogólnym. '
        'Reguła czwarta: sygnał zdrowotny z metryczki, na przykład nadwrażliwość sensoryczna albo choroba przewlekła '
        'wpływająca na uczestnictwo. Reguła piąta: zachowanie powtarzalne, które zagraża uczniowi lub innym — '
        'w tym przypadku uruchamiamy obserwację natychmiast, nie czekając na zespół. Reguła szósta: brak poprawy '
        'mimo udzielanej pomocy przez około trzy miesiące. I reguła siódma, specyficznie szkolna: nagła zmiana '
        'funkcjonowania — obniżenie ocen w co najmniej trzech przedmiotach albo nieobecności przekraczające '
        'dwadzieścia procent zajęć w miesiącu.',
        'Reguły te są naszą decyzją jako rady pedagogicznej, wpisaną do procedury szkoły. Przepis wymaga rozpoznawania '
        'potrzeb i oceny efektywności. Reguły sprawiają, że decyzja nie zależy od tego, kto danego dnia patrzy '
        'na arkusz.',
        'Mamy cztery narzędzia i każde odpowiada na inne pytanie.',
        'Pierwsze narzędzie to model ABC, czyli analiza zachowania, oraz jego rozszerzenie — arkusz analizy funkcjonalnej.',
        'Litera A oznacza poprzednik, czyli to, co działo się bezpośrednio przed zachowaniem. Litera B oznacza zachowanie, '
        'opisane obserwowalnie i mierzalnie. Litera C oznacza następstwo, czyli to, co stało się bezpośrednio potem, '
        'w tym reakcję dorosłych i rówieśników.',
        'Posłuchajmy przykładu poprawnego zapisu, szkolnego. Na trzeciej lekcji matematyki nauczyciel polecił '
        'przepisać z tablicy zadanie o sześciu linijkach. Uczeń pisał przez około dwie minuty, po czym odłożył ołówek, '
        'położył głowę na ławce i przestał reagować na polecenia przez około sześć minut. Nauczyciel podszedł, '
        'podzielił zadanie na trzy części i zaznaczył pierwszą. Uczeń przepisał zaznaczoną część. Każde zdanie '
        'da się sprawdzić. Żadne nie zawiera interpretacji.',
        'A teraz zapis wadliwy, który spotykamy w dokumentacji najczęściej. Uczeń zniechęcił się, bo nie chciało mu się '
        'pracować, i zamanifestował swoją niechęć do przedmiotu. Trzy interpretacje w jednym zdaniu i ani jednego '
        'faktu, który dałoby się policzyć.',
        'Model ABC uruchamiamy, gdy zachowanie trudne jest powtarzalne, zagraża albo przerywa uczestnictwo ucznia '
        'w zajęciach. Zbieramy co najmniej dziesięć do piętnastu zapisów w ciągu dwóch, trzech tygodni, '
        'z różnych lekcji i różnych pór dnia. Szukamy funkcji zachowania: uzyskania uwagi, uzyskania przedmiotu '
        'lub aktywności, uniknięcia trudnego zadania albo regulacji pobudzenia. Interpretację formułujemy dopiero '
        'na spotkaniu zespołu, po analizie wielu zdarzeń, i zapisujemy ją w arkuszu analizy funkcjonalnej: '
        'hipoteza funkcji, dane, które ją potwierdzają, i plan zachowania zastępczego.',
        'To ostatnie jest w szkole kluczowe. Nie wystarczy ustalić, po co uczeń zachowuje się w dany sposób. Trzeba '
        'wskazać, jakim zachowaniem może osiągnąć ten sam cel, i tego zachowania go nauczyć. Jeżeli uczeń krzykiem '
        'unika trudnego zadania, uczymy go karty z napisem „proszę o przerwę” i honorujemy ją natychmiast. '
        'Inaczej odbieramy mu jedyny działający sposób i nie dajemy nic w zamian.',
        'Drugie narzędzie to profil sensoryczny.',
        'To ustrukturyzowana obserwacja reakcji ucznia na bodźce w siedmiu układach: słuchowym, wzrokowym, dotykowym, '
        'węchowym, smakowym, przedsionkowym i proprioceptywnym. Opisujemy wzorzec, nie stawiamy rozpoznania.',
        'Wyróżniamy trzy wzorce. Nadreaktywność, gdy uczeń zakrywa uszy przy dzwonku, unika tłoku na korytarzu, '
        'nie znosi metek i odmawia potraw o określonej konsystencji. Podreaktywność, gdy nie reaguje na wołanie mimo '
        'prawidłowego słuchu, nie zauważa zabrudzonych rąk i wolno rozpoczyna czynności. Poszukiwanie bodźców, gdy '
        'buja się na krześle, wpada na przedmioty, gryzie ubrania i mówi bardzo głośno.',
        'W szkole warto zwrócić uwagę na trzy miejsca, które w przedszkolu nie istnieją albo wyglądają inaczej: '
        'korytarz na przerwie, stołówka i sala gimnastyczna. To tam natężenie bodźców jest największe i tam najczęściej '
        'dochodzi do zachowań trudnych, które potem opisujemy jako zachowania na lekcji.',
        'Profil sensoryczny uruchamiamy na sygnał z metryczki albo przy skupisku niskich ocen w obszarach dbania '
        'o siebie i poruszania się. Granica kompetencji jest jasna. Nauczyciel opisuje obserwowane reakcje. '
        'Rozpoznanie i kwalifikacja do terapii należą do terapeuty integracji sensorycznej. W dokumentacji piszemy: '
        'obserwowany wzorzec poszukiwania bodźców, wskazana konsultacja specjalisty.',
        'Trzecie narzędzie to obserwacja poznania społecznego i teorii umysłu.',
        'Teoria umysłu to zdolność przypisywania sobie i innym stanów umysłu: wiedzy, przekonań, intencji i emocji, '
        'oraz rozumienia, że mogą one różnić się od naszych. To fundament współpracy w grupie, żartu, ironii, '
        'pracy projektowej i rozumienia tekstu literackiego.',
        'Kiedy taka obserwacja jest potrzebna? Gdy obszar wzajemnych kontaktów jest wyraźnie niższy przy zachowanych '
        'obszarach uczenia się i poruszania. Gdy uczeń ma trudność w odczytywaniu mimiki, gestu i tonu głosu. '
        'Gdy rozumie język dosłownie i nie rozpoznaje żartu ani ironii. Gdy nie odróżnia przypadkowego potrącenia '
        'od celowego zaczepienia i reaguje na oba tak samo. Gdy w pracy grupowej nie potrafi przyjąć perspektywy '
        'kolegi. I gdy zespół rozważa wystąpienie do poradni w sprawie całościowych zaburzeń rozwoju.',
        'Wiek ma tu zasadnicze znaczenie. Zadania fałszywego przekonania pierwszego rzędu rozwiązują dzieci typowo '
        'rozwijające się od około czwartego roku życia, a zadania drugiego rzędu, czyli „co Ola myśli, że Kuba myśli”, '
        'od około szóstego, siódmego roku. W szkole podstawowej obserwujemy więc przede wszystkim rozumienie ironii, '
        'żartu, obietnicy, kłamstwa uprzejmościowego i intencji stojącej za zachowaniem rówieśnika.',
        'Obserwacja teorii umysłu w szkole opisuje zachowania. Nie jest testem i nie prowadzi do rozpoznania. '
        'Prowadzi do rzetelnego opisu, który przekażemy poradni.',
        'Czwarte narzędzie to karta obserwacji rozwoju mowy i komunikacji.',
        'Trudności w komunikacji bywają przyczyną wielu innych problemów. Uczeń, który nie rozumie polecenia, wygląda '
        'na nieposłusznego. Uczeń, który nie potrafi powiedzieć, czego chce, krzyczy albo wychodzi z sali. Zanim '
        'uznamy, że problem leży w zachowaniu, sprawdzamy, czy uczeń nas rozumie i czy potrafi się z nami porozumieć.',
        'Karta obejmuje pięć obszarów: rozumienie mowy, mowę czynną i artykulację, słuch fonematyczny, słownictwo '
        'i gramatykę oraz komunikację i budowanie wypowiedzi. To dwadzieścia pięć wskaźników w skali od zera do dwóch. '
        'Zero, gdy zachowanie nie występuje. Jeden, gdy pojawia się częściowo lub z pomocą dorosłego. Dwa, '
        'gdy uczeń robi to samodzielnie. Każdy obszar daje wynik od zera do dziesięciu punktów.',
        'W szkole dodajemy do karty jeden obszar, którego nie ma w wersji przedszkolnej: technikę czytania i pisania '
        'w powiązaniu ze słuchem fonematycznym. Bo trudność w wyodrębnianiu głosek, która w przedszkolu jest jedną '
        'z wielu, w klasie pierwszej staje się barierą w każdym przedmiocie.',
        'Odczyt jest prosty. Od ośmiu do dziesięciu punktów to zasób. Od czterech do siedmiu to obszar wymagający '
        'wsparcia. Od zera do trzech to priorytet. Wynik zasila cele, ocenę wielospecjalistyczną i opinię dla poradni.',
        'Granica kompetencji jest taka sama jak przy pozostałych narzędziach. Nauczyciel opisuje, co słyszy i widzi '
        'w naturalnych sytuacjach. Diagnoza logopedyczna należy do logopedy. W dokumentacji piszemy: obserwowane '
        'trudności w wyodrębnianiu głosek i w budowaniu wypowiedzi, wskazana diagnoza logopedyczna.',
        'Podsumowując. Moduł zamyka karta decyzyjna, jedna na jednego ucznia. Sprawdzamy siedem reguł, wybieramy '
        'narzędzie, wpisujemy, kto obserwuje, od kiedy i kiedy spotyka się zespół. Kartę wpinamy do teczki również '
        'wtedy, gdy decyzja brzmi: nie uruchamiamy.',
        'Zapis decyzji odmownej jest równie ważny jak zapis decyzji pozytywnej. Pokazuje, że zespół sprawę rozważył, '
        'i chroni nas wtedy, gdy pół roku później ktoś zapyta, dlaczego nic nie zrobiono. W module siódmym przejdziemy '
        'od zebranych danych do oceny, programu i ewaluacji.',
    ])

    h2(doc, 'Siedem reguł przekierowania — wystarczy jedna')
    table(doc,
          ['Reguła', 'Kryterium', 'Rekomendowane narzędzie'],
          [['1', 'Wynik któregokolwiek obszaru ≤ 8 pkt w skali 0–20 (Poziom III).',
            'Zależnie od obszaru — patrz kolumna „obszar → narzędzie” poniżej.'],
           ['2', 'Dwa lub więcej twierdzeń ocenionych na 1 lub 2 w tym samym obszarze.',
            'Narzędzie właściwe dla obszaru.'],
           ['3', 'Rozbieżność między oceniającymi ≥ 2 steny w wyniku ogólnym.',
            'Rozmowa z rodzicem + obserwacja w środowisku o wyższym wyniku.'],
           ['4', 'Sygnał zdrowotny z metryczki (nadwrażliwość sensoryczna, choroba przewlekła).',
            'Profil sensoryczny (druk 4).'],
           ['5', 'Zachowanie powtarzalne zagrażające uczniowi lub innym — uruchamiamy NATYCHMIAST.',
            'ABC + arkusz analizy funkcjonalnej (druki 2 i 3).'],
           ['6', 'Brak poprawy mimo udzielanej pomocy przez ok. 3 miesiące.',
            'Karta oceny efektywności + narzędzie właściwe dla obszaru.'],
           ['7', 'Nagła zmiana funkcjonowania: obniżenie ocen w ≥ 3 przedmiotach albo nieobecności > 20% w miesiącu.',
            'Wywiad z uczniem i rodzicem; ocena efektywności; konsultacja psychologa.']],
          widths=[1.5, 8.2, 7.0], size=8.5)

    h2(doc, 'Który obszar KSzOF prowadzi do którego narzędzia')
    table(doc,
          ['Obszar z niskim wynikiem', 'Pytanie', 'Narzędzie', 'Kto prowadzi'],
          [['III. Porozumiewanie się', 'Czy uczeń nas rozumie i czy potrafi się porozumieć?',
            'Karta obserwacji rozwoju mowy i komunikacji (druk 5)', 'wychowawca + logopeda'],
           ['II. Ogólne zadania i obowiązki\nVIII. Edukacja szkolna', 'Dlaczego zachowanie się powtarza?',
            'ABC (druk 2) + arkusz analizy funkcjonalnej FBA (druk 3)', 'wychowawca + psycholog / pedagog specjalny'],
           ['IV. Motoryka\nV. Dbanie o siebie', 'Jak uczeń reaguje na bodźce i czego potrzebuje jego układ nerwowy?',
            'Profil sensoryczny ucznia (druk 4)', 'wychowawca + terapeuta SI (konsultacja)'],
           ['VII. Wzajemne kontakty i związki', 'Czy uczeń rozumie intencje i przekonania innych?',
            'Arkusz obserwacji poznania społecznego i ToM (druk 6)', 'psycholog + wychowawca'],
           ['I. Uczenie się i stosowanie wiedzy', 'Które funkcje uczenia się są obniżone i przy jakim wsparciu rosną?',
            'Karta obserwacji funkcji uczenia się (druk 7)', 'wychowawca + terapeuta pedagogiczny'],
           ['IX. Życie w społeczności lokalnej\nVI. Życie domowe', 'Jakie są zasoby i ograniczenia środowiska?',
            'Ankieta dla rodzica i wywiad środowiskowy (druk 9)', 'pedagog + wychowawca']],
          widths=[4.0, 4.8, 5.2, 3.7], size=8)

    h2(doc, 'Jak przeprowadzić obserwację pogłębioną — krok po kroku')
    steps(doc, [
        '**Sprawdź siedem reguł przekierowania.** Wystarczy jedna spełniona, żeby uruchomić moduł pogłębiony.',
        '**Załóż kartę decyzyjną** — jedną na jednego ucznia (druk 1). Zapisz, która reguła zadziałała i jakie '
        'narzędzie wybieramy. Kartę wpinamy do teczki także wtedy, gdy decyzja brzmi „nie uruchamiamy”.',
        '**Model ABC stosuj wtedy, gdy pytanie brzmi: dlaczego to zachowanie się powtarza.** Notuj trzy kolumny: '
        'co było przed, zachowanie, co nastąpiło po. Zbierz 10–15 zapisów z różnych lekcji i pór dnia.',
        '**Zamknij ABC arkuszem analizy funkcjonalnej:** hipoteza funkcji, dane potwierdzające, zachowanie zastępcze, '
        'którego uczymy, i sposób jego honorowania. Bez zachowania zastępczego plan nie działa.',
        '**Profil sensoryczny stosuj, gdy uczeń reaguje nadmiernie albo zbyt słabo na bodźce.** Opisz wzorzec '
        'i wskaż konsekwencję dla organizacji klasy, korytarza, stołówki i sali gimnastycznej.',
        '**Obserwację poznania społecznego prowadź, gdy trudność dotyczy rozumienia intencji i przekonań innych.** '
        'To nie jest test i nie prowadzi do rozpoznania.',
        '**Kartę obserwacji rozwoju mowy uruchom, gdy uczeń nie rozumie poleceń albo nie buduje wypowiedzi.** '
        'Brak rozumienia bywa przyczyną trudności, które wyglądają na nieposłuszeństwo.',
        '**Zapisz wniosek:** co obserwowaliśmy, jak często, w jakich sytuacjach i co z tego wynika dla wsparcia. '
        'Wnioski przechodzą wprost do sekcji VI–IX wielospecjalistycznej oceny.',
    ])


# ================================================================ CZĘŚĆ 7

def czesc_7(doc):
    page_break(doc)
    band(doc, 'CZĘŚĆ 7', 'WOPF-SP, IPET, PWES, cele SMART, ewaluacja i opinia dla poradni', '14:10  ·  S7.mp4')

    accent_box(doc, [
        'WOPFU i IPET — § 6 rozporządzenia MEN z 9 sierpnia 2017 r. (t.j. Dz.U. 2020 poz. 1309): ust. 1 pkt 1–8 '
        '(zawartość programu), ust. 4 (program po ocenie, z uwzględnieniem zaleceń z orzeczenia), '
        'oraz dalsze ustępy § 6 — ocena co najmniej dwa razy w roku i prawa rodziców do udziału i do kopii. '
        '[[Numeru ustępu nie wpisujemy do dokumentu ucznia bez sprawdzenia w ogłoszonym tekście jednolitym — '
        'ustalenie A1 audytu skryptu przedszkolnego z 5.09.2026.]]',
        'Terminy IPET: do 30 września dla ucznia rozpoczynającego kształcenie z orzeczeniem albo 30 dni od dnia '
        'złożenia orzeczenia w szkole.',
        'Formy pomocy psychologiczno-pedagogicznej i ocena efektywności — rozporządzenie o pomocy pp '
        '(t.j. Dz.U. 2023 poz. 1798); zindywidualizowana ścieżka kształcenia — § 12 tego rozporządzenia.',
        'Opinia o funkcjonowaniu ucznia — § 7 ust. 2–3 rozporządzenia o orzekaniu (Dz.U. 2026 poz. 428): wydaje się '
        'w terminie 10 dni od dnia otrzymania przez dyrektora prośby o jej wydanie; kopię otrzymują rodzice.',
        'Dostosowanie wymagań edukacyjnych i ocena zachowania ucznia z orzeczeniem lub opinią — rozporządzenie '
        'o ocenianiu z 22 lutego 2019 r. ([[t.j. Dz.U. 2023 poz. 2572, z późn. zm.]]).',
        '[[Plan wsparcia edukacyjnego ucznia (PWES) nie jest drukiem wymaganym przez rozporządzenie. Formy, okres '
        'i wymiar pomocy psychologiczno-pedagogicznej ustala dyrektor, a przebieg zajęć dokumentuje się w dziennikach. '
        'PWES wprowadzamy zarządzeniem dyrektora jako narzędzie wewnętrzne — tak samo jak metryczkę.]]',
        'Cele SMART nie są nazwane w rozporządzeniu. Wymagana jest ocena efektywności, a cel z kryterium jest '
        'najprostszym sposobem, żeby ją przeprowadzić.',
    ], title='PODSTAWA PRAWNA CZĘŚCI')

    h3(doc, 'Transkrypcja narracji')
    narration(doc, [
        'Szanowni Państwo, w ostatnim module przechodzimy od zebranych danych do zobowiązania. Omówimy ocenę '
        'wielospecjalistyczną, program, plan wsparcia dla ucznia bez orzeczenia, cele mierzalne, ewaluację '
        'oraz opinię dla poradni.',
        'Po tym module będą Państwo potrafili zbudować ocenę z danych, które już mają, sformułować cel mierzalny '
        'i zaplanować ewaluację w kalendarzu roku szkolnego.',
        'Zaczynamy od wielospecjalistycznej oceny poziomu funkcjonowania ucznia.',
        'Podstawą jest rozporządzenie w sprawie kształcenia specjalnego. Zespół dokonuje oceny co najmniej dwa razy '
        'w roku szkolnym, a ocena jest podstawą opracowania i modyfikacji programu.',
        'Dlaczego ocena jest tak ważna? Po pierwsze, wyznacza punkt startu, bez którego cele byłyby zgadywaniem. '
        'Po drugie, scala perspektywy wychowawcy, nauczycieli przedmiotów, logopedy, psychologa i rodziców w jeden '
        'obraz ucznia. Po trzecie, chroni ucznia przed etykietą, bo dwoje uczniów z tym samym orzeczeniem otrzymuje '
        'dwie różne oceny. Po czwarte, dokumentuje pracę szkoły.',
        'Nasz druk oceny dla szkoły podstawowej ma dwadzieścia trzy sekcje, ale nie jest to dwadzieścia trzy razy '
        'więcej pracy. To jest jeden druk, który zbiera wyniki z druków wcześniejszych, i przy każdej sekcji jest '
        'napisane, skąd te wyniki brać.',
        'Sekcja pierwsza to dane ucznia w zakresie minimalnym, a sekcja pierwsza a to wybór ścieżki: uczeń z orzeczeniem '
        'idzie do programu, uczeń bez orzeczenia do planu wsparcia. Sekcje od drugiej do czwartej to zespół, mapa '
        'dokumentów źródłowych i informacje medyczne istotne dla funkcjonowania. Sekcja piąta to wyniki kwestionariusza '
        'w dziewięciu obszarach, a sekcja piąta a rozpisuje każdy obszar na mocne strony i trudności. Sekcje szósta '
        'do dziesiątej to obserwacja pogłębiona: zachowanie i jego funkcje, poznanie społeczne, mowa i komunikacja, '
        'przetwarzanie sensoryczne. Sekcje jedenasta do trzynastej to synteza rozpoznania: całościowy obraz, '
        'indywidualne potrzeby, przyczyny niepowodzeń i bariery. Sekcje czternasta do dziewiętnastej opisują sposób '
        'pracy. Sekcje dwudziesta do dwudziestej drugiej to decyzja zespołu, cele i ocena efektywności. '
        'A sekcja dwudziesta trzecia, w dwóch kolumnach, wskazuje, co przenosimy do programu, a co do planu wsparcia.',
        'Każdy blok ma swoje źródło w dokumentach, które już omówiliśmy. Mocne strony bierzemy z ocen najwyższych '
        'w kwestionariuszu. Trudności z ocen najniższych i z obserwacji pogłębionej. Bariery i ułatwienia z analizy ABC '
        'i profilu sensorycznego. Efekty wsparcia z poprzedniej karty ewaluacji. Głos ucznia z ankiety Mój głos.',
        'Przechodzimy do indywidualnego programu edukacyjno-terapeutycznego.',
        'Program opracowujemy dla ucznia posiadającego orzeczenie o potrzebie kształcenia specjalnego. Terminy są dwa. '
        'Do trzydziestego września dla ucznia, który rozpoczyna kształcenie z orzeczeniem. Albo trzydzieści dni '
        'od dnia złożenia orzeczenia w szkole, niezależnie od miesiąca.',
        'Pamiętajmy o prawach rodziców. Mają prawo uczestniczyć w spotkaniach zespołu oraz otrzymać kopię programu '
        'i kopię oceny. Dyrektor zawiadamia ich o terminie spotkania w sposób przyjęty w szkole, a przekazanie kopii '
        'odnotowujemy w rejestrze kontaktów.',
        'Co musi znaleźć się w programie? Rozporządzenie wymienia osiem elementów i warto je znać na pamięć, '
        'bo to jest lista kontrolna każdej kontroli. Zakres i sposób dostosowania wymagań. Zintegrowane działania '
        'nauczycieli i specjalistów. Formy, okres i wymiar godzin pomocy. Działania wspierające rodziców oraz zakres '
        'współdziałania z poradniami i instytucjami. Zajęcia rewalidacyjne, resocjalizacyjne i socjoterapeutyczne, '
        'a w klasach starszych także zajęcia związane z wyborem kierunku kształcenia i zawodu. Zakres współpracy '
        'nauczycieli i specjalistów z rodzicami. Rodzaj i sposób dostosowania warunków organizacji kształcenia. '
        'I ewentualnie wskazanie zajęć realizowanych indywidualnie lub w grupie do pięciu osób.',
        'Do programu wpisujemy zalecenia poradni z orzeczenia, jedno po drugim, a przy każdym zapisujemy, jak je '
        'realizujemy: w jakiej formie, kto i w jakim wymiarze. To samo robimy z zaleceniami z oceny wielospecjalistycznej. '
        'Orzeczenie mówi, co uczniowi zalecono. Ocena mówi, co widzimy w szkole. Program pokazuje, jak jedno i drugie '
        'zamieniamy w działanie.',
        'Zatrzymajmy się przy dostosowaniach, bo w szkole wyglądają inaczej niż w przedszkolu. Dostosowanie to zmiana '
        'w tym, jak uczymy i jak sprawdzamy, a nie w tym, czego uczymy. Podstawa programowa pozostaje ta sama. '
        'Wyjątkiem jest uczeń z niepełnosprawnością intelektualną w stopniu umiarkowanym lub znacznym, dla którego '
        'obowiązuje odrębna podstawa programowa.',
        'W szkole dostosowania zapisujemy przedmiotowo, a nie ogólnie. Zapis „wydłużenie czasu pracy” w programie '
        'ucznia klasy szóstej nie mówi nic nauczycielowi geografii. Zapis „na sprawdzianach z geografii: '
        'polecenia czytane na głos, mapa konturowa z pogrubionymi granicami, czas wydłużony o połowę, ocena '
        'za treść bez uwzględniania błędów zapisu” mówi wszystko. Dlatego wprowadzamy kartę dostosowań przedmiotowych, '
        'jedną stronę na przedmiot, którą nauczyciel dostaje we wrześniu.',
        'Pięć rodzajów dostosowań, które warto znać. Zmieniamy sposób podania treści: polecenie krokowe, wsparcie '
        'wizualne, tekst w wersji uproszczonej. Zmieniamy czas: dłuższa chwila na odpowiedź, wydłużony czas pracy '
        'na sprawdzianie. Zmieniamy przestrzeń: miejsce w pierwszej ławce, strefa wyciszenia, dostęp do korytarza '
        'przed dzwonkiem. Zmieniamy sposób sprawdzania wiedzy: odpowiedź ustna zamiast pisemnej, test wyboru zamiast '
        'wypracowania, praca na komputerze. I zmieniamy pomoce: nakładki na przybory, liniatura powiększona, '
        'kalkulator, słuchawki wyciszające, piktogramy. Każde dostosowanie ma źródło w barierze, którą opisaliśmy '
        'w ocenie.',
        'Kolejny wymagany element to zintegrowane działania nauczycieli i specjalistów. Zintegrowane znaczy: jeden cel, '
        'jeden plan, wiele rąk. W szkole wygląda to tak. Logopeda, psycholog, terapeuta, nauczyciel współorganizujący '
        'i nauczyciele przedmiotów pracują nad tymi samymi celami z programu, każdy w swoim czasie i swoimi metodami. '
        'Strategia z gabinetu przechodzi na lekcję. Jeśli psycholog uczy ucznia karty „proszę o przerwę”, '
        'to nauczyciel matematyki honoruje tę kartę na lekcji, a rodzic w domu. Zespół spotyka się w ustalonym rytmie, '
        'prowadzi wspólny zeszyt komunikacji i jedną kartę ewaluacji. Żaden specjalista nie ma osobnych celów '
        'w oderwaniu od programu.',
        'Dwa słowa o klasie w duchu uniwersalnego projektowania. Strefy w sali, materiały dostępne i opisane, '
        'plan lekcji w formie wizualnej, miejsce dobrane do potrzeb sensorycznych, regulowany hałas i światło, '
        'jasna struktura lekcji z zapowiedzią zmian. Taka sala jest dostosowaniem dla ucznia z orzeczeniem '
        'i jednocześnie dobrym środowiskiem dla całej klasy. To jest uniwersalne projektowanie, a przy okazji '
        'realizacja ustawy o zapewnianiu dostępności.',
        'Teraz uczeń bez orzeczenia, bo to najliczniejsza grupa w każdej szkole.',
        'Dla ucznia z opinią poradni albo rozpoznanego przez nauczycieli nie sporządzamy programu. Sporządzamy plan '
        'wsparcia edukacyjnego. Plan ma tę samą logikę: rozpoznana potrzeba, cel z kryterium, forma pomocy, osoba '
        'prowadząca, wymiar godzin, okres i termin oceny efektywności.',
        '[[I tu muszę powiedzieć rzecz, o której łatwo zapomnieć, a która jest istotą pracy Strażnika Prawa. '
        'Rozporządzenie o pomocy psychologiczno-pedagogicznej nie zna dokumentu o nazwie plan wsparcia edukacyjnego. '
        'Nie ma w nim odpowiednika programu dla ucznia bez orzeczenia. Przepis mówi tylko tyle: formy pomocy, okres '
        'ich udzielania i wymiar godzin ustala dyrektor, a nauczyciele i specjaliści oceniają efektywność udzielanej '
        'pomocy i dokumentują zajęcia w dziennikach. Plan wsparcia jest zatem naszym narzędziem wewnętrznym, '
        'wprowadzanym zarządzeniem dyrektora — dokładnie na tej samej zasadzie co metryczka. Nikt nie może wymagać '
        'go od nas w kontroli i my też nie powołujemy się przy nim na rozporządzenie. Robimy go dlatego, że bez '
        'jednego miejsca, w którym zapisano cel z kryterium, ocena efektywności nie ma do czego się odnieść. '
        'To jest uczciwe uzasadnienie i tylko takiego wolno nam używać.]]',
        'Program i plan realizujemy przez cele mierzalne. Zatrzymajmy się na chwilę, bo to pojęcie budzi pytania.',
        'Czym są cele SMART? To sposób formułowania celu, w którym cel jest konkretny, mierzalny, osiągalny, istotny '
        'i określony w czasie. Nazwa to skrót od pierwszych liter tych pięciu cech w języku angielskim. Metoda pochodzi '
        'z zarządzania i została przyjęta w edukacji, ponieważ dobrze odpowiada na potrzebę planowania pracy z uczniem '
        'i sprawdzania jej efektów.',
        'Czy przepis wymaga celów SMART? Nie. Rozporządzenie określa zawartość programu i nazwa SMART nie pada w nim '
        'ani razu. Nie ma obowiązku używania tego skrótu w dokumencie.',
        'Co natomiast jest wymagane? Ocena efektywności. Zespół dokonuje wielospecjalistycznej oceny co najmniej dwa '
        'razy w roku, a nauczyciele i specjaliści oceniają efektywność udzielanej pomocy. Ocenić efektywność można '
        'tylko wtedy, gdy cel ma kryterium, do którego da się porównać wynik. Innymi słowy: nazwa jest dowolna, '
        'mierzalność jest konieczna.',
        'Pięć liter. S, konkretny: jakie zachowanie i w jakiej sytuacji. M, mierzalny: ile razy z ilu prób i przy jakim '
        'wsparciu. A, osiągalny: jeden krok od tego, co uczeń robi dziś. R, istotny: wynika z oceny i zwiększa '
        'uczestnictwo ucznia. T, określony w czasie: do kiedy i kiedy sprawdzamy.',
        'Formuła celu jest jedna. Uczeń, w konkretnej sytuacji, będzie wykonywać obserwowalne zachowanie, w określonej '
        'liczbie prób, przy określonym wsparciu, do określonej daty, z określonym sposobem pomiaru.',
        'Posłuchajmy przykładu. Zapis wyjściowy: doskonalenie techniki pisania. Cel mierzalny: Zofia na lekcjach języka '
        'polskiego przepisze z tablicy zdanie o długości do ośmiu wyrazów bez opuszczeń liter w czterech z pięciu '
        'kolejnych prób, przy zadaniu podzielonym na trzy części oznaczone piktogramami i czasie wydłużonym o połowę, '
        'do osiemnastego grudnia dwa tysiące dwudziestego szóstego roku. Pomiar: karta obserwacji przepisywania, '
        'raz w tygodniu.',
        'I drugi przykład, z obszaru zachowania. Zapis wyjściowy: rozwijanie umiejętności radzenia sobie z emocjami. '
        'Cel mierzalny: Zofia w sytuacji zadania trudnego użyje karty „proszę o przerwę” zamiast położenia głowy '
        'na ławce, w trzech z pięciu takich sytuacji w tygodniu, przy jednokrotnym przypomnieniu wizualnym, '
        'do trzydziestego pierwszego marca dwa tysiące dwudziestego siódmego roku. Pomiar: dzienniczek zachowań '
        'zastępczych prowadzony przez nauczyciela współorganizującego.',
        'Zwróćmy uwagę na jedną rzecz w obu przykładach. Kryterium nie brzmi „w stu procentach”. Cel osiągalny '
        'to jeden krok od tego, co uczeń robi dziś, a nie stan docelowy za trzy lata.',
        'Ewaluacja. Ile razy w roku?',
        'Ocena wielospecjalistyczna co najmniej dwa razy: we wrześniu i w styczniu, a zalecana trzecia w maju. '
        'Modyfikacja programu po każdej ocenie. Ocena efektywności pomocy psychologiczno-pedagogicznej dla uczniów '
        'bez orzeczenia — na bieżąco i na zakończenie form, w praktyce w styczniu i w czerwcu. Kwestionariusz '
        'we wrześniu i w maju.',
        'Rekomendujemy trzy oceny w roku oraz dwa krótkie przeglądy wskaźników, w listopadzie i w marcu. '
        'Piętnaście minut na ucznia i notatka w dzienniku. Dzięki temu o nieskuteczności działania dowiadujemy się '
        'po dwóch miesiącach, a nie po dziesięciu.',
        'Po każdym pomiarze zespół podejmuje jedną z czterech decyzji. Cel osiągnięty: zamykamy i stawiamy kolejny. '
        'Osiągnięty częściowo: kontynuujemy i przesuwamy termin, nie obniżając kryterium. Brak postępu: modyfikujemy '
        'metodę i sprawdzamy bariery środowiskowe. Regres: spotkanie z rodzicami i rozważenie wystąpienia do poradni.',
        'Dwie sprawy szkolne, o których łatwo zapomnieć przy ewaluacji. Listopad klasy ósmej to termin, w którym '
        'porządkujemy dokumentację pod kątem dostosowań na egzaminie ósmoklasisty. Czerwiec klasy trzeciej to termin '
        'karty przekazania do drugiego etapu. Oba dokumenty składamy z danych, które już mamy.',
        'Na koniec opinia o funkcjonowaniu ucznia dla poradni psychologiczno-pedagogicznej.',
        'Wydajemy ją na prośbę przewodniczącego zespołu orzekającego, w terminie dziesięciu dni od dnia otrzymania '
        'prośby przez dyrektora.',
        'Ma siedem punktów: dane formalne, mocne strony i uzdolnienia, funkcjonowanie w obszarach, trudności '
        'z częstotliwością i kontekstem, bariery i ułatwienia, udzielone wsparcie i jego efekty, współpraca z rodzicami.',
        'Piszemy językiem funkcjonalnym i sprawdzalnym. Opisujemy zachowania, ich częstotliwość i kontekst. '
        'Rozpoznania i hipotezy diagnostyczne pozostawiamy poradni.',
        'Informacje pochodzące od rodziców lub specjalistów spoza szkoły oznaczamy jako relację ze wskazaniem źródła. '
        'Kwestionariusz pozostaje naszym materiałem roboczym.',
        'Jedna dobra praktyka na zakończenie. Punkt o mocnych stronach piszemy jako pierwszy i co najmniej tak samo '
        'obszernie jak punkt o trudnościach. Ta proporcja jest najprostszym sprawdzianem, czy nasz opis ucznia jest '
        'naprawdę funkcjonalny.',
        'Podsumowując całe szkolenie. Każdy druk ma swój przepis. Obserwacja wyprzedza pismo z poradni. Cel ma liczbę, '
        'a ewaluacja ma konsekwencję. I jedno zdanie z modułu drugiego, które warto powtórzyć na koniec: '
        'nie zmieniamy dokumentacji dlatego, że ktoś nam kazał, tylko dlatego, że stara dokumentacja przestała '
        'odpowiadać na pytania, które nam teraz zadają. Dziękuję Państwu za uwagę i życzę spokojnego roku szkolnego.',
    ])

    h2(doc, 'Osiem obowiązkowych elementów IPET — karta kontrolna rozporządzenia')
    table(doc,
          ['#', 'Element (§ 6 ust. 1 rozp. t.j. Dz.U. 2020 poz. 1309)', 'Gdzie w naszym druku', 'Jest?'],
          [['1', 'Zakres i sposób dostosowania wymagań edukacyjnych do potrzeb i możliwości ucznia — '
            'w szczególności metody i formy pracy.', 'IPET sekcja 5 + karta dostosowań przedmiotowych (druk 11)', '☐'],
           ['2', 'Zintegrowane działania nauczycieli i specjalistów (przy niepełnosprawności — ukierunkowane '
            'na poprawę funkcjonowania, w tym komunikowanie się).', 'IPET sekcja 4 — pole „zintegrowane działania” w każdej sferze', '☐'],
           ['3', 'Formy i okres udzielania pomocy psychologiczno-pedagogicznej oraz wymiar godzin.',
            'IPET sekcja 6 — tabela form, osób i wymiaru', '☐'],
           ['4', 'Działania wspierające rodziców oraz zakres współdziałania z poradniami, placówkami doskonalenia '
            'nauczycieli, organizacjami i instytucjami.', 'IPET sekcja 9', '☐'],
           ['5', 'Zajęcia rewalidacyjne, resocjalizacyjne i socjoterapeutyczne oraz inne odpowiednie ze względu '
            'na potrzeby — w tym doradztwo zawodowe i zajęcia związane z wyborem kierunku kształcenia.',
            'IPET sekcja 6A i 6B', '☐'],
           ['6', 'Zakres współpracy nauczycieli i specjalistów z rodzicami ucznia.', 'IPET sekcja 9 + rejestr kontaktów', '☐'],
           ['7', 'Rodzaj i sposób dostosowania warunków organizacji kształcenia (w zależności od potrzeb).',
            'IPET sekcja 7 — organizacja, nauczyciel współorganizujący, miejsce w klasie', '☐'],
           ['8', 'W przypadku uczniów wymagających komunikacji wspomagającej — wskazanie zajęć realizowanych '
            'indywidualnie lub w grupie do 5 osób.', 'IPET sekcja 7 — pole „forma realizacji zajęć”', '☐']],
          widths=[0.9, 7.6, 6.4, 1.1], size=8)
    para(doc, 'Kartę kontrolną wypełnia koordynator przed podpisaniem programu. Brak choćby jednego elementu '
              'oznacza dokument niepełny w rozumieniu rozporządzenia. Treść przepisu sprawdzamy w obowiązującym '
              'tekście jednolitym. ⚑', size=8, italic=True, color=GREY, after=8)

    h2(doc, 'Dwie ścieżki — czym różni się IPET od planu wsparcia (PWES)')
    table(doc,
          ['Kryterium', 'IPET — kształcenie specjalne', 'PWES — pomoc psychologiczno-pedagogiczna'],
          [['Dla kogo', 'Uczeń z orzeczeniem o potrzebie kształcenia specjalnego.',
            'Uczeń z opinią poradni albo rozpoznany przez nauczycieli, bez orzeczenia.'],
           ['Podstawa prawna', 'Art. 127 Prawa oświatowego (t.j. Dz.U. 2026 poz. 820); § 6 rozp. t.j. Dz.U. 2020 '
            'poz. 1309. Druk IPET jest wymagany przez rozporządzenie.',
            'Rozp. o pomocy pp, t.j. Dz.U. 2023 poz. 1798 — [[reguluje formy, okres i wymiar pomocy oraz ocenę '
            'efektywności, ale NIE przewiduje druku planu. PWES to narzędzie wewnętrzne szkoły.]]'],
           ['Poprzedza go', 'Wielospecjalistyczna ocena poziomu funkcjonowania (obowiązkowo).',
            'Rozpoznanie potrzeb przez nauczycieli; KSzOF jako narzędzie rozpoznania.'],
           ['Termin', 'Do 30 września albo 30 dni od złożenia orzeczenia w szkole.',
            'Niezwłocznie po ustaleniu form pomocy przez dyrektora.'],
           ['Kto opracowuje', 'Zespół nauczycieli i specjalistów; koordynator wyznaczony przez dyrektora.',
            'Nauczyciele i specjaliści prowadzący zajęcia z uczniem; [[formy, okres i wymiar pomocy ustala '
            'dyrektor]].'],
           ['Obowiązkowa zawartość', 'Osiem elementów z § 6 ust. 1.',
            '[[Zakres ustala szkoła. Z przepisu wynikają: forma pomocy, okres i wymiar godzin (ustala dyrektor) '
            'oraz ocena efektywności.]]'],
           ['Ewaluacja', 'WOPFU co najmniej 2× w roku; modyfikacja programu po każdej ocenie.',
            'Ocena efektywności na bieżąco i na zakończenie formy pomocy.'],
           ['Prawa rodziców', 'Udział w spotkaniach zespołu; kopia oceny i kopia programu.',
            'Informacja o formach, wymiarze i okresie pomocy.'],
           ['Zajęcia rewalidacyjne', 'Tak — § 5 rozporządzenia.', 'Nie. Rewalidacja przysługuje tylko z orzeczeniem.'],
           ['Nauczyciel współorganizujący', 'Obowiązkowo [[wyłącznie przy autyzmie (w tym zespole Aspergera) '
            'i niepełnosprawnościach sprzężonych — § 7 ust. 2. W pozostałych przypadkach, także przy '
            'niedostosowaniu społecznym, tylko za zgodą organu prowadzącego — § 7 ust. 3.]]', 'Nie dotyczy.']],
          widths=[3.2, 6.4, 6.4], size=8)

    h2(doc, 'Cel wyjściowy a cel mierzalny — trzy pary do porównania')
    table(doc,
          ['Zapis wyjściowy (nieewaluowalny)', 'Cel mierzalny', 'Sposób pomiaru'],
          [['Doskonalenie techniki pisania.',
            'Zofia na lekcjach języka polskiego przepisze z tablicy zdanie do 8 wyrazów bez opuszczeń liter '
            'w 4 z 5 kolejnych prób, przy zadaniu podzielonym na 3 części oznaczone piktogramami i czasie '
            'wydłużonym o 50%, do 18.12.2026 r.', 'Karta obserwacji przepisywania — raz w tygodniu.'],
           ['Rozwijanie umiejętności radzenia sobie z emocjami.',
            'Zofia w sytuacji zadania trudnego użyje karty „proszę o przerwę” zamiast położenia głowy na ławce, '
            'w 3 z 5 takich sytuacji w tygodniu, przy jednokrotnym przypomnieniu wizualnym, do 31.03.2027 r.',
            'Dzienniczek zachowań zastępczych — nauczyciel współorganizujący.'],
           ['Poprawa funkcjonowania w grupie rówieśniczej.',
            'Zofia podczas pracy w parze na lekcji przyjmie przydzieloną rolę i wykona swoją część zadania '
            'przez co najmniej 10 minut, w 3 z 5 lekcji w tygodniu, przy wprowadzeniu do zadania przez nauczyciela, '
            'do końca kwietnia 2027 r.', 'Karta obserwacji pracy w parze — raz w tygodniu.']],
          widths=[4.5, 8.0, 3.5], size=8)

    h2(doc, 'Jak przygotować wielospecjalistyczną ocenę (WOPF-SP)')
    steps(doc, [
        '**Zwołaj zespół:** wychowawca, nauczyciele przedmiotów, psycholog, pedagog specjalny, logopeda, terapeuta, '
        'nauczyciel współorganizujący. Dyrektor zawiadamia rodziców o terminie spotkania.',
        '**Wypełnij sekcję I i I a:** dane ucznia w zakresie minimalnym, rodzaj oceny (pierwsza przed programem, '
        'śródroczna, roczna) i ścieżkę dokumentacyjną: IPET albo PWES.',
        '**Wypełnij sekcję III — mapę dokumentów źródłowych.** To lista druków, z których ta ocena powstaje. '
        'Jeśli któregoś brakuje, wiadomo, czego szukać, zanim zespół usiądzie.',
        '**Przepisz wyniki KSzOF do sekcji V** — dziewięć obszarów, punkty, skala 0–20, sten, poziom wsparcia. '
        'W sekcji V a rozpisz każdy obszar na mocne strony i trudności.',
        '**Przenieś wnioski z obserwacji pogłębionej:** ABC i FBA do sekcji VI, poznanie społeczne do VII, '
        'mowa i komunikacja do VIII, przetwarzanie sensoryczne do IX.',
        '**Opisz trudności językiem funkcjonalnym,** z częstotliwością i kontekstem — nie w formie etykiety.',
        '**Zaznacz czynniki środowiskowe i opisz bariery oraz ułatwienia.** To one uzasadniają dostosowania w programie. '
        'Bariera bez opisu = dostosowanie bez uzasadnienia.',
        '**Zamknij ocenę wnioskami.** Każde zalecenie przechodzi do programu jako wiersz „zalecenie z WOPF → realizacja”.',
        '**Przekaż rodzicom kopię oceny** i odnotuj to w rejestrze kontaktów — z datą i podpisem.',
    ])

    h2(doc, 'Jak przygotować IPET')
    steps(doc, [
        '**Sprawdź termin:** do 30 września albo 30 dni od dnia złożenia orzeczenia w szkole. Datę wpływu odczytaj '
        'z metryczki.',
        '**Przepisz zalecenia z orzeczenia poradni,** jedno po drugim, a przy każdym zapisz sposób realizacji: '
        'forma, kto prowadzi, wymiar godzin, od kiedy.',
        '**Zrób to samo z zaleceniami z oceny wielospecjalistycznej.** Orzeczenie mówi, co uczniowi zalecono; '
        'ocena mówi, co widzimy w szkole.',
        '**Opisz zakres i sposób dostosowania przedmiotowo** — osobno dla każdego przedmiotu: sposób podania treści, '
        'czas, przestrzeń, sposób sprawdzania wiedzy, pomoce. Każde dostosowanie ma źródło w barierze opisanej w ocenie.',
        '**Zapisz zintegrowane działania nauczycieli i specjalistów:** jeden cel, jeden plan, wiele rąk — '
        'z rolą rodziców w domu i wskazaniem, kto co honoruje na lekcji.',
        '**Sformułuj cele w postaci mierzalnej:** uczeń, w konkretnej sytuacji, wykona obserwowalne zachowanie, '
        'w określonej liczbie prób, przy określonym wsparciu, do określonej daty, z określonym sposobem pomiaru.',
        '**Wypełnij kartę kontrolną ośmiu elementów** przed podpisaniem programu.',
        '**Zaplanuj ewaluację i wpisz ją do kalendarza.** Rodzice otrzymują kopię programu — potwierdzenie '
        'w rejestrze kontaktów.',
    ])

    h2(doc, 'Jak przygotować opinię o funkcjonowaniu ucznia dla poradni')
    steps(doc, [
        '**Zanotuj datę otrzymania prośby przez dyrektora** — od niej liczy się dziesięć dni.',
        '**Zbierz dane, które już masz:** metryczka, KSzOF, obserwacja pogłębiona, karty ewaluacji, rejestr kontaktów '
        'z rodzicami, karty dostosowań przedmiotowych.',
        '**Napisz siedem punktów:** dane formalne, mocne strony i uzdolnienia, funkcjonowanie w obszarach, trudności '
        'z częstotliwością i kontekstem, bariery i ułatwienia, udzielone wsparcie i jego efekty, współpraca z rodzicami.',
        '**Punkt o mocnych stronach napisz jako pierwszy** i co najmniej tak samo obszernie jak punkt o trudnościach.',
        '**Używaj języka funkcjonalnego:** co uczeń robi, w jakich warunkach, przy jakim wsparciu, jak często. '
        'Rozpoznania i hipotezy diagnostyczne zostaw poradni.',
        '**Zbierz informacje od nauczycieli przedmiotów** — w II etapie edukacyjnym opinia napisana wyłącznie przez '
        'wychowawcę jest opinią niepełną.',
        '**Informacje od rodziców i specjalistów spoza szkoły oznacz jako relację** ze wskazaniem źródła.',
        '**Przekaż kopię opinii rodzicom** i odnotuj to w dokumentacji.',
    ])


# ================================================================ ZAŁĄCZNIKI A i B

def zal_A_B(doc):
    page_break(doc)
    band(doc, 'ZAŁĄCZNIK A', 'Kalendarz dokumentacji na rok szkolny — szkoła podstawowa')
    table(doc,
          ['Miesiąc', 'Co robimy', 'Dokument', 'Odpowiedzialny'],
          [['sierpień / wrzesień', 'Zakładamy teczkę ucznia, uzupełniamy metryczkę, zbieramy zgody rodziców, '
            'odnotowujemy datę wpływu orzeczenia.', 'Metryczka ucznia', 'wychowawca'],
           ['wrzesień', 'KSzOF dla wszystkich uczniów w oddziale — pomiar bazowy; arkusz rodzica przekazany '
            'na pierwszym zebraniu.', 'KSzOF (I–III / IV–VI / VII–VIII)', 'wychowawca + rodzic'],
           ['do 30 września', 'IPET dla ucznia rozpoczynającego kształcenie z orzeczeniem; pierwsza '
            'wielospecjalistyczna ocena.', 'WOPF-SP + IPET', 'koordynator zespołu'],
           ['wrzesień / październik', 'Obserwacja pogłębiona tam, gdzie zadziałała reguła przekierowania; '
            'karty dostosowań przedmiotowych przekazane nauczycielom.',
            'Karta decyzyjna + wybrane narzędzie + druk 11', 'zespół'],
           ['październik', 'Zebranie informacji od nauczycieli przedmiotów o uczniach objętych wsparciem (II etap).',
            'Arkusz informacji nauczyciela przedmiotu (druk 10)', 'wychowawca'],
           ['listopad', 'Krótki przegląd wskaźników — 15 minut na ucznia. W klasach VIII: uporządkowanie '
            'dokumentacji pod dostosowania egzaminacyjne. ⚑ terminy z komunikatu dyrektora CKE na dany rok',
            'Karta celu SMART + karta dostosowań egzaminacyjnych', 'koordynator / dyrektor'],
           ['styczeń', 'Druga wielospecjalistyczna ocena; ewaluacja półroczna celów; ocena efektywności pomocy pp.',
            'WOPF-SP + IPET + karta ewaluacji + karta oceny efektywności', 'zespół'],
           ['marzec', 'Krótki przegląd wskaźników.', 'Karta celu SMART', 'koordynator'],
           ['kwiecień / maj', 'Egzamin ósmoklasisty — realizacja przyznanych dostosowań zgodnie z dokumentacją.',
            'Karta dostosowań egzaminacyjnych', 'dyrektor'],
           ['maj', 'KSzOF po raz drugi — pomiar kontrolny na tym samym arkuszu, innym kolorem; trzecia, zalecana '
            'ocena wielospecjalistyczna.', 'KSzOF + WOPF-SP', 'wychowawca + zespół'],
           ['czerwiec', 'Ocena efektywności udzielanej pomocy na zakończenie form. W klasach III: karta przekazania '
            'informacji o uczniu do II etapu.', 'Karta ewaluacji + karta przekazania (druk 16)', 'zespół'],
           ['cały rok', 'Opinia dla poradni na prośbę przewodniczącego zespołu orzekającego — w terminie 10 dni '
            'od wpływu prośby do dyrektora.', 'Opinia o funkcjonowaniu ucznia', 'dyrektor + zespół']],
          widths=[2.8, 6.2, 4.6, 3.4], size=8)

    h2(doc, 'ZAŁĄCZNIK B — Język funkcjonalny: cztery pytania do każdego zdania')
    accent_box(doc, [
        'CO UCZEŃ ROBI? — zachowanie, które można zobaczyć i policzyć, a nie cecha ani rozpoznanie.',
        'W JAKICH WARUNKACH? — przedmiot, pora dnia, wielkość grupy, poziom hałasu, rodzaj zadania, miejsce w sali.',
        'PRZY JAKIM WSPARCIU? — samodzielnie, po podpowiedzi słownej, po pokazie, z pomocą dorosłego, '
        'z pomocą „ręka na rękę”.',
        'JAK CZĘSTO? — liczba prób, dni w tygodniu, epizody w obserwowanym okresie, czas trwania.',
    ], title='CZTERY PYTANIA')
    table(doc,
          ['Zapis, który nie działa', 'Ten sam uczeń, zapis funkcjonalny'],
          [['Ma trudności w koncentracji uwagi wynikające z zaburzeń rozwojowych; wymaga stałej pomocy nauczyciela.',
            'Pracuje samodzielnie przy zadaniu na karcie pracy ok. 8 minut, jeśli siedzi w pierwszej ławce '
            'i zadanie jest podzielone na 3 kroki oznaczone piktogramami. Przy pracy z tablicą utrzymuje uwagę '
            'ok. 2 minut i wymaga podpowiedzi słownej średnio 4 razy na lekcji. W klasie o podwyższonym hałasie '
            'czas skraca się do 1 minuty.'],
           ['Jest niesamodzielna i wymaga ciągłego nadzoru.',
            'Przygotowuje przybory na lekcję po jednej podpowiedzi słownej w 4 z 5 dni; pakuje plecak po lekcjach '
            'z listą obrazkową samodzielnie w 3 z 5 dni.'],
           ['Zachowuje się agresywnie wobec rówieśników.',
            'W tygodniu obserwacji trzykrotnie odepchnęła kolegę — za każdym razem w kolejce do szatni, '
            'przy zbliżeniu na odległość mniejszą niż pół metra. Po wprowadzeniu wyjścia 2 minuty przed dzwonkiem '
            'zdarzenia nie wystąpiły przez 2 tygodnie.'],
           ['Nie radzi sobie z matematyką.',
            'Wykonuje dodawanie w zakresie 20 na konkretach samodzielnie; na zapisie symbolicznym wymaga '
            'przypomnienia algorytmu przy każdym zadaniu. Zadania tekstowe rozwiązuje po przeczytaniu treści '
            'na głos i podkreśleniu danych — wtedy 3 z 5 zadań poprawnie.'],
           ['Odmawia pracy na lekcji.',
            'Po poleceniu przepisania z tablicy tekstu dłuższego niż 4 linijki odkłada ołówek i kładzie głowę '
            'na ławce średnio 2 razy dziennie, najczęściej na 3. i 5. lekcji. Po podzieleniu zadania na 3 części '
            'podejmuje pracę w 4 z 5 prób.']],
          widths=[6.4, 10.3], size=8)


# ================================================================ ZAŁĄCZNIK C — DRUKI

def _meta(doc, kto, kiedy, podstawa, zasila):
    table(doc, ['Kto wypełnia', 'Kiedy', 'Podstawa / status druku', 'Zasila'],
          [[kto, kiedy, podstawa, zasila]],
          widths=[3.4, 3.4, 5.6, 4.3], size=7.5, align_first_bold=False)

def _naglowek_ucznia(doc, dodatkowe=''):
    t = table(doc, ['Uczeń — imię i nazwisko', 'Klasa / oddział', 'Okres obserwacji', 'Wypełniający i rola'],
              [['', '', 'od ......... do .........', '']],
              widths=[5.4, 2.6, 4.2, 4.5], size=8, align_first_bold=False)
    if dodatkowe:
        para(doc, dodatkowe, size=7.5, italic=True, color=GREY, after=6)
    return t

def zal_C_1(doc):
    page_break(doc)
    band(doc, 'ZAŁĄCZNIK C', 'Osiemnaście druków do powielenia — kwestionariusze, karty i arkusze')
    para(doc, 'Druki 1–18 są gotowe do powielenia. Każdy zawiera nagłówek z metryczką obserwacji, wskazanie, '
              'kto i kiedy go wypełnia, oraz informację, który dokument zasila. Druki nie powielają danych osobowych: '
              'dziedziczą je z metryczki ucznia. Numeracja druków jest ta sama, co w częściach 5–7 skryptu.',
         after=8)
    table(doc,
          ['Nr', 'Druk', 'Sygnatura', 'Uruchamiamy, gdy…'],
          [['1', 'Karta decyzyjna obserwacji pogłębionej', 'KD-SP', 'zadziałała którakolwiek z siedmiu reguł przekierowania'],
           ['2', 'Karta obserwacji zachowania — model ABC', 'ABC-SP', 'zachowanie trudne jest powtarzalne lub zagrażające'],
           ['3', 'Arkusz analizy funkcjonalnej zachowania', 'FBA-SP', 'zebrano 10–15 zapisów ABC i zespół formułuje hipotezę'],
           ['4', 'Profil sensoryczny ucznia', 'PS-SP', 'sygnał z metryczki albo niskie wyniki w obszarach IV i V'],
           ['5', 'Karta obserwacji rozwoju mowy i komunikacji', 'KM-SP', 'niski wynik w obszarze III albo trudność w rozumieniu poleceń'],
           ['6', 'Arkusz obserwacji poznania społecznego i teorii umysłu', 'ToM-SP', 'niski wynik w obszarze VII przy zachowanych I i IV'],
           ['7', 'Karta obserwacji funkcji uczenia się', 'KFU-SP', 'niski wynik w obszarze I lub VIII'],
           ['8', 'Ankieta „Mój głos” — perspektywa ucznia', 'MG-SP', 'zawsze przed zespołem — dla każdego ucznia objętego wsparciem'],
           ['9', 'Ankieta dla rodzica i wywiad środowiskowy', 'AR-SP', 'wrzesień — równolegle z arkuszem wychowawcy'],
           ['10', 'Arkusz informacji od nauczyciela przedmiotu', 'NP-SP', 'II etap edukacyjny — przed każdą oceną wielospecjalistyczną'],
           ['11', 'Karta dostosowań przedmiotowych', 'KDP-SP', 'po opracowaniu IPET lub PWES — jedna karta na przedmiot'],
           ['12', 'Karta dostosowań warunków egzaminu ósmoklasisty', 'AD-E8', 'listopad klasy VIII ⚑ terminy z komunikatu CKE'],
           ['13', 'Karta celu i ewaluacji', 'SMART-SP', 'przy każdym celu z IPET lub PWES'],
           ['14', 'Karta oceny efektywności udzielanej pomocy', 'OE-SP', 'styczeń i czerwiec oraz na zakończenie formy pomocy'],
           ['15', 'Plan wsparcia edukacyjnego ucznia [[(narzędzie wewnętrzne)]]', 'PWES',
            'uczeń bez orzeczenia objęty pomocą pp'],
           ['16', 'Karta przekazania informacji o uczniu (III → IV)', 'KP-SP', 'czerwiec klasy III'],
           ['17', 'Arkusz audytu dokumentacji kształcenia specjalnego', 'AUD-SP', 'sierpień / wrzesień — raz w roku, przed radą pedagogiczną'],
           ['18', 'Rejestr kontaktów z rodzicami', 'RK-SP', 'prowadzony na bieżąco przez cały rok']],
          widths=[1.0, 6.4, 2.0, 7.3], size=8)

    # ---------------------------------------------------------------- DRUK 1
    page_break(doc)
    form_header(doc, 'DRUK 1 · KD-SP', 'Karta decyzyjna obserwacji pogłębionej',
                'Jedna karta na jednego ucznia. Wpinamy do teczki TAKŻE wtedy, gdy decyzja brzmi „nie uruchamiamy”.')
    _meta(doc, 'koordynator zespołu', 'po analizie wyników KSzOF',
          'decyzja rady pedagogicznej wpisana do procedury szkoły; dokumentacja badań i czynności uzupełniających ⚑',
          'sekcje VI–IX oceny WOPF-SP')
    _naglowek_ucznia(doc)
    para(doc, 'I. KTÓRA REGUŁA PRZEKIEROWANIA ZADZIAŁAŁA — zaznacz wszystkie spełnione',
         size=9, bold=True, color=PURPLE, before=8, after=4)
    checkboxes(doc, [
        'R1 — wynik obszaru ≤ 8 pkt w skali 0–20 (Poziom III). Obszar: ...............',
        'R2 — dwa lub więcej twierdzeń ocenionych na 1 lub 2 w tym samym obszarze',
        'R3 — rozbieżność między oceniającymi ≥ 2 steny w wyniku ogólnym',
        'R4 — sygnał zdrowotny z metryczki: ...............',
        'R5 — zachowanie powtarzalne zagrażające uczniowi lub innym (uruchamiamy natychmiast)',
        'R6 — brak poprawy mimo pomocy przez ok. 3 miesiące',
        'R7 — nagła zmiana funkcjonowania: oceny w ≥ 3 przedmiotach / nieobecności > 20% w miesiącu',
        'żadna reguła nie została spełniona — obserwacji pogłębionej NIE uruchamiamy',
    ], cols=1)
    para(doc, 'II. WYBRANE NARZĘDZIE', size=9, bold=True, color=PURPLE, before=6, after=4)
    checkboxes(doc, ['ABC + FBA (druki 2 i 3)', 'Profil sensoryczny (druk 4)',
                     'Karta mowy i komunikacji (druk 5)', 'Poznanie społeczne / ToM (druk 6)',
                     'Karta funkcji uczenia się (druk 7)', 'Inne: ...............'], cols=2)
    table(doc, ['Kto obserwuje', 'Od kiedy — do kiedy', 'Sytuacje objęte obserwacją', 'Termin spotkania zespołu'],
          [['', '', '', ''], ['', '', '', '']],
          widths=[3.7, 3.4, 6.0, 3.6], size=8, align_first_bold=False)
    form_lines(doc, 'UZASADNIENIE DECYZJI (także decyzji odmownej — dlaczego zespół uznał, że obserwacja nie jest potrzebna)', 3)
    table(doc, ['Data', 'Koordynator — podpis', 'Wychowawca — podpis', 'Rodzic poinformowany (data, podpis)'],
          [['', '', '', '']], widths=[2.6, 4.4, 4.4, 5.3], size=8, align_first_bold=False)

    # ---------------------------------------------------------------- DRUK 2
    page_break(doc)
    form_header(doc, 'DRUK 2 · ABC-SP', 'Karta obserwacji zachowania — model ABC',
                'Zbieramy 10–15 zapisów w ciągu 2–3 tygodni, z różnych lekcji i różnych pór dnia. '
                'Zapisujemy fakty, nie interpretacje. Interpretację formułuje zespół w druku 3.')
    _meta(doc, 'nauczyciele uczący ucznia, nauczyciel współorganizujący',
          'bezpośrednio po zdarzeniu — nie z pamięci wieczorem',
          'rozpoznawanie potrzeb (t.j. Dz.U. 2023 poz. 1798); dokumentacja czynności uzupełniających ⚑',
          'druk 3 (FBA) → sekcja VI oceny WOPF-SP')
    _naglowek_ucznia(doc)
    accent_box(doc, [
        'A — POPRZEDNIK: co działo się bezpośrednio przed zachowaniem (polecenie, zmiana, hałas, zadanie, osoba obok).',
        'B — ZACHOWANIE: co uczeń zrobił — opis obserwowalny i mierzalny, z czasem trwania i liczbą.',
        'C — NASTĘPSTWO: co stało się bezpośrednio potem, w tym reakcja dorosłych i rówieśników.',
        'ZAKAZ INTERPRETACJI: nie piszemy „zdenerwował się”, „nie chciało mu się”, „manifestował niechęć”. '
        'Piszemy to, co da się nagrać kamerą.',
    ], title='JAK ZAPISYWAĆ')
    table(doc,
          ['Data / godz.', 'Lekcja / miejsce', 'A — poprzednik', 'B — zachowanie (czas, liczba)',
           'C — następstwo', 'Obserwator'],
          [['', '', '', '', '', ''] for _ in range(8)],
          widths=[2.0, 2.2, 3.4, 3.6, 3.4, 2.1], size=8, align_first_bold=False, zebra=True)
    para(doc, 'Kontynuacja na kolejnym egzemplarzu druku. Zapisy numerujemy narastająco w obrębie jednego okresu obserwacji.',
         size=7.5, italic=True, color=GREY)

    # ---------------------------------------------------------------- DRUK 3
    page_break(doc)
    form_header(doc, 'DRUK 3 · FBA-SP', 'Arkusz analizy funkcjonalnej zachowania',
                'Wypełnia ZESPÓŁ po zebraniu 10–15 zapisów ABC. Bez zaplanowanego zachowania zastępczego plan nie działa.')
    _meta(doc, 'zespół: psycholog / pedagog specjalny + wychowawca', 'po zamknięciu okresu obserwacji ABC',
          'rozpoznawanie potrzeb i ocena efektywności (t.j. Dz.U. 2023 poz. 1798)',
          'sekcja VI oceny WOPF-SP → cele w IPET / PWES')
    _naglowek_ucznia(doc)
    form_lines(doc, 'ZACHOWANIE DOCELOWE — jedno zachowanie, opisane obserwowalnie (co, jak długo, jak często)', 2)
    para(doc, 'HIPOTEZA FUNKCJI — po co uczeń to robi? (zaznacz jedną główną)',
         size=9, bold=True, color=PURPLE, before=6, after=4)
    checkboxes(doc, ['uzyskanie uwagi dorosłego lub rówieśników',
                     'uzyskanie przedmiotu lub aktywności',
                     'uniknięcie / odroczenie trudnego zadania',
                     'ucieczka z sytuacji przeciążającej sensorycznie',
                     'regulacja pobudzenia (samostymulacja)',
                     'komunikat, którego uczeń nie umie wyrazić inaczej'], cols=2)
    table(doc, ['Dane, które potwierdzają hipotezę (numery zapisów ABC)', 'Dane, które jej przeczą', 'Wniosek zespołu'],
          [['', '', '']], widths=[6.0, 5.0, 5.7], size=8, align_first_bold=False)
    table(doc,
          ['Zachowanie zastępcze — czego uczymy', 'Kto uczy i kiedy', 'Jak je honorujemy (kto, jak szybko)',
           'Zmiana w otoczeniu (usunięcie poprzednika)'],
          [['', '', '', ''], ['', '', '', '']],
          widths=[4.4, 3.4, 4.6, 4.3], size=8, align_first_bold=False)
    form_lines(doc, 'PROCEDURA REAGOWANIA — co robimy, gdy zachowanie wystąpi (jednakowa dla wszystkich nauczycieli)', 3)
    table(doc, ['Wskaźnik skuteczności (liczba epizodów / tydzień)', 'Wartość wyjściowa', 'Wartość docelowa',
                'Data pomiaru'],
          [['', '', '', '']], widths=[6.4, 3.2, 3.2, 3.9], size=8, align_first_bold=False)

    # ---------------------------------------------------------------- DRUK 4
    page_break(doc)
    form_header(doc, 'DRUK 4 · PS-SP', 'Profil sensoryczny ucznia — obserwacja szkolna',
                'Opisujemy WZORZEC reakcji, nie stawiamy rozpoznania. Kwalifikacja do terapii należy '
                'do terapeuty integracji sensorycznej.')
    _meta(doc, 'wychowawca + nauczyciel współorganizujący; konsultacja terapeuty SI',
          'sygnał z metryczki albo niskie wyniki w obszarach IV i V KSzOF',
          'obserwacja pedagogiczna; granica kompetencji nauczyciela',
          'sekcja IX oceny WOPF-SP → dostosowania w IPET')
    _naglowek_ucznia(doc)
    para(doc, 'Przy każdym układzie zaznacz obserwowany wzorzec: N = nadreaktywność, P = podreaktywność, '
              'S = poszukiwanie bodźców, — = reakcje typowe. W kolumnie „przykład” wpisz konkretną sytuację szkolną.',
         size=8, italic=True, color=GREY, after=6)
    table(doc,
          ['Układ', 'Na co patrzymy w szkole', 'N', 'P', 'S', '—', 'Przykład obserwowanej sytuacji'],
          [['Słuchowy', 'dzwonek, gwar na korytarzu, hałas na stołówce, praca w grupach', '☐', '☐', '☐', '☐', ''],
           ['Wzrokowy', 'jarzeniówki, tablica multimedialna, natłok bodźców na ścianach klasy', '☐', '☐', '☐', '☐', ''],
           ['Dotykowy', 'kolejka do szatni, przypadkowe dotknięcia, metki, kleje i farby', '☐', '☐', '☐', '☐', ''],
           ['Węchowy', 'stołówka, sala chemiczna, perfumy, środki czystości', '☐', '☐', '☐', '☐', ''],
           ['Smakowy', 'obiad w stołówce, konsystencje, temperatura potraw', '☐', '☐', '☐', '☐', ''],
           ['Przedsionkowy', 'schody, WF, bujanie się na krześle, karuzela na przerwie', '☐', '☐', '☐', '☐', ''],
           ['Proprioceptywny', 'nacisk na ołówek, siła uścisku, wpadanie na przedmioty, pozycja przy ławce',
            '☐', '☐', '☐', '☐', '']],
          widths=[2.5, 5.4, 0.8, 0.8, 0.8, 0.8, 5.6], size=8)
    para(doc, 'TRZY MIEJSCA SZKOLNE O NAJWYŻSZYM NATĘŻENIU BODŹCÓW — opisz funkcjonowanie ucznia w każdym',
         size=9, bold=True, color=PURPLE, before=6, after=4)
    table(doc, ['Korytarz na przerwie', 'Stołówka', 'Sala gimnastyczna / szatnia WF'],
          [['', '', '']], widths=[5.6, 5.5, 5.6], size=8, align_first_bold=False)
    form_lines(doc, 'WNIOSKI DLA ORGANIZACJI PRACY (miejsce w klasie, przerwy, pomoce, strefa wyciszenia)', 3)
    form_lines(doc, 'ZALECANA KONSULTACJA SPECJALISTY — zakres i uzasadnienie', 2)


def zal_C_2(doc):
    # ---------------------------------------------------------------- DRUK 5
    page_break(doc)
    form_header(doc, 'DRUK 5 · KM-SP', 'Karta obserwacji rozwoju mowy i komunikacji ucznia',
                'Skala: 0 = zachowanie nie występuje · 1 = występuje częściowo lub z pomocą dorosłego · '
                '2 = uczeń robi to samodzielnie. Każdy obszar daje 0–10 punktów.')
    _meta(doc, 'wychowawca; konsultacja logopedy', 'niski wynik w obszarze III KSzOF albo trudność w rozumieniu poleceń',
          'obserwacja pedagogiczna; diagnoza logopedyczna należy do logopedy',
          'sekcja VIII oceny WOPF-SP → cele w sferze komunikacji')
    _naglowek_ucznia(doc)
    table(doc,
          ['Obszar', 'Wskaźnik', '0', '1', '2'],
          [['I. Rozumienie mowy', '1. Reaguje na swoje imię i na polecenie skierowane do całej klasy', '☐', '☐', '☐'],
           ['', '2. Wykonuje polecenie jednoetapowe bez wsparcia wizualnego', '☐', '☐', '☐'],
           ['', '3. Wykonuje polecenie dwu- i trzyetapowe', '☐', '☐', '☐'],
           ['', '4. Rozumie treść czytanego tekstu na poziomie dosłownym', '☐', '☐', '☐'],
           ['', '5. Rozumie pytania o przyczynę i cel („dlaczego”, „po co”)', '☐', '☐', '☐'],
           ['II. Mowa czynna i artykulacja', '6. Wypowiada się zdaniami złożonymi', '☐', '☐', '☐'],
           ['', '7. Mowa jest zrozumiała dla osoby spoza najbliższego otoczenia', '☐', '☐', '☐'],
           ['', '8. Realizuje poprawnie głoski właściwe dla wieku', '☐', '☐', '☐'],
           ['', '9. Utrzymuje płynność wypowiedzi (bez zacięć i powtórzeń)', '☐', '☐', '☐'],
           ['', '10. Reguluje natężenie i tempo mowy odpowiednio do sytuacji', '☐', '☐', '☐'],
           ['III. Słuch fonematyczny, czytanie i pisanie', '11. Wyodrębnia pierwszą i ostatnią głoskę w wyrazie', '☐', '☐', '☐'],
           ['', '12. Dzieli wyraz na głoski i scala głoski w wyraz', '☐', '☐', '☐'],
           ['', '13. Różnicuje głoski opozycyjne (p–b, s–z, k–g)', '☐', '☐', '☐'],
           ['', '14. Czyta wyrazy i zdania techniką właściwą dla etapu', '☐', '☐', '☐'],
           ['', '15. Zapisuje wyrazy ze słuchu bez opuszczeń i przestawień liter', '☐', '☐', '☐'],
           ['IV. Słownictwo i gramatyka', '16. Nazywa przedmioty, czynności i pojęcia z zakresu programu', '☐', '☐', '☐'],
           ['', '17. Stosuje poprawne formy fleksyjne (przypadki, liczba, rodzaj)', '☐', '☐', '☐'],
           ['', '18. Buduje zdania z użyciem przyimków i spójników', '☐', '☐', '☐'],
           ['', '19. Rozumie i stosuje pojęcia nadrzędne i podrzędne', '☐', '☐', '☐'],
           ['', '20. Wyjaśnia znaczenie nowego słowa własnymi słowami', '☐', '☐', '☐'],
           ['V. Komunikacja i budowanie wypowiedzi', '21. Inicjuje i podtrzymuje rozmowę z rówieśnikiem', '☐', '☐', '☐'],
           ['', '22. Prosi o pomoc słowem, gestem lub kartą komunikacyjną', '☐', '☐', '☐'],
           ['', '23. Opowiada zdarzenie z zachowaniem kolejności', '☐', '☐', '☐'],
           ['', '24. Odpowiada na pytania nauczyciela na forum klasy', '☐', '☐', '☐'],
           ['', '25. Stosuje zasady dialogu (czeka na swoją kolej, nie przerywa)', '☐', '☐', '☐']],
          widths=[4.2, 9.1, 1.1, 1.1, 1.2], size=7.5)
    table(doc,
          ['Obszar', 'I. Rozumienie', 'II. Mowa czynna', 'III. Słuch fon. / pisanie', 'IV. Słownictwo', 'V. Komunikacja'],
          [['Wynik (0–10)', '', '', '', '', '']],
          widths=[3.0, 2.7, 2.7, 3.2, 2.6, 2.5], size=8, align_first_bold=True)
    accent_box(doc, [
        'ODCZYT: 8–10 pkt = zasób · 4–7 pkt = obszar wymagający wsparcia · 0–3 pkt = priorytet.',
        'GRANICA KOMPETENCJI: w dokumentacji piszemy „obserwowane trudności w wyodrębnianiu głosek '
        'i w budowaniu wypowiedzi — wskazana diagnoza logopedyczna”. Nie wpisujemy rozpoznania.',
    ], title='JAK ODCZYTAĆ WYNIK')
    form_lines(doc, 'OBSERWACJE JAKOŚCIOWE — konkretne przykłady wypowiedzi ucznia (cytaty)', 3)

    # ---------------------------------------------------------------- DRUK 6
    page_break(doc)
    form_header(doc, 'DRUK 6 · ToM-SP', 'Arkusz obserwacji poznania społecznego i teorii umysłu',
                'To NIE jest test i nie prowadzi do rozpoznania. Opisujemy zachowania obserwowane '
                'w naturalnych sytuacjach szkolnych.')
    _meta(doc, 'psycholog szkolny + wychowawca', 'niski wynik w obszarze VII przy zachowanych obszarach I i IV',
          'obserwacja pedagogiczna; rozpoznanie należy do poradni',
          'sekcja VII oceny WOPF-SP → cele w sferze społeczno-emocjonalnej')
    _naglowek_ucznia(doc)
    para(doc, 'Skala: 0 = nie obserwowano · 1 = obserwowano sporadycznie lub z pomocą dorosłego · '
              '2 = obserwowano regularnie i samodzielnie · N = brak możliwości obserwacji',
         size=8, italic=True, color=GREY, after=6)
    table(doc,
          ['Poziom', 'Wskaźnik obserwacyjny', '0', '1', '2', 'N'],
          [['Wskaźniki wczesne', '1. Nawiązuje kontakt wzrokowy adekwatnie do sytuacji', '☐', '☐', '☐', '☐'],
           ['', '2. Dzieli uwagę z drugą osobą (pokazuje coś, żeby się podzielić, nie żeby dostać)', '☐', '☐', '☐', '☐'],
           ['', '3. Sprawdza reakcję dorosłego w sytuacji niejasnej', '☐', '☐', '☐', '☐'],
           ['Rozumienie emocji', '4. Rozpoznaje i nazywa podstawowe emocje z twarzy i tonu głosu', '☐', '☐', '☐', '☐'],
           ['', '5. Wskazuje przyczynę emocji („jest smutny, bo…”)', '☐', '☐', '☐', '☐'],
           ['', '6. Reaguje adekwatnie na emocję rówieśnika (pociesza, ustępuje)', '☐', '☐', '☐', '☐'],
           ['Przekonania i intencje', '7. Rozumie, że ktoś może nie wiedzieć tego, co on wie', '☐', '☐', '☐', '☐'],
           ['', '8. Rozwiązuje zadanie fałszywego przekonania I rzędu („gdzie Ola będzie szukać?”)', '☐', '☐', '☐', '☐'],
           ['', '9. Rozwiązuje zadanie II rzędu („co Ola myśli, że Kuba myśli?”)', '☐', '☐', '☐', '☐'],
           ['', '10. Odróżnia przypadkowe potrącenie od celowego zaczepienia', '☐', '☐', '☐', '☐'],
           ['Język niedosłowny', '11. Rozpoznaje żart i nie odbiera go jako obrazy', '☐', '☐', '☐', '☐'],
           ['', '12. Rozumie ironię i przenośnię w tekście i w rozmowie', '☐', '☐', '☐', '☐'],
           ['', '13. Rozumie kłamstwo uprzejmościowe i konwencje grzecznościowe', '☐', '☐', '☐', '☐'],
           ['Współpraca w grupie', '14. Przyjmuje przydzieloną rolę w pracy zespołowej', '☐', '☐', '☐', '☐'],
           ['', '15. Uwzględnia pomysł kolegi przy planowaniu wspólnego zadania', '☐', '☐', '☐', '☐'],
           ['', '16. Negocjuje i ustępuje w sporze bez udziału dorosłego', '☐', '☐', '☐', '☐']],
          widths=[3.4, 9.4, 1.0, 1.0, 1.0, 0.9], size=7.5)
    accent_box(doc, [
        'WIEK MA ZNACZENIE: zadania I rzędu rozwiązują dzieci typowo rozwijające się od ok. 4. r.ż., '
        'zadania II rzędu — od ok. 6.–7. r.ż. W szkole podstawowej kluczowe są wskaźniki 10–16.',
        'CO ZAPISUJEMY W OCENIE: opis zachowań i sytuacji, w których wystąpiły. Nie zapisujemy hipotez '
        'diagnostycznych ani nazw zaburzeń.',
    ], title='ZASTRZEŻENIA')
    form_lines(doc, 'PRZYKŁADY KONKRETNYCH SYTUACJI (data, lekcja/przerwa, przebieg)', 3)

    # ---------------------------------------------------------------- DRUK 7
    page_break(doc)
    form_header(doc, 'DRUK 7 · KFU-SP', 'Karta obserwacji funkcji uczenia się',
                'Uruchamiamy przy niskim wyniku w obszarze I (uczenie się) lub VIII (edukacja szkolna). '
                'Szukamy nie tego, czego uczeń nie umie, lecz warunków, w których zaczyna umieć.')
    _meta(doc, 'wychowawca + terapeuta pedagogiczny', 'po analizie profilu KSzOF',
          'rozpoznawanie potrzeb (t.j. Dz.U. 2023 poz. 1798)',
          'sekcje V a i XI oceny WOPF-SP → cele w sferze poznawczej')
    _naglowek_ucznia(doc)
    table(doc,
          ['Funkcja', 'Co mierzymy', 'Bez wsparcia', 'Przy wsparciu — jakim?'],
          [['Uwaga', 'Czas pracy samodzielnej nad zadaniem (minuty)', '', ''],
           ['', 'Liczba podpowiedzi potrzebnych na jednej lekcji', '', ''],
           ['Pamięć', 'Liczba elementów polecenia zapamiętanych po jednym usłyszeniu', '', ''],
           ['', 'Odtworzenie treści lekcji na kolejnych zajęciach (0–2)', '', ''],
           ['Funkcje wykonawcze', 'Rozpoczęcie zadania po poleceniu (czas w sekundach)', '', ''],
           ['', 'Doprowadzenie zadania do końca (na ile z 5 prób)', '', ''],
           ['', 'Przygotowanie przyborów / spakowanie plecaka (na ile z 5 dni)', '', ''],
           ['Tempo', 'Liczba zadań wykonanych w czasie przewidzianym dla klasy', '', ''],
           ['', 'Liczba linijek przepisanych z tablicy bez opuszczeń', '', ''],
           ['Przeniesienie', 'Zastosowanie poznanej reguły w nowym zadaniu (na ile z 5 prób)', '', '']],
          widths=[3.4, 6.8, 3.0, 3.5], size=8)
    para(doc, 'WARUNKI, KTÓRE PODNOSZĄ WYNIK — zaznacz sprawdzone w obserwacji',
         size=9, bold=True, color=PURPLE, before=6, after=4)
    checkboxes(doc, ['podział zadania na kroki', 'wsparcie wizualne / piktogramy', 'wydłużony czas',
                     'miejsce w pierwszej ławce', 'polecenie powtórzone indywidualnie',
                     'zmniejszona liczba zadań', 'praca na konkretach', 'przerwa ruchowa co ok. 15 minut',
                     'praca na komputerze', 'praca w parze z wyznaczonym kolegą',
                     'ciszej / słuchawki wyciszające', 'inne: ...............'], cols=3)
    form_lines(doc, 'WNIOSEK — przy jakich warunkach uczeń pracuje najskuteczniej (to jest gotowe źródło dostosowań)', 3)

    # ---------------------------------------------------------------- DRUK 8
    page_break(doc)
    form_header(doc, 'DRUK 8 · MG-SP', 'Ankieta „Mój głos” — perspektywa ucznia',
                'Wypełniamy Z UCZNIEM, w sposób dostosowany do jego możliwości komunikacyjnych. '
                'Zapisujemy jego słowami albo przez wskazanie.')
    _meta(doc, 'wychowawca lub psycholog wspólnie z uczniem', 'przed każdym spotkaniem zespołu',
          'element ponad wymóg rozporządzenia — decyzja rady pedagogicznej',
          'sekcje XI–XIII oceny WOPF-SP; trafność celów w IPET / PWES')
    _naglowek_ucznia(doc)
    para(doc, 'SPOSÓB POZYSKANIA GŁOSU UCZNIA — zaznacz zastosowane', size=9, bold=True, color=PURPLE, before=6, after=4)
    checkboxes(doc, ['rozmowa dostosowana do możliwości ucznia', 'obserwacja na lekcjach i przerwach',
                     'wybór z obrazków / piktogramów / kart wyboru', 'komunikacja wspomagająca i alternatywna (AAC)',
                     'odczytanie z gestu, mimiki, zachowania', 'skala buziek / termometr emocji',
                     'wskazanie „tak / nie” na symbolach', 'informacje od rodziców'], cols=2)
    form_lines(doc, 'CO LUBIĘ ROBIĆ I W CZYM JESTEM DOBRY / DOBRA?', 2)
    form_lines(doc, 'Z CZYM MAM NAJWIĘKSZĄ TRUDNOŚĆ W SZKOLE?', 2)
    form_lines(doc, 'KTÓRA LEKCJA JEST DLA MNIE NAJTRUDNIEJSZA I DLACZEGO?', 2)
    form_lines(doc, 'CO SIĘ DZIEJE NA PRZERWACH? Z KIM LUBIĘ SPĘDZAĆ CZAS?', 2)
    para(doc, 'CO MI NAJBARDZIEJ POMAGA NA LEKCJACH? — zaznacz', size=9, bold=True, color=PURPLE, before=6, after=4)
    checkboxes(doc, ['cisza', 'więcej czasu', 'ruch i przerwy', 'obrazki i piktogramy',
                     'praca na komputerze', 'krótsze zadania', 'pomoc pani / pana', 'praca z kolegą',
                     'przykład pokazany przed zadaniem', 'możliwość odpowiedzi ustnej'], cols=2)
    form_lines(doc, 'CZEGO CHCIAŁBYM / CHCIAŁABYM SIĘ NAUCZYĆ W TYM ROKU?', 2)
    para(doc, 'Zapis sporządzono: ☐ słowami ucznia (cytat)   ☐ przez wskazanie / symbole   ☐ z pomocą AAC   '
              '☐ na podstawie obserwacji, gdy uczeń nie komunikuje się werbalnie',
         size=8, color=INK, before=4)

    # ---------------------------------------------------------------- DRUK 9
    page_break(doc)
    form_header(doc, 'DRUK 9 · AR-SP', 'Ankieta dla rodzica i wywiad środowiskowy',
                'Przekazywana rodzicom we wrześniu, równolegle z arkuszem wychowawcy. Rodzic wypełnia '
                'kwestionariusz KSzOF niezależnie — ta ankieta jest jego uzupełnieniem.')
    _meta(doc, 'rodzic / opiekun prawny; wywiad prowadzi pedagog lub wychowawca', 'wrzesień oraz przed każdą oceną',
          'współpraca z rodzicami — § 6 rozp. t.j. Dz.U. 2020 poz. 1309',
          'sekcje X–XIII oceny WOPF-SP; obszar VI KSzOF')
    _naglowek_ucznia(doc)
    form_lines(doc, 'CO MOJE DZIECKO LUBI ROBIĆ W DOMU? W CZYM JEST DOBRE?', 2)
    form_lines(doc, 'JAK WYGLĄDA ODRABIANIE LEKCJI — ile czasu, przy jakiej pomocy, co przeszkadza?', 2)
    form_lines(doc, 'CO DZIECKO OPOWIADA O SZKOLE? CO JE CIESZY, A CO NIEPOKOI?', 2)
    para(doc, 'SAMODZIELNOŚĆ W DOMU (obszar VI KSzOF) — zaznacz', size=9, bold=True, color=PURPLE, before=6, after=4)
    table(doc, ['Czynność', 'Samodzielnie', 'Z przypomnieniem', 'Z pomocą', 'Nie wykonuje'],
          [['Utrzymuje porządek w swoim miejscu do nauki', '☐', '☐', '☐', '☐'],
           ['Pomaga w drobnych pracach domowych', '☐', '☐', '☐', '☐'],
           ['Przygotowuje tornister na następny dzień', '☐', '☐', '☐', '☐'],
           ['Dba o higienę osobistą', '☐', '☐', '☐', '☐'],
           ['Korzysta z telefonu / komputera w ustalonych granicach', '☐', '☐', '☐', '☐']],
          widths=[7.4, 2.4, 2.6, 2.2, 2.2], size=8)
    form_lines(doc, 'CZY W OSTATNIM ROKU WYDARZYŁO SIĘ COŚ WAŻNEGO DLA DZIECKA? (zmiana, strata, choroba, przeprowadzka)', 2)
    form_lines(doc, 'JAKIEGO WSPARCIA OCZEKUJEMY OD SZKOŁY? CO SPRAWDZIŁO SIĘ WCZEŚNIEJ?', 2)
    form_lines(doc, 'SPECJALIŚCI SPOZA SZKOŁY, POD KTÓRYCH OPIEKĄ JEST DZIECKO (zakres, częstotliwość — bez kopii dokumentacji medycznej)', 2)
    table(doc, ['Data', 'Podpis rodzica', 'Osoba przyjmująca ankietę', 'Odnotowano w rejestrze kontaktów'],
          [['', '', '', '☐ tak']], widths=[2.6, 4.4, 5.0, 4.7], size=8, align_first_bold=False)


def zal_C_3(doc):
    # ---------------------------------------------------------------- DRUK 10
    page_break(doc)
    form_header(doc, 'DRUK 10 · NP-SP', 'Arkusz informacji od nauczyciela przedmiotu (II etap edukacyjny)',
                'Druk, którego nie ma w przedszkolu i bez którego opinia szkoły w II etapie jest niepełna. '
                'Jedna strona na jednego nauczyciela — zbiera wychowawca.')
    _meta(doc, 'każdy nauczyciel uczący ucznia', 'październik oraz przed każdą oceną wielospecjalistyczną',
          'zintegrowane działania nauczycieli — § 6 ust. 1 pkt 2 rozp. t.j. Dz.U. 2020 poz. 1309',
          'sekcje V a i XI oceny WOPF-SP; opinia dla poradni')
    table(doc, ['Uczeń', 'Klasa', 'Przedmiot', 'Nauczyciel', 'Liczba godzin tygodniowo z uczniem'],
          [['', '', '', '', '']], widths=[4.2, 1.8, 3.4, 4.0, 3.3], size=8, align_first_bold=False)
    table(doc,
          ['Obszar', 'Pytanie', 'Odpowiedź (konkretnie, z częstotliwością)'],
          [['Mocne strony', 'Co uczeń robi na mojej lekcji dobrze? Kiedy widać go z najlepszej strony?', ''],
           ['Uczestnictwo', 'W jakiej części lekcji uczestniczy aktywnie? Co go z niej wytrąca?', ''],
           ['Rozumienie poleceń', 'Czy wykonuje polecenie po pierwszym usłyszeniu? Ile podpowiedzi potrzebuje?', ''],
           ['Tempo pracy', 'Ile z zadań przewidzianych dla klasy wykonuje w tym samym czasie?', ''],
           ['Prace pisemne', 'Jak wygląda praca pisemna: czytelność, opuszczenia, długość, tempo?', ''],
           ['Sprawdziany', 'Jak wypada na sprawdzianie w porównaniu z pracą bieżącą? Przy jakiej formie wypada lepiej?', ''],
           ['Relacje w klasie', 'Jak pracuje w parze i w grupie? Jak reaguje na uwagę i na ocenę?', ''],
           ['Skuteczne dostosowania', 'Które z zastosowanych przeze mnie dostosowań realnie pomogły?', ''],
           ['Bariery', 'Co na mojej lekcji utrudnia uczniowi pracę (hałas, tablica, tempo, forma zadania)?', ''],
           ['Potrzeba', 'Czego potrzebuję, żeby lepiej pracować z tym uczniem?', '']],
          widths=[3.2, 6.8, 6.7], size=8)
    table(doc, ['Data', 'Podpis nauczyciela', 'Przyjął (wychowawca / koordynator)'],
          [['', '', '']], widths=[3.0, 6.8, 6.9], size=8, align_first_bold=False)

    # ---------------------------------------------------------------- DRUK 11
    page_break(doc)
    form_header(doc, 'DRUK 11 · KDP-SP', 'Karta dostosowań przedmiotowych',
                'Jedna karta na jeden przedmiot. Nauczyciel otrzymuje ją we wrześniu. Zapis „wydłużenie czasu pracy” '
                'w programie nie mówi nic nauczycielowi geografii — ta karta mówi wszystko.')
    _meta(doc, 'koordynator zespołu na podstawie IPET / PWES', 'wrzesień; aktualizacja po każdej ocenie',
          '§ 6 ust. 1 pkt 1 rozp. t.j. Dz.U. 2020 poz. 1309; rozp. o ocenianiu [[t.j. Dz.U. 2023 poz. 2572, '
          'z późn. zm.]]',
          'praca bieżąca nauczyciela; podstawa dostosowań egzaminacyjnych (druk 12)')
    table(doc, ['Uczeń', 'Klasa', 'Przedmiot', 'Podstawa dostosowania'],
          [['', '', '', '☐ orzeczenie   ☐ opinia PPP   ☐ rozpoznanie nauczycieli (zapisane w dokumentacji)']],
          widths=[3.8, 1.8, 3.4, 7.7], size=8, align_first_bold=False)
    table(doc,
          ['Rodzaj dostosowania', 'Co konkretnie robimy na TYM przedmiocie', 'Bariera z WOPF, z której wynika'],
          [['Sposób podania treści', '', ''],
           ['Czas', '', ''],
           ['Przestrzeń i miejsce', '', ''],
           ['Sposób sprawdzania wiedzy', '', ''],
           ['Pomoce i materiały', '', ''],
           ['Forma i objętość pracy domowej', '', ''],
           ['Zasady oceniania (co oceniamy, czego nie)', '', '']],
          widths=[4.2, 6.6, 5.9], size=8)
    accent_box(doc, [
        'DOSTOSOWANIE ≠ OBNIŻENIE WYMAGAŃ. Zmieniamy sposób, w jaki uczymy i sprawdzamy, a nie zakres podstawy '
        'programowej. Wyjątek: uczeń z niepełnosprawnością intelektualną w stopniu umiarkowanym lub znacznym '
        'realizuje odrębną podstawę programową.',
        'KAŻDE DOSTOSOWANIE MA ŹRÓDŁO. Jeżeli w kolumnie „bariera” nic nie stoi, wracamy do oceny — '
        'albo bariery nie opisaliśmy, albo dostosowanie jest zbędne.',
        'OCENA ZACHOWANIA. Przy ustalaniu oceny zachowania ucznia z orzeczeniem lub opinią uwzględnia się wpływ '
        'stwierdzonych zaburzeń na jego zachowanie — zapis o tym umieszczamy w dokumentacji '
        '([[rozp. o ocenianiu, t.j. Dz.U. 2023 poz. 2572, z późn. zm.]]).',
    ], title='TRZY ZASADY DOSTOSOWAŃ')
    table(doc, ['Data przekazania nauczycielowi', 'Podpis nauczyciela', 'Data przeglądu skuteczności', 'Wynik przeglądu'],
          [['', '', '', '']], widths=[4.4, 4.0, 4.0, 4.3], size=8, align_first_bold=False)

    # ---------------------------------------------------------------- DRUK 12
    page_break(doc)
    form_header(doc, 'DRUK 12 · AD-E8', 'Karta dostosowań warunków i form egzaminu ósmoklasisty',
                'Zakładamy w listopadzie klasy VIII. ⚑ Katalog dostosowań i terminy każdorazowo sprawdzamy '
                'w komunikacie dyrektora Centralnej Komisji Egzaminacyjnej na dany rok szkolny.')
    _meta(doc, 'dyrektor / koordynator na podstawie dokumentacji ucznia', 'listopad klasy VIII ⚑ termin z komunikatu CKE',
          'ustawa o systemie oświaty — art. 44zzr; [[komunikat dyrektora CKE na dany rok szkolny — jedyne źródło '
          'katalogu dostosowań i terminów; sprawdź przed wypełnieniem ⚑]]',
          'protokół rady pedagogicznej; organizacja egzaminu')
    table(doc, ['Uczeń', 'Klasa', 'Rok szkolny', 'Podstawa uprawnienia do dostosowania'],
          [['', '', '', '☐ orzeczenie o potrzebie kształcenia specjalnego, nr ...............\n'
                       '☐ orzeczenie o potrzebie indywidualnego nauczania, nr ...............\n'
                       '☐ opinia poradni psychologiczno-pedagogicznej, nr ...............\n'
                       '☐ zaświadczenie lekarskie o stanie zdrowia\n'
                       '☐ pozytywna opinia rady pedagogicznej (uczeń objęty pomocą pp), data uchwały ...............']],
          widths=[3.6, 1.6, 2.4, 9.1], size=8, align_first_bold=False)
    para(doc, 'DOKUMENTACJA, NA KTÓREJ OPIERAMY WNIOSEK — wskaż druki i daty', size=9, bold=True,
         color=PURPLE, before=6, after=4)
    table(doc, ['WOPF-SP (data)', 'IPET / PWES (nr, data)', 'Karty dostosowań przedmiotowych (przedmioty)',
                'Karty ewaluacji (daty)'],
          [['', '', '', '']], widths=[3.4, 4.0, 5.4, 3.9], size=8, align_first_bold=False)
    para(doc, 'WNIOSKOWANE SPOSOBY DOSTOSOWANIA — zaznacz i uzasadnij każdą pozycję',
         size=9, bold=True, color=PURPLE, before=6, after=4)
    table(doc,
          ['Sposób dostosowania', 'Wniosk.', 'Uzasadnienie — bariera opisana w dokumentacji ucznia'],
          [['Przedłużenie czasu pracy', '☐', ''],
           ['Zaznaczanie odpowiedzi w arkuszu (bez karty odpowiedzi)', '☐', ''],
           ['Arkusz dostosowany do rodzaju niepełnosprawności', '☐', ''],
           ['Pomoc nauczyciela wspomagającego w czytaniu / pisaniu', '☐', ''],
           ['Zapisywanie odpowiedzi na komputerze', '☐', ''],
           ['Oddzielna sala', '☐', ''],
           ['Przerwy w trakcie egzaminu', '☐', ''],
           ['Pomoce optyczne / techniczne, sprzęt medyczny', '☐', ''],
           ['Nieprzenoszenie odpowiedzi / dostosowanie kryteriów oceniania', '☐', ''],
           ['Inne wynikające z komunikatu CKE: ...............', '☐', '']],
          widths=[6.4, 1.4, 8.9], size=8)
    table(doc, ['Data poinformowania rodziców', 'Oświadczenie rodziców (data)', 'Uchwała rady pedagogicznej (data, nr)',
                'Podpis dyrektora'],
          [['', '', '', '']], widths=[4.4, 4.0, 4.4, 3.9], size=8, align_first_bold=False)
    para(doc, 'UWAGA: terminy poinformowania rodziców i złożenia oświadczenia określa komunikat dyrektora CKE '
              'oraz harmonogram na dany rok. Sprawdź je przed wypełnieniem karty. ⚑',
         size=8, italic=True, color=ORANGE, before=4)

    # ---------------------------------------------------------------- DRUK 13
    page_break(doc)
    form_header(doc, 'DRUK 13 · SMART-SP', 'Karta celu i ewaluacji',
                'Jedna karta na jeden cel. Kryterium zapisane w celu jest gotowym wskaźnikiem ewaluacji — '
                'nie tworzymy do niej osobnego narzędzia.')
    _meta(doc, 'nauczyciel / specjalista odpowiedzialny za cel', 'przy opracowaniu IPET lub PWES; pomiar wg planu',
          'ocena efektywności — § 6 rozp. t.j. Dz.U. 2020 poz. 1309',
          'karta oceny efektywności (druk 14); modyfikacja IPET / PWES')
    table(doc, ['Uczeń', 'Klasa', 'Sfera / obszar KSzOF', 'Nr celu', 'Okres realizacji'],
          [['', '', '', '', '']], widths=[4.0, 1.6, 4.6, 1.6, 4.9], size=8, align_first_bold=False)
    form_lines(doc, 'STAN WYJŚCIOWY — co uczeń robi DZIŚ (zachowanie, warunki, wsparcie, częstotliwość)', 2)
    form_lines(doc, 'CEL — uczeń, w sytuacji ……, wykona ……, w …… z …… prób, przy wsparciu ……, do dnia ……', 3)
    table(doc,
          ['Litera', 'Sprawdzenie', 'Tak'],
          [['S — konkretny', 'Czy zapisano, jakie zachowanie i w jakiej sytuacji?', '☐'],
           ['M — mierzalny', 'Czy zapisano liczbę prób i poziom wsparcia?', '☐'],
           ['A — osiągalny', 'Czy cel jest JEDNYM krokiem od tego, co uczeń robi dziś?', '☐'],
           ['R — istotny', 'Czy cel wynika z oceny i zwiększa uczestnictwo ucznia?', '☐'],
           ['T — terminowy', 'Czy zapisano datę końcową i terminy pomiarów?', '☐'],
           ['Test dwóch osób', 'Czy dwie różne osoby, patrząc na to samo dziecko, ocenią cel tak samo?', '☐']],
          widths=[3.4, 11.0, 2.3], size=8)
    table(doc,
          ['Termin pomiaru', 'Wartość osiągnięta', 'Kto mierzył', 'Decyzja zespołu'],
          [['', '', '', '☐ cel osiągnięty — zamykamy   ☐ częściowo — kontynuujemy\n'
                       '☐ brak postępu — modyfikujemy metodę   ☐ regres — spotkanie z rodzicami'],
           ['', '', '', '☐ cel osiągnięty — zamykamy   ☐ częściowo — kontynuujemy\n'
                       '☐ brak postępu — modyfikujemy metodę   ☐ regres — spotkanie z rodzicami'],
           ['', '', '', '☐ cel osiągnięty — zamykamy   ☐ częściowo — kontynuujemy\n'
                       '☐ brak postępu — modyfikujemy metodę   ☐ regres — spotkanie z rodzicami']],
          widths=[3.0, 3.0, 3.0, 7.7], size=8, align_first_bold=False)
    para(doc, 'ZASADA: przy wyniku „częściowo” przesuwamy termin, ale NIE obniżamy kryterium. '
              'Przy „braku postępu” zmieniamy metodę i sprawdzamy bariery środowiskowe, a nie cel.',
         size=8, italic=True, color=GREY, before=4)


def zal_C_4(doc):
    # ---------------------------------------------------------------- DRUK 14
    page_break(doc)
    form_header(doc, 'DRUK 14 · OE-SP', 'Karta oceny efektywności udzielanej pomocy',
                'Ocena efektywności jest obowiązkiem, nie dobrą praktyką. Wypełniamy w styczniu i w czerwcu '
                'oraz na zakończenie każdej formy pomocy.')
    _meta(doc, 'nauczyciele i specjaliści prowadzący zajęcia z uczniem', 'styczeń, czerwiec, koniec formy pomocy',
          'rozp. o pomocy pp (t.j. Dz.U. 2023 poz. 1798); § 6 rozp. t.j. Dz.U. 2020 poz. 1309',
          'modyfikacja IPET / PWES; kolejna ocena wielospecjalistyczna; opinia dla poradni')
    table(doc, ['Uczeń', 'Klasa', 'Okres objęty oceną', 'Ścieżka'],
          [['', '', '', '☐ kształcenie specjalne (IPET)   ☐ pomoc pp (PWES)']],
          widths=[4.4, 1.8, 4.4, 6.1], size=8, align_first_bold=False)
    table(doc,
          ['Forma pomocy / zajęcia', 'Prowadzący', 'Wymiar (h/tydz.)', 'Frekwencja', 'Efekt — co się zmieniło (mierzalnie)',
           'Decyzja'],
          [['', '', '', '', '', ''] for _ in range(6)],
          widths=[3.6, 2.4, 1.8, 1.6, 5.2, 2.1], size=8, align_first_bold=False)
    form_lines(doc, 'CO ZADZIAŁAŁO — metody, formy i warunki, które przyniosły efekt (przenosimy do kolejnego okresu)', 2)
    form_lines(doc, 'CO NIE ZADZIAŁAŁO — i jaka jest hipoteza dlaczego (metoda? wymiar? bariera środowiskowa?)', 2)
    form_lines(doc, 'BARIERY, KTÓRE UJAWNIŁY SIĘ W TYM OKRESIE', 2)
    para(doc, 'WNIOSEK KOŃCOWY — zaznacz', size=9, bold=True, color=PURPLE, before=6, after=4)
    checkboxes(doc, ['kontynuujemy wsparcie w dotychczasowym zakresie',
                     'zwiększamy wymiar lub zmieniamy formę pomocy',
                     'zmieniamy metody pracy przy tym samym wymiarze',
                     'usuwamy barierę środowiskową i mierzymy ponownie',
                     'kończymy formę pomocy — cel osiągnięty',
                     'występujemy do poradni (za zgodą rodziców) — brak poprawy mimo pomocy'], cols=2)
    table(doc, ['Data', 'Podpisy nauczycieli i specjalistów', 'Rodzic poinformowany (data, podpis)'],
          [['', '', '']], widths=[2.6, 8.0, 6.1], size=8, align_first_bold=False)

    # ---------------------------------------------------------------- DRUK 15
    page_break(doc)
    form_header(doc, 'DRUK 15 · PWES', 'Plan wsparcia edukacyjnego ucznia (bez orzeczenia)',
                'Narzędzie wewnętrzne placówki — nie druk wymagany przepisem. Porządkuje to, co i tak musimy '
                'ustalić i ocenić. Ta sama logika co IPET, ale bez rangi rozporządzenia.')
    _meta(doc, 'nauczyciele i specjaliści prowadzący zajęcia; koordynuje wychowawca',
          'niezwłocznie po ustaleniu form pomocy przez dyrektora',
          '[[NARZĘDZIE WEWNĘTRZNE SZKOŁY — wprowadza zarządzenie dyrektora. Rozporządzenie o pomocy pp '
          '(t.j. Dz.U. 2023 poz. 1798) NIE przewiduje takiego druku.]]',
          'karta oceny efektywności (druk 14); ewentualny wniosek do poradni')
    table(doc, ['Uczeń', 'Klasa', 'Rok szkolny', 'Podstawa objęcia pomocą'],
          [['', '', '', '☐ opinia poradni nr ...............   ☐ rozpoznanie nauczycieli (data, dokument)\n'
                       '☐ wniosek rodzica   ☐ wniosek ucznia   ☐ informacja od poprzedniej szkoły']],
          widths=[3.8, 1.6, 2.4, 8.9], size=8, align_first_bold=False)
    table(doc,
          ['Rozpoznana potrzeba (z KSzOF / obserwacji)', 'Cel z kryterium', 'Forma pomocy', 'Prowadzący',
           'Wymiar i okres', 'Termin oceny efektywności'],
          [['', '', '', '', '', ''] for _ in range(5)],
          widths=[4.0, 3.6, 2.6, 2.2, 2.2, 2.1], size=8, align_first_bold=False)
    form_lines(doc, 'DOSTOSOWANIA W PRACY BIEŻĄCEJ NA LEKCJACH — wspólne dla wszystkich przedmiotów '
                    '(szczegóły w kartach przedmiotowych, druk 11)', 3)
    form_lines(doc, 'DZIAŁANIA WSPIERAJĄCE RODZICÓW I ZAKRES WSPÓŁPRACY', 2)
    table(doc, ['Data sporządzenia', 'Koordynator', 'Rodzic poinformowany o formach, wymiarze i okresie (data, podpis)'],
          [['', '', '']], widths=[3.4, 4.6, 8.7], size=8, align_first_bold=False)
    accent_box(doc, [
        '[[STATUS DRUKU. Rozporządzenie o pomocy psychologiczno-pedagogicznej nie przewiduje planu wsparcia '
        'jako dokumentu obowiązkowego. Z przepisów wynika, że formy pomocy, okres ich udzielania i wymiar godzin '
        'ustala dyrektor, nauczyciele i specjaliści oceniają efektywność udzielanej pomocy, a przebieg zajęć '
        'dokumentuje się w dziennikach zajęć. Ten druk wprowadzamy zarządzeniem dyrektora jako narzędzie wewnętrzne '
        '— tak samo jak metryczkę. W stopce druku i w dokumentacji NIE powołujemy się na rozporządzenie jako '
        'na podstawę jego sporządzenia.]]',
        'ZINDYWIDUALIZOWANA ŚCIEŻKA KSZTAŁCENIA (§ 12 rozp. o pomocy pp) wymaga OPINII publicznej poradni '
        'i wniosku rodziców. Nie stosuje się jej wobec ucznia objętego kształceniem specjalnym ani indywidualnym '
        'nauczaniem. To osobna forma — nie mylimy jej z nauczaniem indywidualnym '
        '([[rozp. MEN z 9 sierpnia 2017 r., t.j. Dz.U. 2023 poz. 2468, z późn. zm.]]).',
    ], title='UWAGA — STATUS DRUKU I ROZRÓŻNIENIE, KTÓRE MYLI SIĘ NAJCZĘŚCIEJ')

    # ---------------------------------------------------------------- DRUK 16
    page_break(doc)
    form_header(doc, 'DRUK 16 · KP-SP', 'Karta przekazania informacji o uczniu — z I do II etapu edukacyjnego',
                'Wypełniamy w czerwcu klasy III. W klasie III ucznia zna jedna wychowawczyni. W klasie IV — '
                'dziesięcioro nauczycieli, z których każdy widzi go dwie godziny w tygodniu.')
    _meta(doc, 'wychowawca klasy III wspólnie ze specjalistami', 'czerwiec klasy III',
          'decyzja rady pedagogicznej wpisana do procedury szkoły; ciągłość wsparcia',
          'dokumentacja II etapu; karty dostosowań przedmiotowych na wrzesień')
    table(doc, ['Uczeń', 'Klasa III (wychowawca)', 'Klasa IV (wychowawca)', 'Data przekazania'],
          [['', '', '', '']], widths=[4.4, 4.0, 4.0, 4.3], size=8, align_first_bold=False)
    table(doc,
          ['Zagadnienie', 'Co przekazujemy'],
          [['Podstawa wsparcia', 'orzeczenie / opinia / rozpoznanie — numer, data, ważność; termin kolejnej oceny'],
           ['Mocne strony i zainteresowania', 'to, co działa — od tego zaczyna się wrzesień w klasie IV'],
           ['Dostosowania, które SIĘ SPRAWDZIŁY', 'konkretnie, z podaniem sytuacji; osobno dla pracy pisemnej i ustnej'],
           ['Dostosowania, które NIE zadziałały', 'żeby nauczyciele klasy IV nie powtarzali tej samej drogi'],
           ['Sygnały sensoryczne i zdrowotne', 'z metryczki — hałas, korytarz, stołówka, leki, procedura'],
           ['Zachowania trudne i procedura', 'funkcja zachowania z FBA, zachowanie zastępcze, jednolita reakcja dorosłych'],
           ['Komunikacja', 'sposób porozumiewania się, karty, AAC, poziom rozumienia poleceń'],
           ['Relacje w klasie', 'z kim współpracuje, kogo unika, jak reaguje na zmianę składu grupy'],
           ['Współpraca z rodzicami', 'preferowana forma kontaktu, ustalenia, których pilnujemy'],
           ['Aktualne cele', 'numery i treść celów, które są w toku, z terminami pomiaru'],
           ['Do zrobienia we wrześniu', 'lista zadań dla nowego wychowawcy — z terminami']],
          widths=[4.4, 12.3], size=8)
    table(doc, ['Podpis wychowawcy klasy III', 'Podpis wychowawcy klasy IV', 'Podpis koordynatora / pedagoga specjalnego'],
          [['', '', '']], widths=[5.6, 5.6, 5.5], size=8, align_first_bold=False)

    # ---------------------------------------------------------------- DRUK 17
    page_break(doc)
    form_header(doc, 'DRUK 17 · AUD-SP', 'Arkusz audytu dokumentacji kształcenia specjalnego',
                'Wypełniamy raz w roku, przed sierpniową radą pedagogiczną. Jeden arkusz na całą szkołę '
                'plus wykaz uczniów wymagających działania.')
    _meta(doc, 'koordynator / pedagog specjalny; zatwierdza dyrektor', 'sierpień / wrzesień',
          'nadzór dyrektora nad dokumentacją — [[art. 68 ust. 1 ustawy Prawo oświatowe (t.j. Dz.U. 2026 poz. 820); '
          'rozp. o nadzorze pedagogicznym, t.j. Dz.U. 2024 poz. 15]]',
          'plan nadzoru pedagogicznego; harmonogram wdrożenia (załącznik E)')
    table(doc,
          ['#', 'Pytanie kontrolne', 'Tak', 'Nie', 'Działanie naprawcze i termin'],
          [['1', 'Czy we wszystkich wzorach druków stoi obowiązujący tekst jednolity, a nie publikator pierwotny?', '☐', '☐', ''],
           ['2', 'Czy każdy uczeń z orzeczeniem ma IPET sporządzony w terminie (30 września / 30 dni od wpływu)?', '☐', '☐', ''],
           ['3', 'Czy data wpływu orzeczenia do szkoły jest odnotowana w metryczce każdego ucznia?', '☐', '☐', ''],
           ['4', 'Czy każdy IPET zawiera wszystkie osiem elementów z § 6 ust. 1?', '☐', '☐', ''],
           ['5', 'Czy przy każdym zaleceniu z orzeczenia wskazano formę, osobę i wymiar realizacji?', '☐', '☐', ''],
           ['6', 'Czy WOPFU poprzedza IPET i czy dokonano jej co najmniej dwa razy w roku?', '☐', '☐', ''],
           ['7', 'Czy każdy wniosek w ocenie ma wskazane źródło (druk, z którego pochodzi)?', '☐', '☐', ''],
           ['8', 'Czy cele w programach mają kryterium liczbowe i datę pomiaru?', '☐', '☐', ''],
           ['9', 'Czy przeprowadzono i udokumentowano ocenę efektywności udzielanej pomocy?', '☐', '☐', ''],
           ['10', 'Czy dla uczniów bez orzeczenia ustalono formy, okres i wymiar pomocy (dyrektor) oraz czy '
            'zajęcia są dokumentowane w dziennikach? [[PWES — jeśli szkoła go wprowadziła zarządzeniem]]',
            '☐', '☐', ''],
           ['11', 'Czy dostosowania zapisano przedmiotowo i przekazano nauczycielom przedmiotów?', '☐', '☐', ''],
           ['12', 'Czy zatrudniono nauczyciela współorganizującego tam, gdzie jest to obowiązkowe?', '☐', '☐', ''],
           ['13', 'Czy w teczkach są potwierdzenia przekazania rodzicom kopii oceny i programu?', '☐', '☐', ''],
           ['14', 'Czy prowadzony jest rejestr kontaktów z rodzicami?', '☐', '☐', ''],
           ['15', 'Czy druki nie powielają danych zbędnych (PESEL, adres, dane o pracy rodziców)?', '☐', '☐', ''],
           ['16', 'Czy teczki uczniów są przechowywane w miejscu wskazanym zarządzeniem dyrektora?', '☐', '☐', ''],
           ['17', 'Czy klauzule informacyjne RODO są podpisane i aktualne?', '☐', '☐', ''],
           ['18', 'Czy dla uczniów klas VIII założono karty dostosowań egzaminacyjnych w terminie? ⚑', '☐', '☐', ''],
           ['19', 'Czy dla uczniów klas III sporządzono karty przekazania do II etapu?', '☐', '☐', ''],
           ['20', 'Czy wyznaczono koordynatorów zespołów i Strażnika Prawa na bieżący rok?', '☐', '☐', '']],
          widths=[0.9, 8.6, 1.1, 1.1, 5.0], size=7.5)
    table(doc, ['Liczba uczniów z orzeczeniem', 'Liczba uczniów objętych pomocą pp', 'Liczba pozycji „nie”',
                'Data audytu', 'Podpis dyrektora'],
          [['', '', '', '', '']], widths=[3.4, 3.6, 2.6, 2.6, 4.5], size=8, align_first_bold=False)

    # ---------------------------------------------------------------- DRUK 18
    page_break(doc)
    form_header(doc, 'DRUK 18 · RK-SP', 'Rejestr kontaktów z rodzicami',
                'Prowadzony na bieżąco. To ten rejestr zasili opinię dla poradni i będzie dowodem w każdej '
                'sytuacji spornej. Rozmowa nieodnotowana nie istnieje.')
    _meta(doc, 'wychowawca i specjaliści pracujący z uczniem', 'na bieżąco przez cały rok szkolny',
          '§ 6 rozp. t.j. Dz.U. 2020 poz. 1309 — prawa rodziców; zakres współpracy',
          'opinia dla poradni; sekcja o współpracy z rodzicami w WOPF-SP')
    table(doc, ['Uczeń', 'Klasa', 'Rok szkolny', 'Wychowawca'],
          [['', '', '', '']], widths=[5.0, 2.0, 3.0, 6.7], size=8, align_first_bold=False)
    table(doc,
          ['Lp.', 'Data', 'Forma', 'Osoba ze strony szkoły', 'Temat i ustalenia', 'Podpis rodzica'],
          [[str(i), '', '', '', '', ''] for i in range(1, 13)],
          widths=[0.9, 2.0, 2.2, 3.4, 5.9, 2.3], size=8, align_first_bold=False)
    para(doc, 'FORMA: R = rozmowa indywidualna · T = kontakt telefoniczny · Z = zebranie · S = spotkanie zespołu · '
              'M = wiadomość w dzienniku elektronicznym · K = przekazanie kopii dokumentu',
         size=8, italic=True, color=GREY, after=4)
    accent_box(doc, [
        'ODNOTOWUJEMY OSOBNO: przekazanie kopii wielospecjalistycznej oceny, przekazanie kopii programu '
        'oraz zawiadomienie o terminie spotkania zespołu. To realizacja praw rodziców wynikających wprost '
        'z rozporządzenia — musi mieć datę i podpis.',
    ], title='TRZY WPISY, KTÓRYCH NIE WOLNO POMINĄĆ')


# ================================================================ ZAŁĄCZNIKI D i E

def zal_D_E(doc):
    page_break(doc)
    band(doc, 'ZAŁĄCZNIK D', 'Wzór zarządzenia dyrektora i procedura dokumentacyjna')
    para(doc, 'Wzór do dostosowania do warunków placówki. Elementy oznaczone ⚑ wymagają sprawdzenia publikatora '
              'przed podpisaniem. Zarządzenie nie tworzy nowych obowiązków wobec uczniów i rodziców — porządkuje '
              'sposób wykonywania obowiązków, które wynikają z przepisów.', size=8.5, italic=True, color=GREY, after=8)

    accent_box(doc, [
        'ZARZĄDZENIE NR …… / 2026 Dyrektora Szkoły Podstawowej nr …… w ……',
        'z dnia …… sierpnia 2026 r.',
        'w sprawie organizacji dokumentacji ucznia objętego kształceniem specjalnym oraz pomocą '
        'psychologiczno-pedagogiczną w roku szkolnym 2026/2027',
        'Na podstawie art. 68 ust. 1 ustawy z dnia 14 grudnia 2016 r. — Prawo oświatowe '
        '([[t.j. Dz.U. 2026 poz. 820, z późn. zm.]]), rozporządzenia MEN z dnia 9 sierpnia 2017 r. w sprawie '
        'warunków organizowania kształcenia (…) (t.j. Dz.U. 2020 poz. 1309), rozporządzenia MEN z dnia 9 sierpnia '
        '2017 r. w sprawie pomocy psychologiczno-pedagogicznej (t.j. Dz.U. 2023 poz. 1798) oraz rozporządzenia MEN '
        'z dnia 25 sierpnia 2017 r. w sprawie sposobu prowadzenia dokumentacji przebiegu nauczania '
        '(t.j. Dz.U. 2024 poz. 50) — zarządzam, co następuje:',
    ], title='NAGŁÓWEK')

    numbered(doc, [
        '**§ 1. Teczka ucznia.** Dla każdego ucznia objętego kształceniem specjalnym albo pomocą '
        'psychologiczno-pedagogiczną zakłada się teczkę ucznia. Teczkę zakłada wychowawca w terminie 7 dni '
        'od dnia wpływu orzeczenia, opinii albo od dnia objęcia ucznia pomocą. Pierwszym dokumentem w teczce '
        'jest metryczka ucznia.',
        '**§ 2. Miejsce przechowywania i dostęp.** Teczki przechowuje się w …… (wskazać pomieszczenie i sposób '
        'zabezpieczenia). Dostęp do teczek mają: dyrektor, wicedyrektor, wychowawca, nauczyciele i specjaliści '
        'pracujący z uczniem oraz koordynator zespołu — w zakresie niezbędnym do wykonywania zadań. Wgląd rodzica '
        'odnotowuje się w rejestrze kontaktów.',
        '**§ 3. Zakres danych.** W drukach pochodnych nie powiela się numeru PESEL, adresu zamieszkania, miejsca '
        'urodzenia ani danych o miejscu pracy rodziców. Dane te pozostają wyłącznie w dokumentacji przebiegu '
        'nauczania. Do teczki nie włącza się kopii dokumentacji medycznej wykraczającej poza informacje niezbędne '
        'do organizacji kształcenia i zapewnienia bezpieczeństwa.',
        '**§ 4. Kwestionariusz obserwacji.** Kwestionariusz Szkolnej Oceny Funkcjonalnej (KSzOF) wypełnia się '
        'dla wszystkich uczniów oddziału dwukrotnie w roku szkolnym: we wrześniu (pomiar bazowy) i w maju '
        '(pomiar kontrolny). Arkusz wypełniają niezależnie wychowawca i rodzic, a w klasach IV–VIII dodatkowo '
        'co najmniej trzej nauczyciele przedmiotów oraz specjalista. Wypełnione arkusze wychowawca przekazuje '
        'koordynatorowi w terminie do …… .',
        '**§ 5. Reguły przekierowania i obserwacja pogłębiona.** Obserwację pogłębioną uruchamia się po spełnieniu '
        'co najmniej jednej z siedmiu reguł przekierowania określonych w procedurze stanowiącej załącznik nr 1 '
        'do niniejszego zarządzenia. Każdą decyzję — również odmowną — dokumentuje się kartą decyzyjną (druk KD-SP).',
        '**§ 6. Zespoły i koordynatorzy.** Dla każdego ucznia objętego kształceniem specjalnym dyrektor wyznacza '
        'koordynatora zespołu. Koordynator odpowiada za terminowość oceny i programu, kompletność dokumentacji '
        'oraz wypełnienie karty kontrolnej ośmiu elementów IPET przed podpisaniem programu.',
        '**§ 7. Terminy.** Program opracowuje się do 30 września albo w terminie 30 dni od dnia złożenia orzeczenia '
        'w szkole. Wielospecjalistycznej oceny dokonuje się co najmniej dwa razy w roku szkolnym — we wrześniu '
        'i w styczniu; zaleca się trzecią ocenę w maju. Przeglądy wskaźników przeprowadza się w listopadzie i marcu.',
        '**§ 8. Opinia dla poradni.** Prośbę przewodniczącego zespołu orzekającego o wydanie opinii o funkcjonowaniu '
        'ucznia rejestruje sekretariat w dniu wpływu i niezwłocznie przekazuje dyrektorowi oraz koordynatorowi. '
        'Termin 10 dni liczy się od dnia otrzymania prośby przez dyrektora. Projekt opinii koordynator przedkłada '
        'dyrektorowi nie później niż w 7. dniu terminu.',
        '**§ 9. Dostosowania.** Dla każdego ucznia objętego wsparciem sporządza się karty dostosowań przedmiotowych '
        'i przekazuje je nauczycielom przedmiotów w terminie do …… września, za potwierdzeniem odbioru.',
        '**[[§ 9a. Nauczyciel współorganizujący kształcenie.]]** [[Dyrektor zapewnia zatrudnienie nauczyciela '
        'współorganizującego kształcenie dla uczniów posiadających orzeczenie wydane ze względu na autyzm, w tym '
        'zespół Aspergera, oraz niepełnosprawności sprzężone (§ 7 ust. 2 rozp. t.j. Dz.U. 2020 poz. 1309). '
        'W pozostałych przypadkach, w tym przy niedostosowaniu społecznym i zagrożeniu niedostosowaniem, dyrektor '
        'występuje do organu prowadzącego o zgodę na dodatkowe zatrudnienie (§ 7 ust. 3) i dokumentuje wystąpienie '
        'oraz rozstrzygnięcie w dokumentacji ucznia.]]',
        '**§ 10. Prawa rodziców.** O terminie spotkania zespołu zawiadamia się rodziców …… (wskazać sposób przyjęty '
        'w szkole) nie później niż …… dni przed spotkaniem. Przekazanie rodzicom kopii wielospecjalistycznej oceny '
        'i kopii programu odnotowuje się w rejestrze kontaktów, z datą i podpisem.',
        '**§ 11. Strażnik Prawa.** Ustanawia się rotacyjną funkcję Strażnika Prawa. Osoba pełniąca funkcję sprawdza '
        'aktualność publikatorów w drukach przed każdym cyklem dokumentacyjnym i przy każdej decyzji zespołu zadaje '
        'pytanie o podstawę prawną. Funkcję pełni się przez …… , zgodnie z harmonogramem stanowiącym załącznik nr 2. '
        '[[Strażnik Prawa prowadzi rejestr przepisów cytowanych w drukach szkoły (załącznik F do skryptu): '
        'dla każdego aktu odnotowuje tytuł, obowiązujący publikator, status, źródło weryfikacji i datę sprawdzenia. '
        'Druk, przy którym w rejestrze brakuje daty sprawdzenia, nie jest dopuszczony do obiegu.]]',
        '**§ 12. Audyt roczny.** W sierpniu przeprowadza się audyt dokumentacji kształcenia specjalnego '
        '(druk AUD-SP). Wyniki audytu dyrektor przedstawia radzie pedagogicznej na zebraniu przed rozpoczęciem '
        'roku szkolnego.',
        '**§ 13. Przepisy przejściowe.** Dokumentacja sporządzona przed dniem wejścia w życie zarządzenia zachowuje '
        'ważność. Ocenę i program dostosowuje się do nowego układu przy najbliższej wielospecjalistycznej ocenie. '
        'Dotychczasowe zapisy pozostają w teczce ucznia jako historia wsparcia.',
        '**§ 14. Wejście w życie.** Zarządzenie wchodzi w życie z dniem 1 września 2026 r.',
    ])
    para(doc, 'Załączniki do zarządzenia: nr 1 — procedura rozpoznawania potrzeb i reguły przekierowania; '
              'nr 2 — harmonogram funkcji Strażnika Prawa; nr 3 — wykaz druków (18 pozycji) obowiązujących w szkole; '
              'nr 4 — kalendarz dokumentacji na rok szkolny.', size=8.5, italic=True, color=GREY, before=6)

    # -------------------------------------------------------------- ZAŁĄCZNIK E
    page_break(doc)
    band(doc, 'ZAŁĄCZNIK E', 'Checklista wdrożenia — 30 / 60 / 90 dni')
    para(doc, 'Harmonogram dla dyrektora i koordynatora. Wdrożenie nie polega na przepisaniu dokumentacji, '
              'lecz na uruchomieniu obiegu, w którym każdy druk zasila następny.', after=8)

    h2(doc, 'Pierwsze 30 dni — porządkowanie podstaw', before=4)
    table(doc,
          ['#', 'Zadanie', 'Kto', 'Dowód wykonania'],
          [['1', 'Audyt wzorów druków — sprawdzenie publikatorów i tekstów jednolitych. ⚑',
            'Strażnik Prawa', 'Arkusz audytu (druk 17), pozycje 1 i 15'],
           ['2', 'Szkolenie rady pedagogicznej — siedem części skryptu.', 'dyrektor', 'protokół rady, lista obecności'],
           ['3', 'Wydanie zarządzenia dyrektora wraz z czterema załącznikami.', 'dyrektor', 'zarządzenie nr ……'],
           ['4', 'Wyznaczenie koordynatorów zespołów i Strażnika Prawa.', 'dyrektor', 'załącznik nr 2 do zarządzenia'],
           ['5', 'Założenie teczek uczniów i uzupełnienie metryczek, w tym dat wpływu orzeczeń.',
            'wychowawcy', 'metryczki w teczkach'],
           ['6', 'Powielenie i rozdanie druków 1–18; przeszkolenie z ich wypełniania.',
            'koordynator', 'potwierdzenia odbioru'],
           ['7', 'Wypełnienie KSzOF dla wszystkich uczniów oddziałów — pomiar bazowy.',
            'wychowawcy + rodzice', 'arkusze u koordynatora'],
           ['8', 'Opracowanie IPET dla uczniów rozpoczynających kształcenie z orzeczeniem — do 30 września.',
            'zespoły', 'programy z kartą kontrolną ośmiu elementów'],
           ['9', 'Sprawdzenie obsady nauczyciela współorganizującego tam, gdzie jest obowiązkowa.',
            'dyrektor', 'arkusz organizacyjny']],
          widths=[0.9, 8.4, 3.2, 4.2], size=8)

    h2(doc, 'Dni 31–60 — uruchomienie obserwacji pogłębionej i dostosowań')
    table(doc,
          ['#', 'Zadanie', 'Kto', 'Dowód wykonania'],
          [['10', 'Analiza profili KSzOF; zastosowanie siedmiu reguł przekierowania.', 'zespoły', 'karty decyzyjne (druk 1)'],
           ['11', 'Uruchomienie obserwacji pogłębionej tam, gdzie reguła zadziałała.',
            'wyznaczeni obserwatorzy', 'druki 2–7 w toku'],
           ['12', 'Sporządzenie kart dostosowań przedmiotowych i przekazanie ich nauczycielom.',
            'koordynator', 'potwierdzenia odbioru (druk 11)'],
           ['13', 'Zebranie arkuszy informacji od nauczycieli przedmiotów (klasy IV–VIII).',
            'wychowawcy', 'druki 10 w teczkach'],
           ['14', 'Sporządzenie planów wsparcia (PWES) dla uczniów bez orzeczenia objętych pomocą.',
            'zespoły', 'druki 15 w teczkach'],
           ['15', 'Wypełnienie ankiet „Mój głos” z uczniami objętymi wsparciem.',
            'wychowawcy / psycholog', 'druki 8 w teczkach'],
           ['16', 'Przegląd wskaźników — listopad, 15 minut na ucznia.', 'koordynatorzy', 'karty celów (druk 13)'],
           ['17', 'Klasy VIII: założenie kart dostosowań egzaminacyjnych. ⚑ terminy z komunikatu CKE',
            'dyrektor', 'druki 12 + uchwała rady']],
          widths=[0.9, 8.4, 3.2, 4.2], size=8)

    h2(doc, 'Dni 61–90 — pierwsza pełna ewaluacja')
    table(doc,
          ['#', 'Zadanie', 'Kto', 'Dowód wykonania'],
          [['18', 'Druga wielospecjalistyczna ocena (styczeń) na nowym druku WOPF-SP z jawnymi źródłami.',
            'zespoły', 'oceny w teczkach'],
           ['19', 'Modyfikacja programów i planów po ocenie.', 'zespoły', 'zaktualizowane IPET / PWES'],
           ['20', 'Ocena efektywności udzielanej pomocy.', 'nauczyciele i specjaliści', 'druki 14 w teczkach'],
           ['21', 'Przekazanie rodzicom kopii ocen i programów — z potwierdzeniem.',
            'koordynatorzy', 'rejestr kontaktów (druk 18)'],
           ['22', 'Sprawdzenie kompletności teczek — powtórny audyt wybranych pozycji.',
            'Strażnik Prawa', 'druk 17 z datą kontroli'],
           ['23', 'Podsumowanie wdrożenia na radzie pedagogicznej: co działa, co poprawiamy.',
            'dyrektor', 'protokół rady'],
           ['24', 'Ustalenie terminów: przegląd marcowy, ocena majowa, karty przekazania w czerwcu.',
            'dyrektor', 'kalendarz dokumentacji']],
          widths=[0.9, 8.4, 3.2, 4.2], size=8)

    accent_box(doc, [
        'MIARA SUKCESU WDROŻENIA: w dowolnym dniu roku szkolnego, po otrzymaniu prośby z poradni, zespół jest '
        'w stanie sporządzić rzetelną opinię o funkcjonowaniu ucznia w ciągu dziesięciu dni — nie zbierając '
        'żadnych nowych danych, tylko składając opinię z druków, które już są w teczce.',
        'DRUGA MIARA: nauczyciel, który we wrześniu po raz pierwszy wchodzi do klasy z uczniem objętym wsparciem, '
        'dostaje jedną kartę dostosowań na swój przedmiot i wie z niej, co ma robić — bez czytania '
        'dwudziestostronicowego programu.',
    ], title='PO CZYM POZNAMY, ŻE SIĘ UDAŁO')

    page_break(doc)
    para(doc, '', after=40)
    para(doc, 'EduPlaner 2026 · PCTP', size=9, bold=True, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(doc, 'Skrypt dla nauczycieli — szkoła podstawowa', size=16, bold=True, color=PURPLE,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(doc, 'Kształcenie specjalne i pomoc psychologiczno-pedagogiczna · rok szkolny 2026/2027',
         size=10, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    hairline(doc)
    para(doc, 'Każdy druk ma swój przepis. Obserwacja wyprzedza pismo z poradni. '
              'Cel ma liczbę, a ewaluacja ma konsekwencję.',
         size=11, italic=True, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    hairline(doc)
    para(doc, 'Materiał szkoleniowy. Przed wpisaniem podstawy prawnej do dokumentu ucznia sprawdź publikator '
              'w Internetowym Systemie Aktów Prawnych — dotyczy w szczególności pozycji oznaczonych znakiem ⚑ '
              'oraz terminów i katalogu dostosowań ogłaszanych corocznie w komunikacie dyrektora CKE.',
         size=8, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)


# ================================================================ MAIN

def main():
    doc, sec = build()
    raport_straznika(doc)
    czesc_1(doc)
    czesc_2(doc)
    czesc_3(doc)
    czesc_4(doc)
    czesc_5(doc)
    czesc_6(doc)
    czesc_7(doc)
    zal_A_B(doc)
    zal_C_1(doc)
    zal_C_2(doc)
    zal_C_3(doc)
    zal_C_4(doc)
    zal_D_E(doc)
    zal_F(doc)
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       'Skrypt_dla_nauczycieli_SZKOLA_PODSTAWOWA_EduPlaner_2026.docx')
    doc.save(out)
    print('Zapisano:', out)
    return out



# ================================================ RAPORT STRAŻNIKA PRAWA

def raport_straznika(doc):
    page_break(doc)
    band(doc, 'AUDYT', 'Raport Strażnika Prawa — weryfikacja podstaw prawnych wydania 1', '05.09.2026')

    para(doc, 'Wydanie 1 skryptu zostało poddane audytowi podstaw prawnych. Sprawdzono każdy publikator i każdą '
              'tezę prawną w źródłach: w rejestrze przepisów placówki (stan weryfikacji ELI z 21 sierpnia 2026 r.) '
              'oraz w ogłoszonych tekstach aktów. Poniżej wynik — łącznie dwanaście ustaleń. Fragmenty zmienione '
              'w dokumencie oznaczono kolorem niebieskim.', after=8)

    accent_box(doc, [
        'Audyt potwierdził pięć błędów wydania 1. Cztery z nich to naruszenie zasady, którą sam skrypt ogłasza: '
        'cytowanie publikatora pierwotnego zamiast obowiązującego tekstu jednolitego. Piąty to błąd merytoryczny '
        'w opisie obowiązku zatrudnienia nauczyciela współorganizującego.',
        'Audyt odrzucił dwa zarzuty postawione skryptowi. Rozporządzenie ME z 2 marca 2026 r. (Dz.U. 2026 poz. 428) '
        'istnieje, obowiązuje od 14 kwietnia 2026 r. i z tym dniem uchyliło rozporządzenie z 7 września 2017 r.; '
        'dziesięciodniowy termin na opinię jest w nim zapisany. Prawo oświatowe ma tekst jednolity Dz.U. 2026 '
        'poz. 820 — wskazywany w zarzucie Dz.U. 2024 poz. 737 wygasł 31 lipca 2025 r.',
        'Audyt wykrył pięć dalszych rozbieżności, których w zgłoszeniu nie było — w tym brak podstawy programowej '
        'obowiązującej od 1 września 2026 r. i nieprecyzyjne przypisanie obowiązku finansowego.',
        'Zasada na przyszłość: przepis wpisany do druku bez daty sprawdzenia jest przepisem niesprawdzonym. '
        'Rejestr w załączniku F zamyka tę lukę.',
    ], title='WYNIK W CZTERECH ZDANIACH')

    h2(doc, 'A. Błędy potwierdzone — poprawione w wydaniu 2')
    table(doc,
          ['#', 'Gdzie', 'Było w wydaniu 1', 'Jest w wydaniu 2', 'Na jakiej podstawie'],
          [['A1', 'Cz. 1 i 7, tabela dwóch ścieżek, zał. D',
            'W szkole ogólnodostępnej nauczyciela współorganizującego zatrudnia się dla ucznia z autyzmem, '
            'zespołem Aspergera, niepełnosprawnościami sprzężonymi ORAZ niedostosowanego społecznie — „to obowiązek”.',
            'Obowiązek dotyczy wyłącznie autyzmu (w tym zespołu Aspergera) i niepełnosprawności sprzężonych. '
            'Przy niedostosowaniu społecznym i pozostałych niepełnosprawnościach — tylko za zgodą organu prowadzącego. '
            'Dodano § 9a do wzoru zarządzenia.',
            '§ 7 ust. 2 („zatrudnia się dodatkowo”) i § 7 ust. 3 („za zgodą organu prowadzącego można zatrudniać”) '
            'rozporządzenia MEN z 9.08.2017 r., t.j. Dz.U. 2020 poz. 1309.'],
           ['A2', 'Cz. 1 i 2 — podstawa prawna, narracja, tabela powodów',
            'Ustawa o finansowaniu zadań oświatowych, „art. 8 ust. 16–17: środki naliczone na kształcenie specjalne '
            'wydatkuje się na zadania związane z organizacją tego kształcenia”.',
            'Art. 8 ust. 1: obowiązek przeznaczenia środków nie niższych niż kwota z podziału części oświatowej '
            'subwencji ogólnej spoczywa na jednostce samorządu terytorialnego, nie na szkole. Argument '
            'przeformułowano na pośredni: dokumentacja jest źródłem danych do planowania i rozliczenia.',
            'Ustawa z 27.10.2017 r. o finansowaniu zadań oświatowych, art. 8 ust. 1; t.j. Dz.U. 2026 poz. 650 '
            '(obwieszczenie z 12.05.2026 r.).'],
           ['A3', 'Cz. 1 i 7, druki 11 i 12, tabela „było → ma być”',
            'Rozporządzenie o ocenianiu — „Dz.U. 2019 poz. 373 ze zm.”, czyli publikator pierwotny.',
            'Tekst jednolity: Dz.U. 2023 poz. 2572, z późn. zm. (m.in. Dz.U. 2025 poz. 778).',
            'Obwieszczenie MEiN z 10.11.2023 r. — t.j. rozporządzenia MEN z 22.02.2019 r.'],
           ['A4', 'Cz. 1 — podstawa prawna, druk 15',
            'Rozporządzenie o nauczaniu indywidualnym — „Dz.U. 2017 poz. 1616 ze zm.”, czyli publikator pierwotny.',
            'Tekst jednolity: Dz.U. 2023 poz. 2468, z późn. zm. (zm. Dz.U. 2024 poz. 1714).',
            'Obwieszczenie MEiN z 27.10.2023 r. Uwaga: wskazana w zgłoszeniu pozycja 2724 dotyczy innego aktu.'],
           ['A5', 'Cz. 2 i 7, druk 15, spis druków',
            'Plan wsparcia edukacyjnego ucznia (PWES) przedstawiony jako druk o podstawie prawnej w rozporządzeniu '
            'o pomocy psychologiczno-pedagogicznej.',
            'PWES opisany jako narzędzie wewnętrzne szkoły, wprowadzane zarządzeniem dyrektora — tak jak metryczka. '
            'W stopce druku nie powołujemy się na rozporządzenie. Dodano akapit narracji i ramkę „status druku”.',
            'Rozporządzenie o pomocy pp (t.j. Dz.U. 2023 poz. 1798) nie przewiduje takiego dokumentu: formy, okres '
            'i wymiar pomocy ustala dyrektor, zajęcia dokumentuje się w dziennikach, a przepis wymaga oceny '
            'efektywności — nie planu.']],
          widths=[0.9, 3.3, 4.4, 4.4, 3.7], size=7.5)

    h2(doc, 'B. Zarzuty odrzucone — zapis wydania 1 był prawidłowy')
    table(doc,
          ['#', 'Zarzut', 'Ustalenie audytu', 'Źródło'],
          [['B1', 'Rozporządzenie ME z 2 marca 2026 r. (Dz.U. 2026 poz. 428) jest fikcyjne / projektowe; jedynym '
            'obowiązującym aktem pozostaje rozporządzenie MEN z 7 września 2017 r. (t.j. Dz.U. 2023 poz. 2061).',
            'Zarzut nietrafny — i jego przyjęcie wprowadziłoby błąd. Rozporządzenie zostało ogłoszone w Dz.U. 2026 '
            'poz. 428, obowiązuje od 14 kwietnia 2026 r. i z tym dniem uchyliło rozporządzenie z 7 września 2017 r. '
            'Cytowanie t.j. Dz.U. 2023 poz. 2061 jako obowiązującego byłoby dziś błędem. Zapis wydania 1 '
            'pozostaje bez zmian; usunięto natomiast znak ⚑ jako niepotrzebny.',
            'Rejestr przepisów placówki, wpis „dz-u-2026-428”, status: obowiązuje, weryfikacja ELI 21.08.2026; '
            'ogłoszenie w Dz.U. 2026 poz. 428; komunikaty kuratoriów oświaty.'],
           ['B2', 'Dziesięciodniowy termin na opinię szkoły o funkcjonowaniu ucznia nie istnieje w obowiązującym '
            'prawie i pochodzi z projektów reformy.',
            'Zarzut nietrafny. Termin jest zapisany w § 7 ust. 3 rozporządzenia: opinię wydaje się w terminie 10 dni '
            'od dnia otrzymania przez dyrektora prośby o jej wydanie, a kopię otrzymują rodzice. Wymóg odniesienia '
            'opisu do aktywności i uczestniczenia w rozumieniu ICF zawiera § 7 ust. 7, wchodzący 1 września 2026 r.',
            'Tekst rozporządzenia Dz.U. 2026 poz. 428, § 7 ust. 2–3 i ust. 6–7 oraz § 8; § 33 określa zakres opinii '
            'do 31 sierpnia 2026 r.'],
           ['B3', 'Prawo oświatowe należy cytować jako t.j. Dz.U. 2024 poz. 737.',
            'Zarzut nietrafny i wewnętrznie sprzeczny z zasadą, na którą się powołuje. Tekst jednolity Dz.U. 2024 '
            'poz. 737 wygasł 31 lipca 2025 r., zastąpił go Dz.U. 2025 poz. 1043 (wygasł 21 czerwca 2026 r.), '
            'a obowiązujący to Dz.U. 2026 poz. 820, z późn. zm. Wydanie 1 cytowało prawidłowo; w wydaniu 2 '
            'publikator uzupełniono tam, gdzie go brakowało.',
            'Rejestr przepisów placówki, wpisy „dz-u-2024-737” (uchylony), „dz-u-2025-1043” (uchylony) '
            'i „dz-u-2026-820” (tekst jednolity), weryfikacja ELI 21.08.2026.']],
          widths=[0.9, 4.6, 6.2, 5.0], size=7.5)

    h2(doc, 'C. Ustalenia własne audytu — nieobjęte zgłoszeniem')
    table(doc,
          ['#', 'Ustalenie', 'Co zmieniono w wydaniu 2'],
          [['C1', 'Wydanie 1 nie wskazywało podstawy programowej obowiązującej od 1 września 2026 r. i odsyłało '
            'do sprawdzenia publikatora. To luka w dokumencie, który uczy powoływania podstaw prawnych.',
            'Wpisano rozporządzenie ME z 11 marca 2026 r. (Dz.U. 2026 poz. 378, zm. Dz.U. 2026 poz. 958) '
            'z harmonogramem wdrażania: od 1.09.2026 klasy I i IV, w kolejnych latach następne klasy; dla uczniów '
            'z niepełnosprawnością intelektualną w stopniu umiarkowanym lub znacznym — od 1.09.2026. Dodano '
            'planszę narracji w części 1.'],
           ['C2', 'Skutek praktyczny wdrażania etapami: uczennica z przykładu (klasa III w roku 2026/2027) '
            'nie jest objęta nową podstawą programową.',
            'W narracji dopisano wprost, że w jej programie powołujemy się na podstawę z Dz.U. 2017 poz. 356, '
            'z późn. zm., a nie na nową.'],
           ['C3', 'Rozporządzenie o organizacji publicznych szkół cytowano bez zmian po tekście jednolitym.',
            'Uzupełniono: t.j. Dz.U. 2023 poz. 2736, z późn. zm. (Dz.U. 2025 poz. 849, Dz.U. 2026 poz. 130, '
            'a od 1.09.2026 Dz.U. 2026 poz. 1090).'],
           ['C4', 'Podstawa nadzoru wskazana ogólnikowo („art. 55 i art. 60”).',
            'Zastąpiono: art. 55 oraz art. 68 ust. 1 Prawa oświatowego (t.j. Dz.U. 2026 poz. 820) i rozporządzenie '
            'MEN z 25.08.2017 r. w sprawie nadzoru pedagogicznego (t.j. Dz.U. 2024 poz. 15).'],
           ['C6', 'Numeracja ustępów § 6 rozporządzenia o kształceniu specjalnym (ocena okresowa, prawa '
            'rodziców) nie jest potwierdzona w ogłoszonym tekście — ustalenie wspólne z audytem skryptu '
            'przedszkolnego (A1).',
            'W całym dokumencie cytujemy „§ 6 rozporządzenia” i opisujemy obowiązek słowami. Numer ustępu '
            'wpisujemy dopiero po sprawdzeniu w ogłoszonym tekście jednolitym.'],
           ['C5', 'Brak mechanizmu, który utrwala wynik weryfikacji — bez niego następny audyt zaczyna od zera.',
            'Dodano załącznik F (rejestr przepisów z datą sprawdzenia) oraz obowiązek jego prowadzenia przez '
            'Strażnika Prawa w § 11 wzoru zarządzenia.']],
          widths=[0.9, 7.4, 8.4], size=7.5)

    accent_box(doc, [
        'Czego audyt NIE obejmował: treści merytorycznych narzędzia KSzOF (normy stenowe i twierdzenia pochodzą '
        'z opracowania Z. Gajdzicy, E. Widawskiej, S. Byry i in. z 2024 r.), progów przeliczeniowych i reguł '
        'przekierowania — te pozostają decyzją rady pedagogicznej i skrypt tak je opisuje, oraz katalogu dostosowań '
        'egzaminacyjnych, który ogłasza corocznie dyrektor CKE.',
        'Czego audyt nie może zastąpić: sprawdzenia publikatora w dniu wpisania go do dokumentu ucznia. '
        'Data weryfikacji w załączniku F to 5 września 2026 r. — po tej dacie odpowiada za nią Strażnik Prawa.',
    ], title='GRANICE AUDYTU', accent=ORANGE)


# ================================================ ZAŁĄCZNIK F — REJESTR

def zal_F(doc):
    page_break(doc)
    band(doc, 'ZAŁĄCZNIK F', 'Rejestr przepisów cytowanych w skrypcie — status i data weryfikacji')
    para(doc, 'Rejestr prowadzi Strażnik Prawa. Zasada jest jedna: druk, przy którym w rejestrze brakuje daty '
              'sprawdzenia, nie jest dopuszczony do obiegu. Przy każdym akcie podajemy zapis kanoniczny — dokładnie '
              'w tej postaci wpisujemy go do dokumentu ucznia. Stan na 5 września 2026 r.', after=8)
    table(doc,
          ['Akt i zapis kanoniczny', 'Status', 'Gdzie w skrypcie', 'Uwaga'],
          [['Ustawa z 14.12.2016 r. — Prawo oświatowe (t.j. Dz.U. 2026 poz. 820, z późn. zm.)',
            'obowiązuje', 'Cz. 1, 2, 7; zał. D; druki 12, 17',
            'Teksty jednolite Dz.U. 2024 poz. 737 i Dz.U. 2025 poz. 1043 WYGASŁY — nie cytować.'],
           ['Rozporządzenie MEN z 9.08.2017 r. w sprawie warunków organizowania kształcenia (…) '
            '(t.j. Dz.U. 2020 poz. 1309)',
            'obowiązuje', 'Cz. 1–7; druki 11, 13, 14, 17, 18; zał. D',
            'Publikator pierwotny Dz.U. 2017 poz. 1578 — nie cytować. § 6 ust. 1 pkt 1–8 i ust. 4; dalsze ustępy '
            '§ 6 (ocena okresowa, prawa rodziców) — numeracja do sprawdzenia w ogłoszonym tekście; '
            '§ 5; § 7 ust. 2 i 3.'],
           ['Rozporządzenie MEN z 9.08.2017 r. w sprawie zasad organizacji i udzielania pomocy '
            'psychologiczno-pedagogicznej (t.j. Dz.U. 2023 poz. 1798)',
            'obowiązuje', 'Cz. 1, 3, 5, 6, 7; druki 2, 3, 7, 14, 15',
            'Poprzedni t.j. Dz.U. 2020 poz. 1280 wygasł 5.09.2023. § 12 — zindywidualizowana ścieżka kształcenia.'],
           ['Rozporządzenie ME z 2.03.2026 r. w sprawie orzeczeń i opinii wydawanych przez zespoły orzekające '
            '(Dz.U. 2026 poz. 428)',
            'obowiązuje', 'Cz. 1, 2, 3, 5, 7',
            'Obowiązuje od 14.04.2026; uchyliło rozp. z 7.09.2017 r. (t.j. Dz.U. 2023 poz. 2061). '
            '§ 7 ust. 6–7 i § 8 wchodzą 1.09.2026; § 33 — zakres opinii do 31.08.2026.'],
           ['Rozporządzenie MEN z 25.08.2017 r. w sprawie sposobu prowadzenia dokumentacji przebiegu nauczania (…) '
            '(t.j. Dz.U. 2024 poz. 50)',
            'obowiązuje', 'Cz. 1, 3, 4, 5, 6; stopki druków obserwacyjnych',
            'Nie mylić z Dz.U. 2024 poz. 1640 — to akt o zupełnie innej treści; rozjazd odnotowany w rejestrze '
            'placówki.'],
           ['Rozporządzenie MEN z 22.02.2019 r. w sprawie oceniania, klasyfikowania i promowania uczniów '
            'i słuchaczy w szkołach publicznych (t.j. Dz.U. 2023 poz. 2572, z późn. zm.)',
            'obowiązuje', 'Cz. 1, 2, 7; druki 11, 12',
            'Publikator pierwotny Dz.U. 2019 poz. 373 — nie cytować. Zmiana m.in. Dz.U. 2025 poz. 778.'],
           ['Rozporządzenie MEN z 9.08.2017 r. w sprawie indywidualnego obowiązkowego rocznego przygotowania '
            'przedszkolnego dzieci i indywidualnego nauczania dzieci i młodzieży '
            '(t.j. Dz.U. 2023 poz. 2468, z późn. zm.)',
            'obowiązuje', 'Cz. 1; druk 15',
            'Publikator pierwotny Dz.U. 2017 poz. 1616 — nie cytować. Zmiana Dz.U. 2024 poz. 1714.'],
           ['Rozporządzenie MEN z 28.02.2019 r. w sprawie szczegółowej organizacji publicznych szkół i publicznych '
            'przedszkoli (t.j. Dz.U. 2023 poz. 2736, z późn. zm.)',
            'obowiązuje', 'Cz. 1',
            'Po tekście jednolitym: Dz.U. 2025 poz. 849, Dz.U. 2026 poz. 130; Dz.U. 2026 poz. 1090 od 1.09.2026.'],
           ['Rozporządzenie ME z 11.03.2026 r. w sprawie podstawy programowej wychowania przedszkolnego oraz '
            'podstawy programowej kształcenia ogólnego dla szkoły podstawowej (…) (Dz.U. 2026 poz. 378)',
            'obowiązuje', 'Cz. 1, 7',
            'Zmiana Dz.U. 2026 poz. 958. Wdrażanie: 1.09.2026 klasy I i IV, w kolejnych latach następne klasy; '
            'uczniowie z NI w stopniu umiarkowanym lub znacznym — od 1.09.2026.'],
           ['Rozporządzenie MEN z 14.02.2017 r. w sprawie podstawy programowej (…) '
            '(Dz.U. 2017 poz. 356, z późn. zm.)',
            'stosowane\nprzejściowo', 'Cz. 1, 7',
            'Nadal właściwe dla klas nieobjętych jeszcze wdrożeniem nowej podstawy — m.in. dla klasy III '
            'w roku 2026/2027.'],
           ['Rozporządzenie MEN z 25.08.2017 r. w sprawie nadzoru pedagogicznego (t.j. Dz.U. 2024 poz. 15)',
            'obowiązuje', 'Cz. 2; druk 17', 'Nie mylić z rozporządzeniem o dokumentacji z tej samej daty.'],
           ['Ustawa z 27.10.2017 r. o finansowaniu zadań oświatowych (t.j. Dz.U. 2026 poz. 650)',
            'obowiązuje', 'Cz. 1, 2',
            'Art. 8 ust. 1 — obowiązek jednostki samorządu terytorialnego. Artykuł 8 nie zawiera ustępów 16–17. '
            'Wskazywany w zgłoszeniu t.j. Dz.U. 2024 poz. 754 jest nieaktualny.'],
           ['Ustawa z 19.07.2019 r. o zapewnianiu dostępności osobom ze szczególnymi potrzebami '
            '(t.j. Dz.U. 2024 poz. 1411)',
            'obowiązuje', 'Cz. 1, 2, 7', 'Tekst jednolity uwzględnia zmianę Dz.U. 2024 poz. 731.'],
           ['Rozporządzenie (UE) 2016/679 — RODO, art. 5 ust. 1 lit. c i art. 9',
            'obowiązuje', 'Cz. 1, 2, 4; druki 4, 9, 12', 'Dane o zdrowiu ucznia — dane szczególnej kategorii.'],
           ['Komunikat dyrektora Centralnej Komisji Egzaminacyjnej o dostosowaniach egzaminu ósmoklasisty',
            'ogłaszany\ncorocznie ⚑', 'Cz. 1, 2, 7; druk 12; zał. A',
            'Jedyne źródło katalogu dostosowań i terminów. Sprawdzać każdego roku — nie da się zweryfikować raz '
            'na zawsze.'],
           ['ICF — Międzynarodowa Klasyfikacja Funkcjonowania, Niepełnosprawności i Zdrowia (WHO, 2001)',
            'standard', 'Cz. 2, 5, 6, 7; druki 1–7',
            'Nie jest aktem prawa; obowiązek odniesienia opisu do aktywności i uczestniczenia wynika z § 7 ust. 7 '
            'rozp. Dz.U. 2026 poz. 428.'],
           ['Z. Gajdzica, E. Widawska, S. Byra i in. (2024), Kwestionariusz Szkolnej Oceny Funkcjonalnej '
            '(KSzOF-I-III)',
            'opracowanie\nnaukowe', 'Cz. 5; druk źródłowy',
            'Źródło twierdzeń, skali i norm stenowych. Nie jest aktem prawa i nie podlega weryfikacji w ISAP.']],
          widths=[5.6, 1.9, 3.6, 5.6], size=7.5)

    table(doc,
          ['Rejestr sprawdził', 'Data weryfikacji', 'Źródła', 'Termin kolejnego przeglądu'],
          [['Strażnik Prawa — ……………………', '5 września 2026 r.',
            'Rejestr przepisów placówki (weryfikacja ELI 21.08.2026); ogłoszone teksty aktów; obwieszczenia '
            'o tekstach jednolitych.',
            'przed każdym cyklem dokumentacyjnym, nie rzadziej niż raz na kwartał']],
          widths=[4.2, 3.0, 6.5, 3.0], size=8, align_first_bold=False)


if __name__ == '__main__':
    main()
